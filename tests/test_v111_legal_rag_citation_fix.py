from __future__ import annotations

import io

from docx import Document

from services.accounting_ai_enterprise import normalize_extracted_text_for_rag
from services.rag_storage_v101 import (
    _build_rag_answer_from_sources,
    _citation_document_title,
    _extract_document_intelligence,
    _split_compound_questions,
)
from services.rag_v66_v67 import read_docx_bytes


def _legal_docx_bytes() -> bytes:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "BỘ TÀI CHÍNH"
    table.cell(0, 1).text = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    table.cell(1, 0).text = "Số: 15/2026/TT-BTC"
    table.cell(1, 1).text = "Hà Nội, ngày 04 tháng 03 năm 2026"
    doc.add_paragraph("THÔNG TƯ")
    doc.add_paragraph("HƯỚNG DẪN NGUYÊN TẮC KẾ TOÁN TÀI SẢN MÃ HÓA")
    doc.add_paragraph("Căn cứ Luật Kế toán số 88/2015/QH13;")
    doc.add_paragraph("Điều 1. Phạm vi điều chỉnh")
    doc.add_paragraph("2. Thông tư này không áp dụng cho việc xác định nghĩa vụ thuế.")
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def test_docx_reader_preserves_header_table_before_cited_laws() -> None:
    text = read_docx_bytes(_legal_docx_bytes())
    assert text.index("15/2026/TT-BTC") < text.index("88/2015/QH13")


def test_document_intelligence_uses_own_number_and_issue_date() -> None:
    text = normalize_extracted_text_for_rag(read_docx_bytes(_legal_docx_bytes()))
    intel = _extract_document_intelligence("15_2026_TT-BTC.docx", "test", text, "accounting_law")
    assert intel["document_number"] == "15/2026/TT-BTC"
    assert intel["issue_date"] == "04/03/2026"


def test_citation_repairs_old_wrong_metadata_from_filename() -> None:
    source = {
        "title": "test",
        "metadata": {
            "filename": "15_2026_TT-BTC_696722(1).docx",
            "document_intelligence": {"document_number": "88/2015/QH13"},
        },
    }
    assert _citation_document_title(source) == "Thông tư 15/2026/TT-BTC"


def test_compound_admin_question_is_split_and_answered_per_article() -> None:
    question = (
        "Thông tư này áp dụng Thông tư này có dùng để xác định nghĩa vụ thuế không?"
        "cho những đối tượng nào?"
    )
    subquestions = _split_compound_questions(question)
    assert subquestions == [
        "Thông tư này áp dụng cho những đối tượng nào?",
        "Thông tư này có dùng để xác định nghĩa vụ thuế không?",
    ]

    metadata = {
        "filename": "15_2026_TT-BTC.docx",
        "document_intelligence": {"document_number": "15/2026/TT-BTC"},
    }
    sources = [
        {
            "matched_subquery": subquestions[0],
            "title": "test",
            "document_id": "doc-1",
            "chunk_id": "chunk-2",
            "chunk_no": 2,
            "heading": "Điều 2. Đối tượng áp dụng",
            "content": (
                "Điều 2. Đối tượng áp dụng\n"
                "Thông tư này áp dụng đối với tổ chức cung cấp dịch vụ tài sản mã hóa, "
                "tổ chức phát hành tài sản mã hóa và nhà đầu tư trong nước là tổ chức."
            ),
            "metadata": metadata,
        },
        {
            "matched_subquery": subquestions[1],
            "title": "test",
            "document_id": "doc-1",
            "chunk_id": "chunk-1",
            "chunk_no": 1,
            "heading": "Điều 1. Phạm vi điều chỉnh",
            "content": (
                "Điều 1. Phạm vi điều chỉnh\n"
                "2. Thông tư này không áp dụng cho việc xác định nghĩa vụ thuế của các tổ chức "
                "tham gia thị trường tài sản mã hóa với ngân sách nhà nước."
            ),
            "metadata": metadata,
        },
    ]
    result = _build_rag_answer_from_sources(
        question,
        sources,
        storage_label="Supabase RAG",
        answer_mode="short",
    )
    assert "Điều 2" in result["answer"]
    assert "không áp dụng cho việc xác định nghĩa vụ thuế" in result["answer"]
    assert all(c["document_title"] == "Thông tư 15/2026/TT-BTC" for c in result["citations"])


def test_normalizer_does_not_split_sentence_case_thong_tu() -> None:
    normalized = normalize_extracted_text_for_rag("THÔNG TƯ\nThông tư này áp dụng cho tổ chức.")
    assert "Thông tư\nnày" not in normalized
    assert "Thông tư này áp dụng" in normalized

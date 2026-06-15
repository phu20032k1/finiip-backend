"""Finiip V100 - Backend-only Admin RAG UI helpers.

This layer is intentionally small and offline-first. It gives the backend owner
an HTML admin console for uploading and managing official RAG knowledge without
building a separate frontend.

Design rule:
- Admin/owner uploads official knowledge into RAG.
- Normal users upload invoices/reports/files only for temporary processing.
- User uploads must not contaminate the official RAG knowledge base.
"""
from __future__ import annotations

import csv
import hashlib
import html
import io
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from services.accounting_ai_enterprise import (
    ROOT_DIR,
    add_uploaded_document as local_add_uploaded_document,
    answer_with_enterprise_rag as local_answer_with_enterprise_rag,
    audit,
    extract_text_from_bytes,
    load_store,
    normalize_extracted_text_for_rag,
    save_store,
    search_documents as local_search_documents,
    split_document_into_chunks,
)
from services.rag_storage_v101 import (
    V101_VERSION,
    SUPABASE_SCHEMA_SQL,
    add_uploaded_document_supabase,
    answer_with_supabase_rag,
    delete_document_supabase,
    get_document_supabase,
    list_documents_supabase,
    reindex_document_supabase,
    search_documents_supabase,
    evaluate_rag_test_cases_supabase,
    clear_supabase_chat_memory,
    supabase_config,
    supabase_is_active,
)

V100_VERSION = "v100_backend_rag_admin_ui"

ADMIN_RAG_SOURCE_TYPES = [
    "tax_legal",
    "accounting_law",
    "internal_process",
    "chart_of_accounts",
    "payroll_bhxh",
    "invoice_policy",
    "audit_policy",
    "knowledge",
]

ADMIN_RAG_ALLOWED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".xlsx",
    ".xlsm",
    ".html",
    ".htm",
}

MAX_ADMIN_UPLOAD_BYTES = int(os.getenv("FINIIP_ADMIN_RAG_MAX_MB", "30")) * 1024 * 1024

# V67: temporary file reader/output generator. These files are NOT inserted
# into the official RAG knowledge base unless the admin uploads them via the
# normal Upload & Index form.
ADMIN_FILE_OUTPUT_DIR = ROOT_DIR / "data" / "admin_file_outputs"
ADMIN_FILE_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

ADMIN_FILE_OUTPUT_FORMATS = ["md", "txt", "json", "csv", "docx", "xlsx"]
ADMIN_FILE_PROCESS_TASKS = [
    ("extract", "Đọc & trích xuất text"),
    ("summary", "Tóm tắt tài liệu"),
    ("qa", "Trả lời câu hỏi theo file"),
    ("accounting_review", "Review kế toán/kiểm soát"),
    ("questions", "Tạo câu hỏi ôn tập + đáp án"),
]


def admin_key_is_valid(admin_key: Optional[str]) -> bool:
    """Simple admin-key check for backend-only HTML forms.

    In local development, if neither FINIIP_ADMIN_KEY nor FINIIP_API_KEY is set,
    the UI is open so the developer can test quickly. In deployment, set one of
    those env vars and open /admin/rag-ui?key=YOUR_KEY.
    """
    expected = os.getenv("FINIIP_ADMIN_KEY") or os.getenv("FINIIP_API_KEY") or ""
    if not expected:
        return True
    return (admin_key or "") == expected


def _safe(value: Any) -> str:
    return html.escape("" if value is None else str(value), quote=True)


def _short(value: Any, max_len: int = 120) -> str:
    text = "" if value is None else str(value)
    text = text.replace("\n", " ").strip()
    return text if len(text) <= max_len else text[: max_len - 1] + "…"


def _chunk_page(chunk: Dict[str, Any]) -> Optional[int]:
    """Best-effort page detector for chunk cards."""
    if chunk.get("page"):
        try:
            return int(chunk.get("page"))
        except Exception:
            pass
    text = "\n".join([str(chunk.get("heading") or ""), str(chunk.get("content") or "")])
    match = re.search(r"---\s*page\s*(\d+)\s*---", text, flags=re.I)
    if match:
        try:
            return int(match.group(1))
        except Exception:
            return None
    return None


def _chunk_structure_label(chunk: Dict[str, Any]) -> str:
    text = normalize_extracted_text_for_rag("\n".join([str(chunk.get("heading") or ""), str(chunk.get("content") or "")]))
    for pattern in [r"\bĐiều\s+\d+[^\n\.]*\.?", r"\bKhoản\s+\d+[^\n\.]*\.?", r"\bPHỤ\s+LỤC\s+[IVXLC\d]+"]:
        match = re.search(pattern, text, flags=re.I)
        if match:
            return _short(match.group(0).strip(), 90)
    heading = str(chunk.get("heading") or "").strip()
    heading = normalize_extracted_text_for_rag(heading).split("\n", 1)[0]
    return _short(heading, 90) if heading else "không có heading"


def _display_chunk_text(text: Any, max_len: int = 6000) -> str:
    clean = normalize_extracted_text_for_rag(str(text or ""))
    # Page marker is metadata in the UI, so do not let it dominate the content block.
    clean = re.sub(r"^---\s*page\s*\d+\s*---\s*", "", clean, flags=re.I).strip()
    if len(clean) > max_len:
        clean = clean[:max_len].rstrip() + "\n\n… [đã rút gọn trong UI, dữ liệu gốc vẫn còn trong Supabase]"
    return clean


def _render_doc_quality_panel(doc: Dict[str, Any], chunks: List[Dict[str, Any]]) -> str:
    metadata = doc.get("metadata") or {}
    parser = metadata.get("parser") or "unknown"
    warnings = metadata.get("warnings") or []
    pages = sorted({p for p in (_chunk_page(c) for c in chunks) if p})
    char_count = int(doc.get("char_count") or sum(len(str(c.get("content") or "")) for c in chunks))
    avg_chars = int(char_count / max(1, len(chunks)))
    advice: List[str] = []
    if parser == "pypdf":
        advice.append("PDF text-layer đã đọc được bằng pypdf; nếu chữ vẫn dính, bấm Re-index để chạy bộ làm sạch spacing mới.")
    if warnings:
        advice.append("Có cảnh báo parser, nên kiểm tra lại file gốc hoặc upload bản PDF có text layer rõ hơn.")
    if avg_chars > 1800:
        advice.append("Chunk hơi dài; nên giảm kích thước chunk khi nâng cấp hybrid/rerank để câu trả lời chính xác hơn.")
    if len(chunks) > 120:
        advice.append("Tài liệu nhiều chunk; nên bật hybrid search + rerank cho câu hỏi dài.")
    if not advice:
        advice.append("Index ổn cho test RAG cơ bản.")
    warnings_html = "".join(f"<li>{_safe(w)}</li>" for w in warnings) or "<li>Không có cảnh báo parser.</li>"
    advice_html = "".join(f"<li>{_safe(a)}</li>" for a in advice)
    pages_text = f"{pages[0]}–{pages[-1]} ({len(pages)} trang phát hiện)" if pages else "chưa phát hiện page marker"
    return f"""
      <div class="doc-health">
        <div><b>Parser:</b> {_safe(parser)}</div>
        <div><b>Pages:</b> {_safe(pages_text)}</div>
        <div><b>Avg chunk:</b> {_safe(avg_chars)} ký tự</div>
        <details>
          <summary>Đánh giá index / khuyến nghị</summary>
          <ul>{advice_html}</ul>
          <p><b>Cảnh báo parser:</b></p>
          <ul>{warnings_html}</ul>
        </details>
      </div>
    """


def _render_chunk_items(chunks: List[Dict[str, Any]], limit: int = 80) -> str:
    cards: List[str] = []
    for idx, c in enumerate(chunks[:limit], 1):
        chunk_no = c.get("chunk_no") or idx
        page = _chunk_page(c)
        label = _chunk_structure_label(c)
        raw = str(c.get("content") or "")
        clean = _display_chunk_text(raw)
        words = c.get("word_count") or len(clean.split())
        chunk_id = f"chunk-{_safe(chunk_no)}-{idx}"
        search_text = " ".join([str(chunk_no), str(page or ""), label, clean]).lower()
        page_chip = f'<span class="chip">trang {page}</span>' if page else '<span class="chip muted-chip">không rõ trang</span>'
        cards.append(f"""
        <details class="chunk-card" data-chunk-text="{_safe(search_text)}">
          <summary>
            <b>Chunk {_safe(chunk_no)}</b>
            {page_chip}
            <span class="chip">{_safe(words)} từ</span>
            <span class="summary-title">{_safe(label)}</span>
          </summary>
          <div class="chunk-actions">
            <button type="button" class="copy-chunk" data-copy-target="{chunk_id}">Copy đoạn sạch</button>
          </div>
          <pre id="{chunk_id}" class="chunk-body">{_safe(clean)}</pre>
          <details class="raw-chunk">
            <summary>Raw gốc để debug parser</summary>
            <pre>{_safe(raw)}</pre>
          </details>
        </details>
        """)
    if len(chunks) > limit:
        cards.append(f'<p class="muted">Đang hiển thị {limit}/{len(chunks)} chunks đầu để giao diện nhẹ. Dùng Search chunk hoặc tăng limit trong code nếu cần xem toàn bộ.</p>')
    return "\n".join(cards)


def validate_admin_rag_upload(filename: str, content: bytes) -> None:
    suffix = Path(filename or "").suffix.lower()
    if suffix not in ADMIN_RAG_ALLOWED_EXTENSIONS:
        allowed = ", ".join(sorted(ADMIN_RAG_ALLOWED_EXTENSIONS))
        raise ValueError(f"File {suffix or '(không có đuôi)'} chưa được phép upload vào RAG admin. Cho phép: {allowed}")
    if len(content or b"") <= 0:
        raise ValueError("File upload đang rỗng.")
    if len(content) > MAX_ADMIN_UPLOAD_BYTES:
        mb = MAX_ADMIN_UPLOAD_BYTES // (1024 * 1024)
        raise ValueError(f"File quá lớn. Giới hạn hiện tại: {mb}MB. Có thể chỉnh FINIIP_ADMIN_RAG_MAX_MB.")


def upload_admin_rag_document(
    *,
    filename: str,
    content: bytes,
    workspace_id: str = "default",
    title: Optional[str] = None,
    source_type: str = "knowledge",
    note: Optional[str] = None,
) -> Dict[str, Any]:
    """Upload an official admin-managed RAG document."""
    validate_admin_rag_upload(filename, content)
    if source_type not in ADMIN_RAG_SOURCE_TYPES:
        source_type = "knowledge"
    metadata = {
        "document_scope": "global_knowledge",
        "uploaded_from": V100_VERSION,
        "can_train_ai": True,
        "can_use_for_global_rag": True,
        "admin_note": note or "",
        "uploaded_at": datetime.utcnow().replace(microsecond=0).isoformat() + "Z",
    }
    if supabase_is_active():
        result = add_uploaded_document_supabase(
            filename=filename,
        content=content,
        workspace_id=workspace_id or "default",
        title=title or filename,
        source_type=source_type,
            metadata=metadata,
        )
    else:
        result = local_add_uploaded_document(
            filename=filename,
            content=content,
            workspace_id=workspace_id or "default",
            title=title or filename,
            source_type=source_type,
            metadata=metadata,
        )
    audit("v100_admin_rag_upload", {
        "workspace_id": workspace_id,
        "document_id": result.get("document", {}).get("document_id"),
        "filename": filename,
        "source_type": source_type,
    })
    return result


def list_admin_rag_documents(workspace_id: Optional[str] = None, include_deleted: bool = False) -> Dict[str, Any]:
    if supabase_is_active():
        return list_documents_supabase(workspace_id=workspace_id, include_deleted=include_deleted)
    store = load_store()
    items: List[Dict[str, Any]] = []
    for doc in store.get("documents", {}).values():
        if workspace_id and doc.get("workspace_id") != workspace_id:
            continue
        if not include_deleted and doc.get("status") == "deleted":
            continue
        metadata = doc.get("metadata") or {}
        item = dict(doc)
        item["document_scope"] = metadata.get("document_scope", "knowledge")
        item["filename"] = metadata.get("filename") or metadata.get("saved_path") or doc.get("title")
        item["parser"] = metadata.get("parser")
        item["warnings"] = metadata.get("warnings") or []
        item["admin_note"] = metadata.get("admin_note") or ""
        items.append(item)
    items.sort(key=lambda d: d.get("updated_at") or d.get("created_at") or "", reverse=True)
    return {"version": V100_VERSION, "count": len(items), "items": items}


def get_admin_rag_document(document_id: str) -> Dict[str, Any]:
    if supabase_is_active():
        return get_document_supabase(document_id)
    store = load_store()
    doc = store.get("documents", {}).get(document_id)
    if not doc:
        raise KeyError(f"Không tìm thấy document_id={document_id}")
    chunks = [c for c in store.get("chunks", {}).values() if c.get("document_id") == document_id]
    chunks.sort(key=lambda c: c.get("chunk_no", 0))
    return {"document": doc, "chunks": chunks, "chunk_count": len(chunks)}


def delete_admin_rag_document(document_id: str, hard_delete: bool = False) -> Dict[str, Any]:
    """Remove a document from official RAG search.

    Default is soft-delete document metadata and remove chunks so it cannot be
    retrieved by RAG. Hard delete removes metadata too.
    """
    if supabase_is_active():
        return delete_document_supabase(document_id=document_id, hard_delete=hard_delete)
    store = load_store()
    if document_id not in store.get("documents", {}):
        raise KeyError(f"Không tìm thấy document_id={document_id}")
    chunk_ids = [cid for cid, c in store.get("chunks", {}).items() if c.get("document_id") == document_id]
    for cid in chunk_ids:
        store["chunks"].pop(cid, None)
    if hard_delete:
        store["documents"].pop(document_id, None)
    else:
        doc = store["documents"][document_id]
        doc["status"] = "deleted"
        doc["chunk_count"] = 0
        doc["updated_at"] = datetime.utcnow().replace(microsecond=0).isoformat() + "Z"
    save_store(store)
    audit("v100_admin_rag_delete", {"document_id": document_id, "hard_delete": hard_delete, "chunks_removed": len(chunk_ids)})
    return {"ok": True, "document_id": document_id, "hard_delete": hard_delete, "chunks_removed": len(chunk_ids)}


def reindex_admin_rag_document(document_id: str) -> Dict[str, Any]:
    """Rebuild chunks for one admin RAG document.

    If the original saved file still exists, extraction runs again. Otherwise,
    it rebuilds from existing chunks as a fallback.
    """
    if supabase_is_active():
        return reindex_document_supabase(document_id)
    store = load_store()
    doc = store.get("documents", {}).get(document_id)
    if not doc:
        raise KeyError(f"Không tìm thấy document_id={document_id}")
    metadata = doc.get("metadata") or {}
    saved_path = metadata.get("saved_path")
    content_text = ""
    extraction: Dict[str, Any] = {}

    if saved_path:
        path = ROOT_DIR / saved_path
        if path.exists():
            data = path.read_bytes()
            extraction = extract_text_from_bytes(metadata.get("filename") or doc.get("title") or path.name, data)
            content_text = extraction.get("text") or ""

    if not content_text:
        old_chunks = [c for c in store.get("chunks", {}).values() if c.get("document_id") == document_id]
        old_chunks.sort(key=lambda c: c.get("chunk_no", 0))
        content_text = "\n\n".join(c.get("content", "") for c in old_chunks).strip()

    if not content_text:
        raise ValueError("Không có nội dung để re-index. Hãy upload lại tài liệu.")

    old_chunk_ids = [cid for cid, c in store.get("chunks", {}).items() if c.get("document_id") == document_id]
    for cid in old_chunk_ids:
        store["chunks"].pop(cid, None)

    chunks = split_document_into_chunks(content_text)
    # Keep document_id stable.
    from services.accounting_ai_enterprise import _id, _now, _tokens  # local import: internal helper reuse

    for c in chunks:
        chunk_id = _id("chk")
        store["chunks"][chunk_id] = {
            "chunk_id": chunk_id,
            "document_id": document_id,
            "workspace_id": doc.get("workspace_id", "default"),
            "title": doc.get("title"),
            "source_type": doc.get("source_type"),
            "section": c.get("section"),
            "chunk_no": c.get("chunk_no"),
            "heading": c.get("heading"),
            "content": c.get("content"),
            "tokens": sorted(set(_tokens(c.get("content", ""))))[:500],
            "created_at": _now(),
        }
    doc["status"] = "active"
    doc["chunk_count"] = len(chunks)
    doc["char_count"] = len(content_text)
    doc["updated_at"] = _now()
    metadata["last_reindexed_at"] = _now()
    if extraction:
        metadata["parser"] = extraction.get("parser")
        metadata["warnings"] = extraction.get("warnings") or []
        metadata["sha256"] = extraction.get("sha256") or metadata.get("sha256")
    doc["metadata"] = metadata
    save_store(store)
    audit("v100_admin_rag_reindex", {"document_id": document_id, "chunks_added": len(chunks), "chunks_removed": len(old_chunk_ids)})
    return {"ok": True, "document_id": document_id, "chunks_added": len(chunks), "chunks_removed": len(old_chunk_ids), "extraction": {k: v for k, v in extraction.items() if k != "text"}}


def admin_rag_stats(workspace_id: Optional[str] = None) -> Dict[str, Any]:
    docs = list_admin_rag_documents(workspace_id=workspace_id)["items"]
    active_docs = [d for d in docs if d.get("status") == "active"]
    return {
        "version": V100_VERSION,
        "workspace_id": workspace_id or "all",
        "documents": len(active_docs),
        "chunks": sum(int(d.get("chunk_count") or 0) for d in active_docs),
        "chars": sum(int(d.get("char_count") or 0) for d in active_docs),
        "by_source_type": _count_by(active_docs, "source_type"),
    }


def _count_by(items: List[Dict[str, Any]], key: str) -> Dict[str, int]:
    out: Dict[str, int] = {}
    for item in items:
        value = str(item.get(key) or "unknown")
        out[value] = out.get(value, 0) + 1
    return out


def admin_rag_storage_status() -> Dict[str, Any]:
    status = supabase_config()
    status["v100_ui_version"] = V100_VERSION
    status["schema_sql_available"] = True
    return status


def search_documents(
    query: str,
    workspace_id: str = "default",
    source_types: Optional[List[str]] = None,
    limit: int = 6,
) -> Dict[str, Any]:
    if supabase_is_active():
        return search_documents_supabase(query=query, workspace_id=workspace_id, source_types=source_types, limit=limit)
    result = local_search_documents(query=query, workspace_id=workspace_id, source_types=source_types, limit=limit)
    result["storage_backend"] = "local"
    return result


def answer_admin_rag_question(question: str, workspace_id: str = "default", limit: int = 6, history: str = "", answer_mode: str = "auto", conversation_id: str = "admin", save_memory: bool = True) -> Dict[str, Any]:
    # V109 uses one conversational engine for both Supabase and local mode.
    # This keeps greeting/help, formula solving, source cards and follow-up
    # memory consistent across /ai/v106/chat, Admin UI and /api/v1/chat.
    return answer_with_supabase_rag(
        question=question,
        workspace_id=workspace_id,
        limit=limit,
        history=history,
        answer_mode=answer_mode,
        conversation_id=conversation_id,
        save_memory=save_memory,
    )



def parse_eval_cases(raw_text: str) -> List[Dict[str, str]]:
    """Parse V60 Test Center cases from admin textarea.

    Supported formats:
    - question => expected answer
    - question | expected answer | expected source
    - one question per line (expected answer empty)
    """
    cases: List[Dict[str, str]] = []
    for line in str(raw_text or "").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        expected_source = ""
        if "|" in line:
            parts = [p.strip() for p in line.split("|")]
            question = parts[0] if parts else ""
            expected_answer = parts[1] if len(parts) > 1 else ""
            expected_source = parts[2] if len(parts) > 2 else ""
        elif "=>" in line:
            question, expected_answer = [p.strip() for p in line.split("=>", 1)]
        else:
            question, expected_answer = line, ""
        if question:
            cases.append({"question": question, "expected_answer": expected_answer, "expected_source": expected_source})
    return cases


def run_admin_rag_eval(raw_cases: str, workspace_id: str = "default", answer_mode: str = "short") -> Dict[str, Any]:
    cases = parse_eval_cases(raw_cases)
    if not cases:
        raise ValueError("Chưa có test case. Nhập dạng: Câu hỏi => đáp án mong muốn")
    if supabase_is_active():
        return evaluate_rag_test_cases_supabase(cases, workspace_id=workspace_id, answer_mode=answer_mode)
    # Local fallback: still useful for UI, but does not persist eval result.
    items = []
    for case in cases[:50]:
        ans = answer_admin_rag_question(case["question"], workspace_id=workspace_id, limit=6, answer_mode=answer_mode, save_memory=False)
        actual = ans.get("answer") or ""
        expected = case.get("expected_answer") or ""
        passed = bool(expected and _safe(expected).lower() in _safe(actual).lower()) if expected else bool(actual)
        items.append({**case, "actual_answer": actual, "score": 100 if passed else 0, "passed": passed, "citations": ans.get("citations") or []})
    passed_count = sum(1 for item in items if item.get("passed"))
    return {"version": "v60_eval_center_local", "workspace_id": workspace_id, "count": len(items), "passed": passed_count, "failed": len(items) - passed_count, "avg_score": round(sum(float(i.get("score") or 0) for i in items) / max(1, len(items)), 2), "items": items}


def clear_admin_rag_memory(workspace_id: str = "default", conversation_id: str = "admin") -> Dict[str, Any]:
    if supabase_is_active():
        return clear_supabase_chat_memory(workspace_id=workspace_id, conversation_id=conversation_id)
    return {"ok": True, "storage_backend": "local", "workspace_id": workspace_id, "conversation_id": conversation_id}


def _render_doc_intelligence_panel(doc: Dict[str, Any]) -> str:
    meta = doc.get("metadata") or {}
    intel = meta.get("document_intelligence") or meta.get("legal_intelligence") or {}
    if not intel:
        return '<p class="muted">Chưa có Document Intelligence. Hãy Re-index tài liệu để tạo metadata V64.</p>'
    mods = intel.get("modified_documents") or []
    tags = intel.get("tags") or []
    return f"""
      <div class="doc-intel">
        <div><b>Loại:</b> {_safe(intel.get('document_type') or 'unknown')}</div>
        <div><b>Số văn bản:</b> {_safe(intel.get('document_number') or 'chưa nhận diện')}</div>
        <div><b>Ngày ban hành:</b> {_safe(intel.get('issue_date') or 'chưa rõ')}</div>
        <div><b>Hiệu lực:</b> {_safe(intel.get('effective_date') or 'chưa rõ')}</div>
        <div><b>Tags:</b> {_safe(', '.join(tags) if tags else 'chưa có')}</div>
        <div><b>Sửa đổi/thay thế:</b> {_safe(', '.join(mods) if mods else 'chưa phát hiện')}</div>
      </div>
    """


def _render_eval_result(eval_result: Optional[Dict[str, Any]]) -> str:
    if not eval_result:
        return ""
    rows = []
    for idx, item in enumerate(eval_result.get("items") or [], 1):
        badge = "✅ PASS" if item.get("passed") else "❌ FAIL"
        cites = item.get("citations") or []
        cite_text = "; ".join(f"[{c.get('index')}] {c.get('document_title') or c.get('title')} — {c.get('location')}" for c in cites[:3])
        rows.append(f"""
          <tr>
            <td>{idx}</td>
            <td>{_safe(item.get('question'))}</td>
            <td>{_safe(item.get('expected_answer'))}<br><span class="muted">{_safe(item.get('expected_source'))}</span></td>
            <td><b>{_safe(item.get('score'))}</b><br>{badge}</td>
            <td><pre>{_safe(_short(item.get('actual_answer'), 900))}</pre><div class="muted">{_safe(cite_text)}</div></td>
          </tr>
        """)
    return f"""
    <section class="panel">
      <h2>V60 Test Center Result</h2>
      <div class="stats">
        <div class="stat"><b>{_safe(eval_result.get('count'))}</b><br><span class="muted">Test cases</span></div>
        <div class="stat"><b>{_safe(eval_result.get('passed'))}</b><br><span class="muted">Passed</span></div>
        <div class="stat"><b>{_safe(eval_result.get('failed'))}</b><br><span class="muted">Failed</span></div>
        <div class="stat"><b>{_safe(eval_result.get('avg_score'))}</b><br><span class="muted">Avg score</span></div>
      </div>
      <table><thead><tr><th>#</th><th>Câu hỏi</th><th>Kỳ vọng</th><th>Điểm</th><th>AI trả lời</th></tr></thead><tbody>{''.join(rows)}</tbody></table>
    </section>
    """


def _detect_uploaded_file_profile(filename: str, text: str) -> Dict[str, Any]:
    """V67 best-effort document profile for temporary file processing."""
    clean = normalize_extracted_text_for_rag(text or "")
    lowered = clean.lower()
    suffix = Path(filename or "").suffix.lower().lstrip(".") or "unknown"
    doc_type = "general_document"
    if "thông tư" in lowered or re.search(r"\btt-btc\b", lowered, flags=re.I):
        doc_type = "circular"
    elif "nghị định" in lowered:
        doc_type = "decree"
    elif "hóa đơn" in lowered or "invoice" in lowered:
        doc_type = "invoice"
    elif "báo cáo tài chính" in lowered or "bảng cân đối" in lowered or "báo cáo tình hình tài chính" in lowered:
        doc_type = "financial_statement"
    elif "hợp đồng" in lowered:
        doc_type = "contract"
    elif suffix in {"xlsx", "xlsm", "csv"}:
        doc_type = "spreadsheet"
    doc_no = ""
    m = re.search(r"(?:Số\s*:\s*)?([0-9]{1,4}/[0-9]{4}/[A-ZĐ\-]+)", clean, flags=re.I)
    if m:
        doc_no = m.group(1)
    tags = []
    for key, tag in [
        ("thuế", "tax"), ("hóa đơn", "invoice"), ("báo cáo tài chính", "financial_statement"),
        ("hợp nhất", "consolidation"), ("khấu hao", "depreciation"), ("doanh thu", "revenue"),
        ("chi phí", "expense"), ("công ty con", "subsidiary"), ("kiểm toán", "audit"),
    ]:
        if key in lowered:
            tags.append(tag)
    return {"filename": filename, "extension": suffix, "document_type": doc_type, "document_number": doc_no, "tags": sorted(set(tags))}


def _first_meaningful_lines(text: str, limit: int = 12) -> List[str]:
    lines = []
    for line in normalize_extracted_text_for_rag(text).splitlines():
        line = line.strip()
        if len(line) < 3:
            continue
        if re.fullmatch(r"[-–—\s]+", line):
            continue
        lines.append(line)
        if len(lines) >= limit:
            break
    return lines


def _score_sentence_against_query(sentence: str, query: str) -> int:
    tokens = set(re.findall(r"[\wÀ-ỹĐđ]+", (query or "").lower()))
    stokens = set(re.findall(r"[\wÀ-ỹĐđ]+", (sentence or "").lower()))
    if not tokens or not stokens:
        return 0
    return len(tokens & stokens) * 10 + min(len(sentence), 220) // 80


def _select_relevant_passages(text: str, query: str = "", limit: int = 8) -> List[str]:
    clean = normalize_extracted_text_for_rag(text)
    # Split at legal paragraphs/sentences while keeping Vietnamese legal text readable.
    parts = re.split(r"(?<=[\.?!])\s+|\n{2,}", clean)
    scored = []
    for part in parts:
        part = part.strip()
        if len(part) < 40:
            continue
        score = _score_sentence_against_query(part, query)
        # Useful legal/accounting markers get a small boost.
        if re.search(r"\b(Điều|Khoản|Mã số|Tài khoản|Nợ|Có|Báo cáo|Thông tư|Nghị định)\b", part, flags=re.I):
            score += 3
        scored.append((score, part))
    scored.sort(key=lambda x: x[0], reverse=True)
    if query and scored and scored[0][0] > 0:
        return [p for _, p in scored[:limit]]
    return _first_meaningful_lines(clean, limit=limit)


def _make_summary_from_text(filename: str, text: str, instruction: str = "") -> Dict[str, Any]:
    profile = _detect_uploaded_file_profile(filename, text)
    lines = _first_meaningful_lines(text, limit=10)
    relevant = _select_relevant_passages(text, instruction, limit=6) if instruction else _select_relevant_passages(text, "", limit=6)
    return {
        "title": Path(filename or "upload").name,
        "profile": profile,
        "key_points": relevant,
        "opening_lines": lines,
        "char_count": len(text or ""),
        "word_count": len((text or "").split()),
    }


def _build_file_process_markdown(*, filename: str, task_type: str, question: str, instruction: str, extracted: Dict[str, Any], clean_text: str) -> str:
    summary = _make_summary_from_text(filename, clean_text, instruction or question)
    profile = summary["profile"]
    title = Path(filename or "upload").name
    header = [
        f"# Finiip V67 - Kết quả đọc file: {title}",
        "",
        f"- Loại xử lý: `{task_type}`",
        f"- Parser: `{extracted.get('parser')}`",
        f"- Số ký tự đọc được: {len(clean_text)}",
        f"- Loại tài liệu nhận diện: `{profile.get('document_type')}`",
        f"- Số văn bản nhận diện: {profile.get('document_number') or 'chưa phát hiện'}",
        f"- Tags: {', '.join(profile.get('tags') or []) or 'chưa có'}",
        "",
    ]
    if extracted.get("warnings"):
        header += ["## Cảnh báo parser", ""] + [f"- {w}" for w in extracted.get("warnings") or []] + [""]

    if task_type == "extract":
        body = ["## Text đã trích xuất", "", clean_text[:60000]]
    elif task_type == "qa":
        passages = _select_relevant_passages(clean_text, question or instruction, limit=10)
        body = ["## Câu hỏi", "", question or instruction or "(chưa nhập câu hỏi)", "", "## Trả lời nháp theo file", ""]
        if passages:
            body.append("Dựa trên các đoạn liên quan trong file, nội dung cần chú ý là:")
            body.append("")
            for idx, psg in enumerate(passages, 1):
                body.append(f"{idx}. {psg}")
        else:
            body.append("Chưa tìm thấy đoạn liên quan rõ ràng trong file.")
        body += ["", "## Gợi ý kiểm tra", "", "- Đối chiếu lại file gốc trước khi dùng kết quả.", "- Nếu là nghiệp vụ kế toán/thuế, cần kiểm tra chứng từ và chính sách công ty."]
    elif task_type == "accounting_review":
        passages = _select_relevant_passages(clean_text, instruction or question or "kế toán thuế báo cáo tài chính hóa đơn chi phí doanh thu", limit=10)
        body = [
            "## Review kế toán/kiểm soát", "",
            "### 1. Nhận diện nhanh", "",
            f"- Loại tài liệu: {profile.get('document_type')}",
            f"- Tags: {', '.join(profile.get('tags') or []) or 'chưa có'}", "",
            "### 2. Nội dung quan trọng đọc được", "",
        ]
        body += [f"- {p}" for p in passages]
        body += ["", "### 3. Việc kế toán nên kiểm tra", "", "- Tính đầy đủ của chứng từ gốc.", "- Kỳ kế toán/kỳ thuế áp dụng.", "- Tài khoản, chỉ tiêu báo cáo hoặc mã số liên quan.", "- Quy định nội bộ và phê duyệt của người phụ trách.", "", "### 4. Kết luận", "", "Kết quả này là bản đọc/tổng hợp tự động từ file, dùng để hỗ trợ soát xét trước khi xử lý chính thức."]
    elif task_type == "questions":
        passages = _select_relevant_passages(clean_text, instruction or question, limit=12)
        body = ["## Câu hỏi ôn tập / kiểm tra hiểu tài liệu", ""]
        for idx, p in enumerate(passages[:10], 1):
            short = _short(p, 180)
            body.append(f"### Câu {idx}")
            body.append(f"Hãy giải thích nội dung sau trong tài liệu: {short}")
            body.append("")
            body.append("**Đáp án gợi ý:**")
            body.append(p)
            body.append("")
    else:
        body = ["## Tóm tắt tài liệu", ""]
        for idx, p in enumerate(summary.get("key_points") or [], 1):
            body.append(f"{idx}. {p}")
        body += ["", "## Các dòng đầu tiên", ""] + [f"- {x}" for x in summary.get("opening_lines") or []]

    return "\n".join(header + body).strip() + "\n"


def _write_output_file(job_id: str, output_format: str, markdown_text: str, payload: Dict[str, Any]) -> Path:
    fmt = (output_format or "md").lower().strip(".")
    if fmt not in ADMIN_FILE_OUTPUT_FORMATS:
        fmt = "md"
    path = ADMIN_FILE_OUTPUT_DIR / f"{job_id}.{fmt}"
    if fmt == "md":
        path.write_text(markdown_text, encoding="utf-8")
    elif fmt == "txt":
        # Strip simple markdown markers enough for plain text.
        plain = re.sub(r"^#+\s*", "", markdown_text, flags=re.M)
        plain = plain.replace("**", "").replace("`", "")
        path.write_text(plain, encoding="utf-8")
    elif fmt == "json":
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    elif fmt == "csv":
        with path.open("w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["field", "value"])
            for key, value in payload.items():
                if isinstance(value, (dict, list)):
                    value = json.dumps(value, ensure_ascii=False)
                writer.writerow([key, value])
    elif fmt == "docx":
        try:
            from docx import Document  # type: ignore
            doc = Document()
            for line in markdown_text.splitlines():
                if line.startswith("# "):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:].strip(), level=3)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:].strip(), style="List Bullet")
                elif re.match(r"^\d+\.\s+", line):
                    doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
                elif line.strip():
                    doc.add_paragraph(line.strip().replace("**", "").replace("`", ""))
            doc.save(path)
        except Exception:
            fallback = path.with_suffix(".md")
            fallback.write_text(markdown_text, encoding="utf-8")
            path = fallback
    elif fmt == "xlsx":
        try:
            import openpyxl  # type: ignore
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Finiip result"
            ws.append(["Field", "Value"])
            for key, value in payload.items():
                if isinstance(value, (dict, list)):
                    value = json.dumps(value, ensure_ascii=False)
                ws.append([key, str(value)])
            ws2 = wb.create_sheet("Markdown")
            ws2.append(["Line"])
            for line in markdown_text.splitlines()[:5000]:
                ws2.append([line])
            wb.save(path)
        except Exception:
            fallback = path.with_suffix(".csv")
            with fallback.open("w", encoding="utf-8-sig", newline="") as f:
                writer = csv.writer(f)
                writer.writerow(["content"])
                for line in markdown_text.splitlines():
                    writer.writerow([line])
            path = fallback
    return path


def process_admin_file_to_output(
    *,
    filename: str,
    content: bytes,
    workspace_id: str = "default",
    task_type: str = "summary",
    output_format: str = "md",
    question: str = "",
    instruction: str = "",
) -> Dict[str, Any]:
    """V67: Read a temporary file and return a generated output file.

    This is intentionally separate from Upload & Index. It lets the admin test
    file reading, extract text, make a summary/review/questions, then download a
    result file without adding the upload to official RAG.
    """
    validate_admin_rag_upload(filename, content)
    if task_type not in {x[0] for x in ADMIN_FILE_PROCESS_TASKS}:
        task_type = "summary"
    fmt = (output_format or "md").lower().strip(".")
    if fmt not in ADMIN_FILE_OUTPUT_FORMATS:
        fmt = "md"
    extracted = extract_text_from_bytes(filename, content)
    clean_text = normalize_extracted_text_for_rag(extracted.get("text") or "")
    if not clean_text.strip():
        raise ValueError("Không đọc được text từ file này. Hãy thử PDF có text-layer, DOCX, TXT, CSV hoặc XLSX.")
    job_seed = f"{filename}|{workspace_id}|{task_type}|{fmt}|{datetime.utcnow().isoformat()}|{hashlib.sha256(content).hexdigest()[:12]}"
    job_id = "file_" + hashlib.sha1(job_seed.encode("utf-8")).hexdigest()[:16]
    markdown = _build_file_process_markdown(
        filename=filename,
        task_type=task_type,
        question=question,
        instruction=instruction,
        extracted=extracted,
        clean_text=clean_text,
    )
    profile = _detect_uploaded_file_profile(filename, clean_text)
    payload = {
        "version": "v67_file_reader_return_file",
        "job_id": job_id,
        "workspace_id": workspace_id or "default",
        "filename": filename,
        "task_type": task_type,
        "output_format": fmt,
        "question": question,
        "instruction": instruction,
        "parser": extracted.get("parser"),
        "warnings": extracted.get("warnings") or [],
        "profile": profile,
        "char_count": len(clean_text),
        "word_count": len(clean_text.split()),
        "preview": markdown[:2500],
    }
    path = _write_output_file(job_id, fmt, markdown, payload)
    payload["output_filename"] = path.name
    payload["output_path"] = str(path)
    payload["download_url"] = f"/admin/rag-ui/file/download?job_id={job_id}"
    meta_path = ADMIN_FILE_OUTPUT_DIR / f"{job_id}.json.meta"
    meta_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    audit("v67_admin_file_process", {"workspace_id": workspace_id, "job_id": job_id, "filename": filename, "task_type": task_type, "output_format": path.suffix.lstrip('.')})
    return payload


def resolve_admin_file_output(job_id: str) -> Dict[str, Any]:
    safe_job_id = re.sub(r"[^A-Za-z0-9_\-]", "", job_id or "")
    if not safe_job_id:
        raise FileNotFoundError("Thiếu job_id")
    meta_path = ADMIN_FILE_OUTPUT_DIR / f"{safe_job_id}.json.meta"
    if meta_path.exists():
        try:
            meta = json.loads(meta_path.read_text(encoding="utf-8"))
            out = Path(meta.get("output_path") or "")
            if out.exists() and ADMIN_FILE_OUTPUT_DIR in out.parents:
                return {"path": out, "filename": meta.get("output_filename") or out.name, "metadata": meta}
        except Exception:
            pass
    for ext in ADMIN_FILE_OUTPUT_FORMATS:
        path = ADMIN_FILE_OUTPUT_DIR / f"{safe_job_id}.{ext}"
        if path.exists():
            return {"path": path, "filename": path.name, "metadata": {"job_id": safe_job_id}}
    raise FileNotFoundError(f"Không tìm thấy file kết quả cho job_id={safe_job_id}")


def _render_file_process_result(file_result: Optional[Dict[str, Any]], admin_key: str = "") -> str:
    if not file_result:
        return ""
    dl = f"{file_result.get('download_url')}&key={_safe(admin_key)}" if admin_key else file_result.get("download_url")
    warn = "".join(f"<li>{_safe(w)}</li>" for w in (file_result.get("warnings") or [])) or "<li>Không có cảnh báo parser.</li>"
    return f"""
    <section class="panel result">
      <h2>V67 File Reader Result</h2>
      <p><b>File:</b> {_safe(file_result.get('filename'))} | <b>Task:</b> {_safe(file_result.get('task_type'))} | <b>Parser:</b> {_safe(file_result.get('parser'))}</p>
      <p><b>Ký tự:</b> {_safe(file_result.get('char_count'))} | <b>Từ:</b> {_safe(file_result.get('word_count'))} | <b>Output:</b> {_safe(file_result.get('output_filename'))}</p>
      <p><a href="{_safe(dl)}" target="_blank"><button type="button">Tải file kết quả</button></a></p>
      <details open><summary>Xem nhanh kết quả</summary><pre>{_safe(file_result.get('preview'))}</pre></details>
      <details><summary>Cảnh báo parser</summary><ul>{warn}</ul></details>
    </section>
    """

def render_admin_rag_page(
    *,
    admin_key: str = "",
    workspace_id: str = "default",
    message: str = "",
    error: str = "",
    question: str = "",
    answer_result: Optional[Dict[str, Any]] = None,
    search_query: str = "",
    search_result: Optional[Dict[str, Any]] = None,
    document_detail: Optional[Dict[str, Any]] = None,
    eval_result: Optional[Dict[str, Any]] = None,
    file_result: Optional[Dict[str, Any]] = None,
    eval_cases: str = "",
    answer_mode: str = "auto",
) -> str:
    """Render a plain white backend admin page."""
    docs = list_admin_rag_documents(workspace_id=workspace_id).get("items", [])
    stats = admin_rag_stats(workspace_id=workspace_id)
    storage_status = admin_rag_storage_status()
    key_hidden = f'<input type="hidden" name="admin_key" value="{_safe(admin_key)}" />'
    key_query = f"?key={_safe(admin_key)}&workspace_id={_safe(workspace_id)}" if admin_key else f"?workspace_id={_safe(workspace_id)}"
    source_options = "\n".join(f'<option value="{_safe(s)}">{_safe(s)}</option>' for s in ADMIN_RAG_SOURCE_TYPES)
    answer_modes = [
        ("auto", "Auto"),
        ("short", "Ngắn gọn"),
        ("detailed", "Chi tiết"),
        ("chief_accountant", "Kế toán trưởng"),
        ("with_example", "Có ví dụ"),
        ("with_journal", "Có bút toán/checklist"),
        ("risk", "Rủi ro/kiểm soát"),
        ("source_only", "Chỉ nguồn"),
    ]
    answer_mode_options = "\n".join(
        f'<option value="{_safe(value)}" {"selected" if value == answer_mode else ""}>{_safe(label)}</option>'
        for value, label in answer_modes
    )

    doc_rows = []
    for doc in docs:
        warnings = doc.get("warnings") or []
        warn = f'<div class="muted warn">⚠ {_safe(" | ".join(warnings[:2]))}</div>' if warnings else ""
        doc_rows.append(f"""
        <tr>
          <td><b>{_safe(doc.get('title'))}</b><div class="muted">{_safe(doc.get('document_id'))}</div>{warn}</td>
          <td>{_safe(doc.get('workspace_id'))}</td>
          <td>{_safe(doc.get('source_type'))}</td>
          <td>{_safe(doc.get('chunk_count'))}</td>
          <td>{_safe(doc.get('parser'))}</td>
          <td>{_safe(_short(doc.get('updated_at') or doc.get('created_at'), 22))}</td>
          <td class="actions">
            <form method="get" action="/admin/rag-ui" class="inline" data-ajax="true">
              <input type="hidden" name="key" value="{_safe(admin_key)}" />
              <input type="hidden" name="workspace_id" value="{_safe(workspace_id)}" />
              <input type="hidden" name="detail_document_id" value="{_safe(doc.get('document_id'))}" />
              <button type="submit">Xem</button>
            </form>
            <form method="post" action="/admin/rag-ui/reindex" class="inline" data-ajax="true">
              {key_hidden}
              <input type="hidden" name="workspace_id" value="{_safe(workspace_id)}" />
              <input type="hidden" name="document_id" value="{_safe(doc.get('document_id'))}" />
              <button type="submit">Re-index</button>
            </form>
            <form method="post" action="/admin/rag-ui/delete" class="inline" data-ajax="true" data-confirm="Xóa tài liệu này khỏi RAG?">
              {key_hidden}
              <input type="hidden" name="workspace_id" value="{_safe(workspace_id)}" />
              <input type="hidden" name="document_id" value="{_safe(doc.get('document_id'))}" />
              <button type="submit" class="danger">Xóa</button>
            </form>
          </td>
        </tr>
        """)
    doc_rows_html = "\n".join(doc_rows) if doc_rows else '<tr><td colspan="7" class="muted">Chưa có tài liệu RAG trong workspace này.</td></tr>'

    answer_html = ""
    if answer_result:
        sources = answer_result.get("enterprise_sources") or []
        citations = answer_result.get("citations") or []
        if citations:
            src_html = "".join(
                f"<li><b>[{_safe(c.get('index'))}] {_safe(c.get('document_title') or c.get('title'))}</b> — {_safe(c.get('location'))}<br>"
                f"<span class='muted'>{_safe(_short(c.get('excerpt'), 320))}</span></li>"
                for c in citations
            )
        else:
            src_html = "".join(
                f"<li><b>{_safe(s.get('title'))}</b> — {_safe(s.get('heading'))}<br><span class='muted'>{_safe(_short(s.get('snippet'), 260))}</span></li>"
                for s in sources
            ) or "<li>Chưa tìm thấy nguồn phù hợp trong RAG.</li>"
        answer_mode = answer_result.get("answer_mode") or "rag_answer"
        retrieval_mode = answer_result.get("retrieval_mode") or "local/basic"
        memory_used = "yes" if answer_result.get("conversation_memory_used") else "no"
        answer_html = f"""
        <div class="panel result" id="last-rag-answer" data-question="{_safe(question)}" data-answer="{_safe(answer_result.get('answer'))}">
          <h3>Kết quả hỏi thử RAG</h3>
          <p><b>Câu hỏi:</b> {_safe(question)}</p>
          <p class="muted">Mode: {_safe(answer_mode)} | retrieval={_safe(retrieval_mode)} | memory={_safe(memory_used)} | confidence={_safe(answer_result.get('confidence'))}</p>
          <pre>{_safe(answer_result.get('answer'))}</pre>
          <h4>Nguồn / citation</h4>
          <ol>{src_html}</ol>
        </div>
        """

    search_html = ""
    if search_result:
        results = search_result.get("results") or []
        items = "".join(
            f"<li><b>{_safe(r.get('title'))}</b> | score={_safe(r.get('score'))} | {_safe(r.get('location') or r.get('heading'))}<br>"
            f"<span class='muted'>breakdown={_safe(r.get('score_breakdown') or {})}</span><br>"
            f"<span class='muted'>{_safe(_short(r.get('snippet'), 300))}</span></li>"
            for r in results
        ) or "<li>Không tìm thấy chunk phù hợp.</li>"
        search_html = f"""
        <div class="panel result">
          <h3>Kết quả search chunk</h3>
          <ol>{items}</ol>
        </div>
        """

    detail_html = ""
    if document_detail:
        doc = document_detail.get("document") or {}
        chunks = document_detail.get("chunks") or []
        chunk_items = _render_chunk_items(chunks)
        quality_panel = _render_doc_quality_panel(doc, chunks)
        intelligence_panel = _render_doc_intelligence_panel(doc)
        detail_html = f"""
        <div class="panel detail-panel">
          <h3>Chi tiết tài liệu: {_safe(doc.get('title'))}</h3>
          <p class="muted">ID: {_safe(doc.get('document_id'))} | chunks: {_safe(len(chunks))} | status: {_safe(doc.get('status'))}</p>
          <h4>V64 Document Intelligence</h4>
          {intelligence_panel}
          {quality_panel}
          <details>
            <summary>Metadata kỹ thuật</summary>
            <pre>{_safe(doc.get('metadata'))}</pre>
          </details>
          <h4>Chunks đã làm sạch để đọc/debug</h4>
          <div class="chunk-toolbar">
            <input id="chunk-filter" type="search" placeholder="Lọc trong chunks đang hiển thị: Điều 6, 90 ngày, mã số 429..." />
            <button type="button" data-chunk-action="expand">Mở tất cả</button>
            <button type="button" data-chunk-action="collapse">Thu gọn</button>
            <button type="button" data-chunk-action="raw">Ẩn/hiện raw</button>
          </div>
          {chunk_items or '<p class="muted">Không có chunk.</p>'}
        </div>
        """

    eval_html = _render_eval_result(eval_result)
    file_result_html = _render_file_process_result(file_result, admin_key=admin_key)
    message_html = f'<div class="notice ok">{_safe(message)}</div>' if message else ""
    error_html = f'<div class="notice err">{_safe(error)}</div>' if error else ""

    return f"""<!doctype html>
<html lang="vi">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>Finiip Admin RAG UI</title>
  <style>
    :root {{ color-scheme: light; }}
    body {{ font-family: Arial, sans-serif; margin: 0; background: #fff; color: #111; }}
    header {{ padding: 18px 28px; border-bottom: 1px solid #e5e5e5; }}
    main {{ padding: 24px 28px 60px; max-width: 1280px; margin: 0 auto; }}
    h1 {{ margin: 0 0 6px; font-size: 24px; }}
    h2 {{ font-size: 18px; margin-top: 0; }}
    h3 {{ margin-top: 0; }}
    .muted {{ color: #666; font-size: 13px; }}
    .grid {{ display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 16px; }}
    .panel {{ border: 1px solid #ddd; border-radius: 10px; padding: 16px; margin-bottom: 16px; background: #fff; }}
    .stats {{ display: flex; gap: 12px; flex-wrap: wrap; margin-top: 12px; }}
    .stat {{ border: 1px solid #e0e0e0; border-radius: 8px; padding: 10px 12px; min-width: 120px; }}
    label {{ display: block; font-size: 13px; font-weight: 700; margin: 10px 0 5px; }}
    input, select, textarea {{ width: 100%; padding: 9px; border: 1px solid #ccc; border-radius: 7px; box-sizing: border-box; background: #fff; color: #111; }}
    textarea {{ min-height: 90px; }}
    button {{ padding: 8px 12px; border: 1px solid #111; border-radius: 7px; background: #fff; color: #111; cursor: pointer; }}
    button:hover {{ background: #f5f5f5; }}
    button.danger {{ border-color: #a00; color: #a00; }}
    table {{ width: 100%; border-collapse: collapse; }}
    th, td {{ border-bottom: 1px solid #eee; padding: 10px; text-align: left; vertical-align: top; }}
    th {{ background: #fafafa; font-size: 13px; }}
    .actions {{ white-space: nowrap; }}
    .inline {{ display: inline-block; margin: 0 3px 4px 0; }}
    .inline input {{ width: auto; }}
    .notice {{ padding: 10px 12px; border-radius: 8px; margin-bottom: 14px; }}
    #ajax-status {{ border: 1px solid #b8c7e0; background: #f7fbff; }}
    .ok {{ border: 1px solid #8ac58a; background: #f6fff6; }}
    .err {{ border: 1px solid #d48a8a; background: #fff6f6; }}
    pre {{ white-space: pre-wrap; word-break: break-word; border: 1px solid #eee; border-radius: 8px; padding: 10px; background: #fafafa; }}
    details {{ border: 1px solid #eee; border-radius: 8px; padding: 8px; margin: 8px 0; }}
    .doc-health {{ display: grid; grid-template-columns: repeat(3, minmax(0, 1fr)); gap: 10px; border: 1px solid #e8eef7; background: #fbfdff; border-radius: 10px; padding: 12px; margin: 12px 0; }}
    .doc-intel {{ display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 8px; border: 1px solid #e8eef7; background: #fffefa; border-radius: 10px; padding: 12px; margin: 12px 0; }}
    .doc-health details {{ grid-column: 1 / -1; background: #fff; }}
    .chunk-toolbar {{ display: flex; gap: 8px; flex-wrap: wrap; align-items: center; margin: 10px 0 12px; }}
    .chunk-toolbar input {{ flex: 1 1 340px; }}
    .chunk-card {{ border-color: #e3e8ef; background: #fff; }}
    .chunk-card[open] {{ box-shadow: 0 1px 8px rgba(0,0,0,.04); }}
    .chunk-card summary {{ cursor: pointer; line-height: 1.6; }}
    .chip {{ display: inline-block; margin-left: 6px; padding: 2px 7px; border: 1px solid #d7e0ea; border-radius: 999px; background: #f7fbff; font-size: 12px; color: #32445a; }}
    .ok-chip {{ border-color: #a8d8a8; background: #f6fff6; color: #265b26; margin-top: 6px; }}
    .capability-row {{ display: flex; flex-wrap: wrap; gap: 6px; margin-top: 10px; }}
    .memory-box {{ border: 1px dashed #d7e0ea; border-radius: 9px; padding: 10px; margin-top: 12px; background: #fbfdff; }}
    .memory-item {{ padding: 6px 0; border-bottom: 1px solid #eef2f7; }}
    .memory-item:last-child {{ border-bottom: 0; }}
    .muted-chip {{ color: #777; background: #fafafa; }}
    .summary-title {{ margin-left: 8px; color: #333; }}
    .chunk-actions {{ margin: 10px 0 8px; }}
    .chunk-body {{ font-family: Arial, sans-serif; line-height: 1.55; background: #fff; border-color: #e8eef7; }}
    .raw-chunk {{ display: none; background: #fcfcfc; }}
    .show-raw .raw-chunk {{ display: block; }}
    .warn {{ margin-top: 4px; color: #8a5a00; }}
    @media (max-width: 900px) {{ .grid {{ grid-template-columns: 1fr; }} .actions {{ white-space: normal; }} .doc-health {{ grid-template-columns: 1fr; }} }}
  </style>
</head>
<body>
<header>
  <h1>Finiip Admin RAG UI</h1>
  <div class="muted">Giao diện backend trắng đơn giản — chỉ dành cho bạn/Admin nạp tri thức RAG chính thức. User thường không dùng màn này.</div>
</header>
<main>
  <div id="ajax-status" class="notice" style="display:none"></div>
  {message_html}
  {error_html}

  <div class="panel">
    <form method="get" action="/admin/rag-ui" data-ajax="true">
      <label>Admin key khi deploy</label>
      <input name="key" value="{_safe(admin_key)}" placeholder="FINIIP_ADMIN_KEY hoặc FINIIP_API_KEY" />
      <label>Workspace</label>
      <input name="workspace_id" value="{_safe(workspace_id)}" />
      <button type="submit">Mở workspace</button>
    </form>
    <div class="stats">
      <div class="stat"><b>{_safe(stats.get('documents'))}</b><br><span class="muted">Tài liệu active</span></div>
      <div class="stat"><b>{_safe(stats.get('chunks'))}</b><br><span class="muted">Chunks</span></div>
      <div class="stat"><b>{_safe(stats.get('chars'))}</b><br><span class="muted">Ký tự index</span></div>
      <div class="stat"><b>{_safe(stats.get('by_source_type'))}</b><br><span class="muted">Loại tài liệu</span></div>
      <div class="stat"><b>{_safe(storage_status.get('active_backend'))}</b><br><span class="muted">Storage</span></div>
    </div>
    <p class="muted">Supabase: configured={_safe(storage_status.get('configured'))}, mode={_safe(storage_status.get('mode'))}, bucket={_safe(storage_status.get('bucket'))}. Muốn bật Supabase: đặt <code>RAG_STORAGE_MODE=supabase</code>.</p>
    <div class="capability-row">
      <span class="chip ok-chip">V54 Hybrid Search</span>
      <span class="chip ok-chip">V55 Rerank</span>
      <span class="chip ok-chip">V60 Test Center</span>
      <span class="chip ok-chip">V61 Citation Điều/Khoản/Điểm</span>
      <span class="chip ok-chip">V62 Answer Mode</span>
      <span class="chip ok-chip">V63 Accounting Workflow</span>
      <span class="chip ok-chip">V64 Document Intelligence</span>
      <span class="chip ok-chip">V65 Conflict Checker</span>
      <span class="chip ok-chip">V66 Persistent Memory</span>
      <span class="chip ok-chip">V67 File Reader → Return File</span>
      <span class="chip ok-chip">Formula Engine</span>
    </div>
  </div>

  <div class="grid">
    <section class="panel">
      <h2>Upload tài liệu vào RAG chính thức</h2>
      <form method="post" action="/admin/rag-ui/upload" enctype="multipart/form-data" data-ajax="true">
        {key_hidden}
        <label>Workspace</label>
        <input name="workspace_id" value="{_safe(workspace_id)}" />
        <label>Tiêu đề tài liệu</label>
        <input name="title" placeholder="VD: Thông tư thuế GTGT / Quy trình thanh toán nội bộ" />
        <label>Loại tài liệu</label>
        <select name="source_type">{source_options}</select>
        <label>Ghi chú admin</label>
        <textarea name="note" placeholder="VD: tài liệu chuẩn áp dụng cho toàn bộ workspace này"></textarea>
        <label>File</label>
        <input type="file" name="file" required />
        <p class="muted">Cho phép: PDF, DOCX, TXT, MD, CSV, JSON, XLSX. File này sẽ đi vào RAG chính thức.</p>
        <button type="submit">Upload & Index</button>
      </form>
    </section>

    <section class="panel">
      <h2>Hỏi thử RAG</h2>
      <form method="post" action="/admin/rag-ui/ask" data-ajax="true" id="rag-ask-form">
        {key_hidden}
        <input type="hidden" name="history" id="rag-history-field" value="" />
        <label>Workspace</label>
        <input name="workspace_id" value="{_safe(workspace_id)}" />
        <label>Chế độ trả lời</label>
        <select name="answer_mode">{answer_mode_options}</select>
        <input type="hidden" name="conversation_id" value="admin" />
        <label>Câu hỏi</label>
        <textarea name="question" placeholder="VD: Hãy giải thích chi tiết thời hạn nộp BCTC hợp nhất và căn cứ điều khoản nào?">{_safe(question)}</textarea>
        <button type="submit">Hỏi thử</button>
      </form>
      <div class="memory-box">
        <div><b>Memory hội thoại</b> <span class="muted">V66: dùng cả local tab + Supabase memory nếu đã chạy schema mới.</span></div>
        <div id="rag-memory-list" class="muted">Chưa có lịch sử hỏi thử.</div>
        <button type="button" data-memory-action="clear">Xóa memory tab + server</button>
      </div>
      <hr>
      <form method="post" action="/admin/rag-ui/search" data-ajax="true">
        {key_hidden}
        <label>Search chunk thô</label>
        <input name="query" value="{_safe(search_query)}" placeholder="VAT, chi phí được trừ, khấu hao..." />
        <input type="hidden" name="workspace_id" value="{_safe(workspace_id)}" />
        <button type="submit">Search</button>
      </form>
    </section>
  </div>

  <section class="panel">
    <h2>V67 Đọc file & trả file kết quả</h2>
    <p class="muted">Upload file tạm để Finiip đọc, trích xuất, tóm tắt, review kế toán hoặc tạo câu hỏi rồi trả về file tải xuống. File này <b>không</b> tự đưa vào RAG chính thức.</p>
    <form method="post" action="/admin/rag-ui/file/process" enctype="multipart/form-data" data-ajax="true">
      {key_hidden}
      <input type="hidden" name="workspace_id" value="{_safe(workspace_id)}" />
      <label>File cần đọc/xử lý</label>
      <input type="file" name="file" required />
      <label>Kiểu xử lý</label>
      <select name="task_type">
        <option value="summary">Tóm tắt tài liệu</option>
        <option value="extract">Đọc & trích xuất text</option>
        <option value="qa">Trả lời câu hỏi theo file</option>
        <option value="accounting_review">Review kế toán/kiểm soát</option>
        <option value="questions">Tạo câu hỏi ôn tập + đáp án</option>
      </select>
      <label>Định dạng file trả về</label>
      <select name="output_format">
        <option value="docx">Word .docx</option>
        <option value="xlsx">Excel .xlsx</option>
        <option value="md">Markdown .md</option>
        <option value="txt">Text .txt</option>
        <option value="json">JSON .json</option>
        <option value="csv">CSV .csv</option>
      </select>
      <label>Câu hỏi nếu chọn QA</label>
      <textarea name="question" placeholder="VD: File này nói gì về thời hạn nộp báo cáo? hoặc Hãy tìm các rủi ro kế toán trong file"></textarea>
      <label>Yêu cầu thêm</label>
      <textarea name="instruction" placeholder="VD: trả lời chi tiết, chia mục, nêu căn cứ, tạo bảng checklist..."></textarea>
      <button type="submit">Đọc file & tạo file kết quả</button>
    </form>
  </section>

  <section class="panel">
    <h2>V60 Test Center</h2>
    <p class="muted">Test hàng loạt sau mỗi lần upload/re-index. Nhập mỗi dòng dạng <code>Câu hỏi => đáp án mong muốn</code> hoặc <code>Câu hỏi | đáp án | nguồn mong muốn</code>.</p>
    <form method="post" action="/admin/rag-ui/eval" data-ajax="true">
      {key_hidden}
      <input type="hidden" name="workspace_id" value="{_safe(workspace_id)}" />
      <label>Chế độ test</label>
      <select name="answer_mode">{answer_mode_options}</select>
      <label>Bộ câu hỏi kiểm thử</label>
      <textarea name="cases" style="min-height:150px" placeholder="Thông tư 43/2026 sửa đổi thông tư nào? => Thông tư 202/2014/TT-BTC
BCTC hợp nhất năm nộp chậm nhất bao nhiêu ngày? => 90 ngày
Lợi ích cổ đông không kiểm soát mã số nào? => 429">{_safe(eval_cases)}</textarea>
      <button type="submit">Chạy Test Center</button>
    </form>
  </section>

  {file_result_html}
  {eval_html}
  {answer_html}
  {search_html}
  {detail_html}

  <section class="panel">
    <h2>Danh sách tài liệu RAG</h2>
    <p class="muted">Các tài liệu ở đây là knowledge chính thức do Admin nạp. User upload hóa đơn/báo cáo không nên nằm trong bảng này.</p>
    <table>
      <thead>
        <tr><th>Tài liệu</th><th>Workspace</th><th>Loại</th><th>Chunks</th><th>Parser</th><th>Cập nhật</th><th>Thao tác</th></tr>
      </thead>
      <tbody>{doc_rows_html}</tbody>
    </table>
  </section>

  <p class="muted">Mở nhanh: <code>/admin/rag-ui?key=YOUR_ADMIN_KEY&workspace_id=default</code></p>
</main>

<script>
(function () {{
  const MEMORY_KEY = 'finiip_admin_rag_memory_v1';

  function showAjaxStatus(message, isError) {{
    const box = document.getElementById('ajax-status');
    if (!box) return;
    box.textContent = message || '';
    box.style.display = message ? 'block' : 'none';
    box.className = 'notice ' + (isError ? 'err' : '');
  }}

  function loadMemory() {{
    try {{ return JSON.parse(localStorage.getItem(MEMORY_KEY) || '[]'); }} catch (err) {{ return []; }}
  }}

  function saveMemory(items) {{
    localStorage.setItem(MEMORY_KEY, JSON.stringify((items || []).slice(0, 6)));
  }}

  function compactMemoryText() {{
    return loadMemory().map(function (item, idx) {{
      return '[' + (idx + 1) + '] Q: ' + (item.q || '') + '\nA: ' + (item.a || '').slice(0, 700);
    }}).join('\n\n');
  }}

  function escapeHtml(text) {{
    return String(text || '').replace(/[&<>"]/g, function (ch) {{
      return ({{'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;'}})[ch];
    }});
  }}

  function renderMemory() {{
    const box = document.getElementById('rag-memory-list');
    if (!box) return;
    const items = loadMemory();
    if (!items.length) {{ box.textContent = 'Chưa có lịch sử hỏi thử.'; return; }}
    box.innerHTML = items.map(function (item) {{
      return '<div class="memory-item"><b>Q:</b> ' + escapeHtml((item.q || '').slice(0, 160)) + '<br><b>A:</b> ' + escapeHtml((item.a || '').slice(0, 220)) + '</div>';
    }}).join('');
  }}



  // V58 Admin UX: preserve form values/open panels when AJAX updates the page.
  // This fixes the annoying "bấm nút là trang như bị reset" feeling.
  function fieldKey(field) {{
    const form = field.form;
    const formId = form ? ((form.getAttribute('action') || '') + '|' + (form.getAttribute('method') || 'get')) : 'nofrom';
    return formId + '|' + (field.getAttribute('name') || field.id || 'field') + '|' + Array.prototype.indexOf.call(document.querySelectorAll('input, textarea, select'), field);
  }}

  function collectUiState() {{
    const state = {{ fields: {{}}, details: {{}}, scrollY: window.scrollY, active: null }};
    document.querySelectorAll('input, textarea, select').forEach(function (field) {{
      if (!field.name && !field.id) return;
      if (field.type === 'file') return;
      const key = fieldKey(field);
      if (field.type === 'checkbox' || field.type === 'radio') state.fields[key] = {{ checked: field.checked }};
      else state.fields[key] = {{ value: field.value }};
    }});
    document.querySelectorAll('details').forEach(function (node, idx) {{
      const summary = node.querySelector('summary');
      const label = summary ? summary.textContent.trim().slice(0, 120) : 'details-' + idx;
      state.details[label + '|' + idx] = node.open;
    }});
    const active = document.activeElement;
    if (active && (active.name || active.id)) state.active = fieldKey(active);
    return state;
  }}

  function restoreUiState(state) {{
    if (!state) return;
    document.querySelectorAll('input, textarea, select').forEach(function (field) {{
      if (!field.name && !field.id) return;
      if (field.type === 'file') return;
      const saved = state.fields[fieldKey(field)];
      if (!saved) return;
      if (field.type === 'checkbox' || field.type === 'radio') field.checked = !!saved.checked;
      else field.value = saved.value;
    }});
    document.querySelectorAll('details').forEach(function (node, idx) {{
      const summary = node.querySelector('summary');
      const label = summary ? summary.textContent.trim().slice(0, 120) : 'details-' + idx;
      const key = label + '|' + idx;
      if (Object.prototype.hasOwnProperty.call(state.details, key)) node.open = !!state.details[key];
    }});
    if (state.active) {{
      const activeField = Array.prototype.find.call(document.querySelectorAll('input, textarea, select'), function (field) {{
        return fieldKey(field) === state.active;
      }});
      if (activeField) {{
        try {{ activeField.focus({{ preventScroll: true }}); }} catch (err) {{ try {{ activeField.focus(); }} catch (err2) {{}} }}
      }}
    }}
    window.scrollTo(0, Math.min(state.scrollY || 0, document.body.scrollHeight));
  }}

  function captureAnswerToMemory() {{
    const node = document.getElementById('last-rag-answer');
    if (!node) return;
    const q = node.getAttribute('data-question') || '';
    const a = node.getAttribute('data-answer') || '';
    if (!q || !a) return;
    const items = loadMemory().filter(function (item) {{ return item.q !== q; }});
    items.unshift({{ q: q, a: a, ts: new Date().toISOString() }});
    saveMemory(items);
    renderMemory();
  }}

  async function submitAjaxForm(form, submitter) {{
    const confirmMessage = form.getAttribute('data-confirm');
    if (confirmMessage && !window.confirm(confirmMessage)) return;

    const method = (form.getAttribute('method') || 'get').toUpperCase();
    let url = form.getAttribute('action') || window.location.pathname;
    let options = {{ method, credentials: 'same-origin' }};
    const historyInput = form.querySelector && form.querySelector('#rag-history-field');
    if (historyInput) historyInput.value = compactMemoryText();
    const formData = new FormData(form);

    if (method === 'GET') {{
      const qs = new URLSearchParams(formData);
      url += (url.includes('?') ? '&' : '?') + qs.toString();
      options.method = 'GET';
    }} else {{
      options.body = formData;
    }}

    const uiState = collectUiState();
    const oldText = submitter ? submitter.textContent : '';
    if (submitter) {{
      submitter.disabled = true;
      submitter.textContent = 'Đang xử lý...';
    }}
    showAjaxStatus('Đang xử lý, trang sẽ không reload...', false);
    const scrollY = window.scrollY;

    try {{
      const response = await fetch(url, options);
      const html = await response.text();
      const parsed = new DOMParser().parseFromString(html, 'text/html');
      const incomingMain = parsed.querySelector('main');
      const currentMain = document.querySelector('main');
      if (!incomingMain || !currentMain) {{
        throw new Error('Server không trả về HTML admin hợp lệ.');
      }}
      currentMain.innerHTML = incomingMain.innerHTML;
      restoreUiState(uiState);
      if (method === 'GET') {{
        window.history.replaceState({{}}, '', url);
      }}
      window.scrollTo(0, Math.min(scrollY, document.body.scrollHeight));
      const status = document.getElementById('ajax-status');
      if (status && response.ok) {{
        status.textContent = 'Xong. Nội dung đã cập nhật mà không reload trang.';
        status.style.display = 'block';
        status.className = 'notice ok';
      }}
      captureAnswerToMemory();
      renderMemory();
    }} catch (err) {{
      showAjaxStatus('Lỗi AJAX: ' + (err && err.message ? err.message : err), true);
      if (submitter) {{
        submitter.disabled = false;
        submitter.textContent = oldText;
      }}
    }}
  }}

  document.addEventListener('submit', function (event) {{
    const form = event.target;
    if (!form || form.getAttribute('data-ajax') !== 'true') return;
    event.preventDefault();
    submitAjaxForm(form, event.submitter || form.querySelector('button[type="submit"]'));
  }});

  document.addEventListener('click', async function (event) {{
    const memoryButton = event.target.closest && event.target.closest('[data-memory-action]');
    if (memoryButton && memoryButton.getAttribute('data-memory-action') === 'clear') {{
      saveMemory([]);
      renderMemory();
      try {{
        const fd = new FormData();
        const keyInput = document.querySelector('input[name="admin_key"]') || document.querySelector('input[name="key"]');
        const wsInput = document.querySelector('input[name="workspace_id"]');
        fd.append('admin_key', keyInput ? keyInput.value : '');
        fd.append('workspace_id', wsInput ? wsInput.value : 'default');
        fd.append('conversation_id', 'admin');
        await fetch('/admin/rag-ui/memory/clear', {{ method: 'POST', body: fd, credentials: 'same-origin' }});
      }} catch (err) {{}}
      showAjaxStatus('Đã xóa memory tab và yêu cầu xóa memory server.', false);
      return;
    }}

    const copyButton = event.target.closest && event.target.closest('.copy-chunk');
    if (copyButton) {{
      const targetId = copyButton.getAttribute('data-copy-target');
      const target = targetId ? document.getElementById(targetId) : null;
      if (!target) return;
      try {{
        await navigator.clipboard.writeText(target.textContent || '');
        const old = copyButton.textContent;
        copyButton.textContent = 'Đã copy';
        setTimeout(function () {{ copyButton.textContent = old; }}, 1200);
      }} catch (err) {{
        showAjaxStatus('Không copy được tự động, hãy bôi đen để copy thủ công.', true);
      }}
      return;
    }}

    const actionButton = event.target.closest && event.target.closest('[data-chunk-action]');
    if (!actionButton) return;
    const action = actionButton.getAttribute('data-chunk-action');
    if (action === 'expand' || action === 'collapse') {{
      document.querySelectorAll('.chunk-card').forEach(function (card) {{
        card.open = action === 'expand';
      }});
    }} else if (action === 'raw') {{
      const panel = document.querySelector('.detail-panel');
      if (panel) panel.classList.toggle('show-raw');
    }}
  }});

  document.addEventListener('input', function (event) {{
    if (!event.target || event.target.id !== 'chunk-filter') return;
    const q = (event.target.value || '').toLowerCase().trim();
    document.querySelectorAll('.chunk-card').forEach(function (card) {{
      const text = card.getAttribute('data-chunk-text') || '';
      const matched = !q || text.includes(q);
      card.style.display = matched ? '' : 'none';
      if (matched && q) card.open = true;
    }});
  }});

  captureAnswerToMemory();
  renderMemory();
}})();
</script>

</body>
</html>"""


def render_admin_login_or_unauthorized() -> str:
    return """<!doctype html><html lang="vi"><head><meta charset="utf-8"><title>Finiip Admin RAG UI</title>
<style>body{font-family:Arial,sans-serif;background:#fff;color:#111;margin:40px;max-width:680px}input{width:100%;padding:10px;border:1px solid #ccc;border-radius:8px}button{margin-top:12px;padding:9px 14px;border:1px solid #111;border-radius:8px;background:#fff}</style></head><body>
<h1>Finiip Admin RAG UI</h1><p>Nhập admin key để mở giao diện upload/quản lý RAG.</p>
<form method="get" action="/admin/rag-ui" data-ajax="true"><input name="key" type="password" placeholder="FINIIP_ADMIN_KEY hoặc FINIIP_API_KEY"><input type="hidden" name="workspace_id" value="default"><button type="submit">Mở Admin UI</button></form>
<p><small>Khi deploy, đặt biến môi trường <code>FINIIP_ADMIN_KEY</code> để bảo vệ màn này.</small></p>
</body></html>"""

"""Finiip V68-V72 - Frontend file reader/report generator.

This module is intentionally separate from the official RAG knowledge base:
frontend users can upload one or many files, Finiip reads them, creates a
structured report, stores a job record, and returns a downloadable output file.

V68: public frontend API helpers
V69: async job_id workflow for large files
V70: processing history
V71: PDF export
V72: multi-file report generation
"""
from __future__ import annotations

import csv
import hashlib
import html
import json
import os
import re
import shutil
import unicodedata
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

from services.accounting_ai_enterprise import (
    ROOT_DIR,
    extract_text_from_bytes,
    normalize_extracted_text_for_rag,
)
from services.smart_orchestrator_v110 import analyze_request, build_attachment_context

FILE_REPORT_VERSION = "v110_intelligent_file_report"
FILE_REPORT_ROOT = ROOT_DIR / "data" / "file_report_jobs"
FILE_REPORT_ROOT.mkdir(parents=True, exist_ok=True)
HISTORY_FILE = FILE_REPORT_ROOT / "history.json"

FILE_REPORT_OUTPUT_FORMATS = ["docx", "xlsx", "pdf", "md", "txt", "json", "csv"]
FILE_REPORT_TASK_TYPES = [
    "auto_report",
    "summary",
    "accounting_review",
    "legal_review",
    "financial_report",
    "qa",
    "extract",
    "study_questions",
]
FILE_REPORT_STYLES = ["short", "standard", "detailed", "executive", "accounting_manager"]

MAX_FILE_REPORT_MB = int(os.getenv("FINIIP_FILE_REPORT_MAX_MB", "50"))
MAX_FILE_REPORT_BYTES = MAX_FILE_REPORT_MB * 1024 * 1024
MAX_FILE_REPORT_FILES = int(os.getenv("FINIIP_FILE_REPORT_MAX_FILES", "10"))

ALLOWED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".xlsx",
    ".xlsm",
    ".html",
    ".htm",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".tif",
    ".tiff",
}

VIETNAMESE_ACCOUNTING_TERMS = [
    "doanh thu", "chi phí", "lợi nhuận", "thuế", "gtgt", "tndn", "hóa đơn",
    "chứng từ", "báo cáo tài chính", "báo cáo tình hình tài chính", "lưu chuyển tiền tệ",
    "tài sản", "nợ phải trả", "vốn chủ sở hữu", "công ty mẹ", "công ty con",
    "hợp nhất", "khấu hao", "dự phòng", "hàng tồn kho", "phải thu", "phải trả",
    "định khoản", "ghi sổ", "quyết toán", "kiểm toán", "kiểm soát", "rủi ro",
]

VIETNAMESE_LEGAL_TERMS = [
    "thông tư", "nghị định", "luật", "điều", "khoản", "điểm", "phụ lục",
    "hiệu lực", "sửa đổi", "bổ sung", "bãi bỏ", "quy định", "căn cứ",
]


@dataclass
class FileReportInput:
    filename: str
    content: bytes


def _utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def _safe_filename(name: str, fallback: str = "upload.bin") -> str:
    raw = Path(name or fallback).name
    raw = re.sub(r"[^A-Za-z0-9À-ỹĐđ._()\-\s]", "_", raw).strip(" .")
    return raw or fallback


def _slug(value: str, max_len: int = 80) -> str:
    value = unicodedata.normalize("NFKD", value or "")
    value = "".join(ch for ch in value if not unicodedata.combining(ch))
    value = re.sub(r"[^A-Za-z0-9]+", "-", value).strip("-").lower()
    return (value[:max_len].strip("-") or "file-report")


def _load_history() -> Dict[str, Any]:
    if not HISTORY_FILE.exists():
        return {"version": FILE_REPORT_VERSION, "items": []}
    try:
        data = json.loads(HISTORY_FILE.read_text(encoding="utf-8"))
        if not isinstance(data, dict):
            return {"version": FILE_REPORT_VERSION, "items": []}
        data.setdefault("items", [])
        return data
    except Exception:
        return {"version": FILE_REPORT_VERSION, "items": []}


def _save_history(data: Dict[str, Any]) -> None:
    FILE_REPORT_ROOT.mkdir(parents=True, exist_ok=True)
    HISTORY_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _job_dir(job_id: str) -> Path:
    safe = re.sub(r"[^A-Za-z0-9_\-]", "", job_id or "")
    if not safe:
        raise FileNotFoundError("Thiếu job_id")
    return FILE_REPORT_ROOT / safe


def _job_meta_path(job_id: str) -> Path:
    return _job_dir(job_id) / "job.json"


def _read_job(job_id: str) -> Dict[str, Any]:
    path = _job_meta_path(job_id)
    if not path.exists():
        raise FileNotFoundError(f"Không tìm thấy job_id={job_id}")
    return json.loads(path.read_text(encoding="utf-8"))


def _write_job(job: Dict[str, Any]) -> Dict[str, Any]:
    job_id = job["job_id"]
    folder = _job_dir(job_id)
    folder.mkdir(parents=True, exist_ok=True)
    job["updated_at"] = _utc_now()
    _job_meta_path(job_id).write_text(json.dumps(job, ensure_ascii=False, indent=2), encoding="utf-8")
    _upsert_history(job)
    return job


def _upsert_history(job: Dict[str, Any]) -> None:
    store = _load_history()
    items = [x for x in store.get("items", []) if x.get("job_id") != job.get("job_id")]
    summary = {
        "job_id": job.get("job_id"),
        "status": job.get("status"),
        "created_at": job.get("created_at"),
        "updated_at": job.get("updated_at"),
        "workspace_id": job.get("workspace_id"),
        "user_id": job.get("user_id"),
        "task_type": job.get("task_type"),
        "output_format": job.get("output_format"),
        "files_count": len(job.get("files") or []),
        "title": job.get("title"),
        "download_url": job.get("download_url"),
        "output_filename": job.get("output_filename"),
        "error": job.get("error"),
        "deleted": bool(job.get("deleted")),
    }
    items.insert(0, summary)
    store["items"] = items[:500]
    store["updated_at"] = _utc_now()
    _save_history(store)


def validate_file_report_inputs(files: List[FileReportInput]) -> None:
    if not files:
        raise ValueError("Cần upload ít nhất 1 file")
    if len(files) > MAX_FILE_REPORT_FILES:
        raise ValueError(f"Tối đa {MAX_FILE_REPORT_FILES} file/lần")
    total = 0
    for f in files:
        name = _safe_filename(f.filename)
        suffix = Path(name).suffix.lower()
        if suffix not in ALLOWED_EXTENSIONS:
            raise ValueError(f"Định dạng chưa hỗ trợ: {suffix or '(không có đuôi file)'}")
        size = len(f.content or b"")
        total += size
        if size <= 0:
            raise ValueError(f"File rỗng: {name}")
        if size > MAX_FILE_REPORT_BYTES:
            raise ValueError(f"File quá lớn: {name}. Tối đa {MAX_FILE_REPORT_MB}MB/file")
    if total > MAX_FILE_REPORT_BYTES * max(1, min(3, MAX_FILE_REPORT_FILES)):
        raise ValueError("Tổng dung lượng upload quá lớn cho một job xử lý")


def create_job_id(files: List[FileReportInput], workspace_id: str, task_type: str) -> str:
    seed = "|".join([f.filename + ":" + hashlib.sha256(f.content).hexdigest()[:16] for f in files])
    # A random nonce prevents collisions when the same files are exported to
    # multiple formats within the same second.
    seed += f"|{workspace_id}|{task_type}|{datetime.now(timezone.utc).isoformat(timespec='microseconds')}|{uuid.uuid4().hex}"
    return "fr_" + hashlib.sha1(seed.encode("utf-8")).hexdigest()[:18]


def create_file_report_job(
    *,
    files: List[FileReportInput],
    instruction: str = "",
    question: str = "",
    task_type: str = "auto_report",
    output_format: str = "docx",
    report_style: str = "detailed",
    workspace_id: str = "default",
    user_id: str = "anonymous",
    title: str = "",
    save_inputs: bool = True,
) -> Dict[str, Any]:
    validate_file_report_inputs(files)
    task = (task_type or "auto_report").strip().lower()
    if task not in FILE_REPORT_TASK_TYPES:
        task = "auto_report"
    fmt = (output_format or "docx").lower().strip(".")
    if fmt not in FILE_REPORT_OUTPUT_FORMATS:
        fmt = "docx"
    style = (report_style or "detailed").strip().lower()
    if style not in FILE_REPORT_STYLES:
        style = "detailed"
    job_id = create_job_id(files, workspace_id, task)
    folder = _job_dir(job_id)
    input_dir = folder / "inputs"
    output_dir = folder / "outputs"
    input_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)
    file_meta = []
    for idx, f in enumerate(files, 1):
        name = _safe_filename(f.filename, f"file_{idx}.bin")
        sha = hashlib.sha256(f.content).hexdigest()
        stored_path = ""
        if save_inputs:
            stored = input_dir / f"{idx:02d}_{name}"
            stored.write_bytes(f.content)
            stored_path = str(stored)
        file_meta.append({
            "index": idx,
            "filename": name,
            "size_bytes": len(f.content),
            "sha256": sha,
            "stored_path": stored_path,
        })
    job = {
        "version": FILE_REPORT_VERSION,
        "job_id": job_id,
        "status": "queued",
        "created_at": _utc_now(),
        "updated_at": _utc_now(),
        "workspace_id": workspace_id or "default",
        "user_id": user_id or "anonymous",
        "title": title or _default_report_title(files, task),
        "task_type": task,
        "output_format": fmt,
        "report_style": style,
        "instruction": instruction or "",
        "question": question or "",
        "save_inputs": bool(save_inputs),
        "files": file_meta,
        "output_dir": str(output_dir),
        "download_url": None,
        "output_path": None,
        "output_filename": None,
        "error": None,
        "deleted": False,
    }
    return _write_job(job)


def _default_report_title(files: List[FileReportInput], task_type: str) -> str:
    if len(files) == 1:
        return f"Báo cáo đọc file - {_safe_filename(files[0].filename)}"
    return f"Báo cáo tổng hợp {len(files)} file - {task_type}"


def run_file_report_job(job_id: str) -> Dict[str, Any]:
    job = _read_job(job_id)
    if job.get("deleted"):
        raise FileNotFoundError("Job đã bị xóa")
    job["status"] = "processing"
    job["started_at"] = _utc_now()
    job["error"] = None
    _write_job(job)
    try:
        files = []
        for meta in job.get("files") or []:
            path = Path(meta.get("stored_path") or "")
            if not path.exists():
                raise FileNotFoundError(f"Không tìm thấy input đã lưu: {meta.get('filename')}")
            files.append(FileReportInput(filename=meta.get("filename") or path.name, content=path.read_bytes()))
        result = process_file_report_job(
            job=job,
            files=files,
        )
        job.update(result)
        job["status"] = "done"
        job["finished_at"] = _utc_now()
        return _write_job(job)
    except Exception as exc:
        job["status"] = "failed"
        job["error"] = str(exc)
        job["finished_at"] = _utc_now()
        return _write_job(job)


def process_file_report_job(*, job: Dict[str, Any], files: List[FileReportInput]) -> Dict[str, Any]:
    extracted_files: List[Dict[str, Any]] = []
    for item in files:
        extracted = extract_text_from_bytes(item.filename, item.content)
        clean_text = normalize_extracted_text_for_rag(extracted.get("text") or "")
        profile = detect_document_profile(item.filename, clean_text)
        summary = summarize_text(clean_text, instruction=(job.get("instruction") or job.get("question") or ""))
        extracted_files.append({
            "filename": _safe_filename(item.filename),
            "size_bytes": len(item.content),
            "parser": extracted.get("parser"),
            "warnings": extracted.get("warnings") or [],
            "text": clean_text,
            "char_count": len(clean_text),
            "word_count": len(clean_text.split()),
            "profile": profile,
            "summary": summary,
        })
    deterministic_markdown = build_report_markdown(job, extracted_files)
    ai_report = _llm_report_markdown(job, extracted_files)
    markdown = str((ai_report or {}).get("markdown") or deterministic_markdown)
    analysis_mode = "llm_grounded_file_report" if ai_report else "deterministic_file_report"
    payload = build_report_payload(job, extracted_files, markdown)
    payload["analysis_mode"] = analysis_mode
    payload["ai_analysis"] = {k: v for k, v in (ai_report or {}).items() if k != "markdown"}
    payload.setdefault("report", {})["generated_by"] = analysis_mode
    output_path = write_report_output_file(job, markdown, payload)
    return {
        "output_path": str(output_path),
        "output_filename": output_path.name,
        "download_url": f"/ai/v69/file-report/jobs/{job['job_id']}/download",
        "preview": markdown[:6000],
        "report": payload.get("report"),
        "analysis_mode": analysis_mode,
        "ai_analysis": payload.get("ai_analysis"),
        "files_analyzed": [
            {k: v for k, v in f.items() if k != "text"}
            for f in extracted_files
        ],
    }


def detect_document_profile(filename: str, text: str) -> Dict[str, Any]:
    clean = normalize_extracted_text_for_rag(text or "")
    lowered = clean.lower()
    suffix = Path(filename or "").suffix.lower().lstrip(".") or "unknown"
    doc_type = "general_document"
    if "thông tư" in lowered or re.search(r"\btt-btc\b", lowered, flags=re.I):
        doc_type = "circular"
    elif "nghị định" in lowered:
        doc_type = "decree"
    elif "luật" in lowered and "điều" in lowered:
        doc_type = "law_or_regulation"
    elif "hóa đơn" in lowered or "invoice" in lowered:
        doc_type = "invoice"
    elif "báo cáo tài chính" in lowered or "báo cáo tình hình tài chính" in lowered or "bảng cân đối" in lowered:
        doc_type = "financial_statement"
    elif "hợp đồng" in lowered:
        doc_type = "contract"
    elif suffix in {"xlsx", "xlsm", "csv"}:
        doc_type = "spreadsheet"
    doc_number = ""
    m = re.search(r"(?:Số\s*:\s*)?([0-9]{1,4}/[0-9]{4}/[A-ZĐ\-]+)", clean, flags=re.I)
    if m:
        doc_number = m.group(1)
    issued_date = ""
    m = re.search(r"ngày\s+(\d{1,2})\s+tháng\s+(\d{1,2})\s+năm\s+(\d{4})", clean, flags=re.I)
    if m:
        issued_date = f"{int(m.group(1)):02d}/{int(m.group(2)):02d}/{m.group(3)}"
    tags: List[str] = []
    for term in VIETNAMESE_ACCOUNTING_TERMS + VIETNAMESE_LEGAL_TERMS:
        if term in lowered:
            tags.append(term)
    return {
        "document_type": doc_type,
        "extension": suffix,
        "document_number": doc_number,
        "issued_date": issued_date,
        "tags": tags[:20],
        "has_accounting_terms": any(t in lowered for t in VIETNAMESE_ACCOUNTING_TERMS),
        "has_legal_terms": any(t in lowered for t in VIETNAMESE_LEGAL_TERMS),
    }


def _sentences(text: str) -> List[str]:
    text = re.sub(r"\s+", " ", text or " ").strip()
    raw = re.split(r"(?<=[\.!?。])\s+|\n+", text)
    out = []
    for s in raw:
        s = s.strip(" -\t")
        if len(s) >= 30:
            out.append(s[:650])
    if len(out) < 5:
        chunks = [p.strip() for p in re.split(r"\n{1,}", text) if len(p.strip()) >= 30]
        out.extend(chunks)
    return out[:500]


def _score_sentence(sentence: str, instruction: str = "") -> float:
    low = sentence.lower()
    score = 0.0
    for term in VIETNAMESE_ACCOUNTING_TERMS:
        if term in low:
            score += 2.0
    for term in VIETNAMESE_LEGAL_TERMS:
        if term in low:
            score += 1.5
    for word in re.findall(r"[\wÀ-ỹĐđ]{3,}", (instruction or "").lower()):
        if word in low:
            score += 1.0
    if re.search(r"\d", sentence):
        score += 0.8
    if len(sentence) > 280:
        score -= 0.3
    return score


def summarize_text(text: str, instruction: str = "") -> Dict[str, Any]:
    sents = _sentences(text)
    ranked = sorted(enumerate(sents), key=lambda x: (_score_sentence(x[1], instruction), -x[0]), reverse=True)
    key = [s for _, s in ranked[:10]]
    opening = [s for s in sents[:5]]
    risks = []
    lowered = (text or "").lower()
    if any(t in lowered for t in ["thuế", "hóa đơn", "gtgt", "tndn"]):
        risks.append("Cần đối chiếu kỳ thuế, chứng từ gốc và điều kiện khấu trừ/ghi nhận.")
    if any(t in lowered for t in ["hợp nhất", "công ty mẹ", "công ty con"]):
        risks.append("Cần kiểm tra phạm vi hợp nhất, thời điểm kiểm soát/mất kiểm soát và giao dịch nội bộ.")
    if any(t in lowered for t in ["thông tư", "nghị định", "luật"]):
        risks.append("Cần kiểm tra hiệu lực văn bản và văn bản sửa đổi/bổ sung mới hơn trước khi áp dụng.")
    if not risks:
        risks.append("Cần kiểm tra lại file gốc vì báo cáo này được tạo tự động từ nội dung trích xuất.")
    return {
        "opening_lines": opening,
        "key_points": key,
        "risks": risks,
        "sentence_count": len(sents),
    }


def _llm_report_markdown(job: Dict[str, Any], files: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    """Create a grounded long-form report from selected file chunks."""
    if not os.getenv("OPENAI_API_KEY") or (job.get("task_type") == "extract"):
        return None
    try:
        from openai import OpenAI  # type: ignore

        instruction = str(job.get("instruction") or job.get("question") or "Tóm tắt, phân tích và lập báo cáo từ các tệp.")
        file_context = build_attachment_context(
            instruction,
            [{"filename": f.get("filename"), "text": f.get("text") or ""} for f in files],
            max_total_chars=int(os.getenv("FINIIP_FILE_REPORT_CONTEXT_CHARS", "70000")),
        )
        context = str(file_context.get("context") or "")
        if not context.strip():
            return None
        request_plan = analyze_request(instruction)
        model = os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini")
        max_tokens = int(os.getenv("FINIIP_FILE_REPORT_MAX_OUTPUT_TOKENS", "5000"))
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        task_type = str(job.get("task_type") or "auto_report")
        style = str(job.get("report_style") or "detailed")
        system = (
            "Bạn là Finiip, trợ lý AI thuộc CTCP IIP Việt Nam, chuyên đọc hồ sơ, kế toán, thuế, tài chính và lập báo cáo. "
            "Chỉ sử dụng dữ kiện có trong các đoạn tệp được cung cấp. Không bịa số, điều khoản, văn bản hay kết luận. "
            "Khi dữ liệu thiếu hoặc mâu thuẫn, phải nêu rõ. Viết Markdown tiếng Việt chuyên nghiệp, có tóm tắt điều hành, "
            "phân tích theo từng yêu cầu, bảng số liệu dạng Markdown khi phù hợp, rủi ro, checklist và kết luận. "
            "Không đưa đường dẫn file nội bộ và không nói rằng bạn đã đọc phần không được cung cấp."
        )
        user = (
            f"TIÊU ĐỀ: {job.get('title')}\n"
            f"LOẠI NHIỆM VỤ: {task_type}\n"
            f"PHONG CÁCH: {style}\n"
            f"YÊU CẦU: {instruction}\n"
            f"KẾ HOẠCH YÊU CẦU: {json.dumps(request_plan, ensure_ascii=False)}\n\n"
            f"CÁC ĐOẠN TỆP ĐÃ CHỌN:\n{context}\n\n"
            "Hãy tạo báo cáo hoàn chỉnh. Mỗi số liệu quan trọng phải gắn với tên tệp hoặc đoạn tệp tương ứng trong nội dung."
        )
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
            temperature=0.1,
            max_tokens=max_tokens,
        )
        body = str(response.choices[0].message.content or "").strip()
        if not body:
            return None
        if not body.startswith("#"):
            body = f"# {job.get('title') or 'Báo cáo Finiip'}\n\n" + body
        return {
            "markdown": body,
            "model": model,
            "context_characters": file_context.get("context_characters"),
            "manifest": file_context.get("manifest") or [],
            "request_analysis": request_plan,
        }
    except Exception:
        return None


def build_report_markdown(job: Dict[str, Any], files: List[Dict[str, Any]]) -> str:
    title = job.get("title") or "Báo cáo đọc file Finiip"
    instruction = job.get("instruction") or ""
    question = job.get("question") or ""
    task_type = job.get("task_type") or "auto_report"
    style = job.get("report_style") or "detailed"
    now = _utc_now()
    lines: List[str] = [
        f"# {title}",
        "",
        "## 1. Thông tin xử lý",
        "",
        f"- Job ID: `{job.get('job_id')}`",
        f"- Workspace: `{job.get('workspace_id')}`",
        f"- Người dùng: `{job.get('user_id')}`",
        f"- Kiểu xử lý: `{task_type}`",
        f"- Kiểu báo cáo: `{style}`",
        f"- Thời điểm tạo: {now}",
        f"- Số file: {len(files)}",
    ]
    if instruction:
        lines += [f"- Yêu cầu: {instruction}"]
    if question:
        lines += [f"- Câu hỏi: {question}"]
    lines += [""]

    if task_type == "extract":
        lines += ["## 2. Text trích xuất", ""]
        for f in files:
            lines += [f"### {f['filename']}", "", (f.get("text") or "")[:80000], ""]
        return "\n".join(lines).strip() + "\n"

    lines += ["## 2. Tóm tắt điều hành", ""]
    if len(files) == 1:
        points = files[0].get("summary", {}).get("key_points") or []
        if points:
            lines += [f"- {p}" for p in points[:5]]
        else:
            lines += ["- File đã được đọc nhưng chưa phát hiện điểm nổi bật rõ ràng."]
    else:
        lines += [f"- Đã đọc {len(files)} file và tổng hợp thành một báo cáo chung."]
        for f in files:
            profile = f.get("profile") or {}
            lines.append(f"- {f['filename']}: loại `{profile.get('document_type')}`, {f.get('word_count', 0)} từ.")
    lines += [""]

    lines += ["## 3. Nhận diện tài liệu", ""]
    for f in files:
        profile = f.get("profile") or {}
        lines += [
            f"### {f['filename']}",
            "",
            f"- Parser: `{f.get('parser')}`",
            f"- Số ký tự: {f.get('char_count')}",
            f"- Số từ: {f.get('word_count')}",
            f"- Loại tài liệu: `{profile.get('document_type')}`",
            f"- Số văn bản: {profile.get('document_number') or 'chưa phát hiện'}",
            f"- Ngày ban hành: {profile.get('issued_date') or 'chưa phát hiện'}",
            f"- Tags: {', '.join(profile.get('tags') or []) or 'chưa có'}",
        ]
        if f.get("warnings"):
            lines.append(f"- Cảnh báo parser: {'; '.join(f.get('warnings') or [])}")
        lines.append("")

    if task_type == "qa":
        lines += ["## 4. Trả lời theo file", ""]
        query = question or instruction
        for f in files:
            passages = select_relevant_passages(f.get("text") or "", query, limit=6)
            lines += [f"### {f['filename']}", ""]
            if passages:
                for i, p in enumerate(passages, 1):
                    lines.append(f"{i}. {p}")
            else:
                lines.append("Chưa tìm thấy đoạn liên quan rõ với câu hỏi trong file này.")
            lines.append("")
    elif task_type == "study_questions":
        lines += ["## 4. Câu hỏi ôn tập và đáp án gợi ý", ""]
        n = 1
        for f in files:
            passages = select_relevant_passages(f.get("text") or "", instruction or question, limit=8)
            for p in passages:
                lines += [f"### Câu {n}", f"Hãy giải thích nội dung: {p[:180]}", "", "**Đáp án gợi ý:**", p, ""]
                n += 1
    else:
        lines += ["## 4. Các điểm quan trọng", ""]
        for f in files:
            points = f.get("summary", {}).get("key_points") or []
            lines += [f"### {f['filename']}", ""]
            if points:
                lines += [f"- {p}" for p in points[:10]]
            else:
                lines.append("- Chưa phát hiện điểm nổi bật rõ ràng.")
            lines.append("")

    lines += ["## 5. Review kế toán / kiểm soát", ""]
    all_risks: List[str] = []
    for f in files:
        all_risks.extend(f.get("summary", {}).get("risks") or [])
    seen = set()
    for risk in all_risks:
        if risk not in seen:
            seen.add(risk)
            lines.append(f"- {risk}")
    lines += [
        "- Không tự động ghi sổ hoặc quyết toán chỉ dựa trên báo cáo này.",
        "- Cần đối chiếu file gốc, chứng từ, chính sách công ty và quy định còn hiệu lực.",
        "",
        "## 6. Checklist hành động đề xuất",
        "",
        "1. Kiểm tra file gốc có đủ trang/bảng/phụ lục không.",
        "2. Đối chiếu các con số, ngày tháng, điều khoản và mã chỉ tiêu quan trọng.",
        "3. Nếu là tài liệu pháp lý, kiểm tra văn bản sửa đổi/bổ sung mới hơn.",
        "4. Nếu là tài liệu kế toán, xác định kỳ kế toán, tài khoản/chỉ tiêu liên quan và người phê duyệt.",
        "5. Lưu báo cáo này như bản nháp hỗ trợ soát xét, không thay thế ý kiến chuyên môn cuối cùng.",
    ]
    return "\n".join(lines).strip() + "\n"


def select_relevant_passages(text: str, query: str = "", limit: int = 8) -> List[str]:
    sents = _sentences(text)
    ranked = sorted(sents, key=lambda s: _score_sentence(s, query), reverse=True)
    out: List[str] = []
    for s in ranked:
        clean = s.strip()
        if clean and clean not in out:
            out.append(clean)
        if len(out) >= limit:
            break
    return out


def build_report_payload(job: Dict[str, Any], files: List[Dict[str, Any]], markdown: str) -> Dict[str, Any]:
    public_files = []
    for f in files:
        public_files.append({
            "filename": f.get("filename"),
            "parser": f.get("parser"),
            "warnings": f.get("warnings") or [],
            "char_count": f.get("char_count"),
            "word_count": f.get("word_count"),
            "profile": f.get("profile") or {},
            "summary": f.get("summary") or {},
        })
    return {
        "version": FILE_REPORT_VERSION,
        "job_id": job.get("job_id"),
        "workspace_id": job.get("workspace_id"),
        "user_id": job.get("user_id"),
        "task_type": job.get("task_type"),
        "output_format": job.get("output_format"),
        "report": {
            "title": job.get("title"),
            "markdown": markdown,
            "preview": markdown[:4000],
        },
        "files": public_files,
    }


def write_report_output_file(job: Dict[str, Any], markdown_text: str, payload: Dict[str, Any]) -> Path:
    fmt = (job.get("output_format") or "docx").lower().strip(".")
    if fmt not in FILE_REPORT_OUTPUT_FORMATS:
        fmt = "docx"
    output_dir = Path(job.get("output_dir") or (_job_dir(job["job_id"]) / "outputs"))
    output_dir.mkdir(parents=True, exist_ok=True)
    name = f"{_slug(job.get('title') or job['job_id'])}_{job['job_id']}.{fmt}"
    path = output_dir / name
    if fmt == "md":
        path.write_text(markdown_text, encoding="utf-8")
    elif fmt == "txt":
        plain = re.sub(r"^#+\s*", "", markdown_text, flags=re.M).replace("**", "").replace("`", "")
        path.write_text(plain, encoding="utf-8")
    elif fmt == "json":
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    elif fmt == "csv":
        with path.open("w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["section", "content"])
            current = "report"
            for line in markdown_text.splitlines():
                if line.startswith("#"):
                    current = line.lstrip("# ").strip()
                elif line.strip():
                    writer.writerow([current, line.strip()])
    elif fmt == "docx":
        _write_docx(path, markdown_text)
    elif fmt == "xlsx":
        _write_xlsx(path, payload, markdown_text)
    elif fmt == "pdf":
        _write_pdf(path, markdown_text)
    return path


def _clean_markdown_inline(text: str) -> str:
    return str(text or "").replace("**", "").replace("__", "").replace("`", "").strip()


def _markdown_table_rows(lines: List[str], start: int) -> tuple[List[List[str]], int]:
    rows: List[List[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [_clean_markdown_inline(cell) for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def _write_docx(path: Path, markdown_text: str) -> None:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Cm, Pt  # type: ignore

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        if style_name in doc.styles:
            doc.styles[style_name].font.name = "Arial"

    lines = markdown_text.splitlines()
    index = 0
    first_heading = True
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|"):
            rows, next_index = _markdown_table_rows(lines, index)
            if rows:
                cols = max(len(row) for row in rows)
                table = doc.add_table(rows=0, cols=cols)
                table.style = "Table Grid"
                for row_index, row in enumerate(rows):
                    cells = table.add_row().cells
                    for col_index in range(cols):
                        cells[col_index].text = row[col_index] if col_index < len(row) else ""
                        for run in cells[col_index].paragraphs[0].runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(9.5)
                            run.bold = row_index == 0
                index = next_index
                continue
        if stripped.startswith("# "):
            paragraph = doc.add_heading(_clean_markdown_inline(stripped[2:]), level=0 if first_heading else 1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if first_heading else WD_ALIGN_PARAGRAPH.LEFT
            first_heading = False
        elif stripped.startswith("## "):
            doc.add_heading(_clean_markdown_inline(stripped[3:]), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(_clean_markdown_inline(stripped[4:]), level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(_clean_markdown_inline(stripped[5:]), level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(_clean_markdown_inline(stripped[2:]), style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(_clean_markdown_inline(re.sub(r"^\d+\.\s+", "", stripped)), style="List Number")
        else:
            doc.add_paragraph(_clean_markdown_inline(stripped))
        index += 1

    footer = section.footer.paragraphs[0]
    footer.text = "Finiip — CTCP IIP Việt Nam | Báo cáo tạo tự động, cần kiểm tra trước khi sử dụng chính thức"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
    doc.save(path)


def _write_xlsx(path: Path, payload: Dict[str, Any], markdown_text: str) -> None:
    import openpyxl  # type: ignore
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side  # type: ignore
    from openpyxl.utils import get_column_letter  # type: ignore

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Tong_quan"
    header_fill = PatternFill("solid", fgColor="365F91")
    sub_fill = PatternFill("solid", fgColor="D9EAF7")
    white_font = Font(color="FFFFFF", bold=True)
    bold_font = Font(bold=True)
    thin = Side(style="thin", color="D9E1F2")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws.merge_cells("A1:F1")
    ws["A1"] = str((payload.get("report") or {}).get("title") or "Báo cáo Finiip")
    ws["A1"].font = Font(size=16, bold=True, color="FFFFFF")
    ws["A1"].fill = header_fill
    ws["A1"].alignment = Alignment(horizontal="center")
    summary_rows = [
        ["Trường", "Giá trị"],
        ["Job ID", payload.get("job_id")],
        ["Nhiệm vụ", payload.get("task_type")],
        ["Workspace", payload.get("workspace_id")],
        ["Chế độ phân tích", payload.get("analysis_mode")],
        ["Số file", len(payload.get("files") or [])],
    ]
    for row in summary_rows:
        ws.append(row)
    for cell in ws[2]:
        cell.fill = sub_fill; cell.font = bold_font; cell.border = border
    for row in ws.iter_rows(min_row=3, max_row=ws.max_row, min_col=1, max_col=2):
        for cell in row:
            cell.border = border
    ws.column_dimensions["A"].width = 24
    ws.column_dimensions["B"].width = 55
    ws.freeze_panes = "A3"

    files_ws = wb.create_sheet("Danh_sach_file")
    headers = ["STT", "Tên file", "Loại tài liệu", "Số văn bản", "Số ký tự", "Số từ", "Cảnh báo"]
    files_ws.append(headers)
    for cell in files_ws[1]:
        cell.fill = header_fill; cell.font = white_font; cell.alignment = Alignment(horizontal="center"); cell.border = border
    for index, f in enumerate(payload.get("files") or [], 1):
        profile = f.get("profile") or {}
        files_ws.append([
            index, f.get("filename"), profile.get("document_type"), profile.get("document_number"),
            f.get("char_count"), f.get("word_count"), "; ".join(f.get("warnings") or []),
        ])
    total_row = files_ws.max_row + 1
    files_ws.cell(total_row, 5, "Tổng")
    files_ws.cell(total_row, 6, f"=SUM(F2:F{max(2, total_row - 1)})")
    files_ws.cell(total_row, 5).font = bold_font
    files_ws.cell(total_row, 6).font = bold_font
    files_ws.auto_filter.ref = f"A1:G{max(1, files_ws.max_row - 1)}"
    files_ws.freeze_panes = "A2"
    widths = [8, 34, 22, 20, 14, 14, 48]
    for idx, width in enumerate(widths, 1):
        files_ws.column_dimensions[get_column_letter(idx)].width = width
    for row in files_ws.iter_rows():
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    points_ws = wb.create_sheet("Diem_quan_trong")
    points_ws.append(["File", "Nhóm", "Nội dung"])
    for cell in points_ws[1]:
        cell.fill = header_fill; cell.font = white_font; cell.border = border
    for f in payload.get("files") or []:
        summary = f.get("summary") or {}
        for point in summary.get("key_points") or []:
            points_ws.append([f.get("filename"), "Điểm quan trọng", point])
        for risk in summary.get("risks") or []:
            points_ws.append([f.get("filename"), "Rủi ro/Lưu ý", risk])
    points_ws.freeze_panes = "A2"
    points_ws.auto_filter.ref = f"A1:C{points_ws.max_row}"
    points_ws.column_dimensions["A"].width = 32
    points_ws.column_dimensions["B"].width = 20
    points_ws.column_dimensions["C"].width = 100
    for row in points_ws.iter_rows():
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    report_ws = wb.create_sheet("Bao_cao")
    report_ws.append(["Dòng", "Nội dung"])
    for cell in report_ws[1]:
        cell.fill = header_fill; cell.font = white_font; cell.border = border
    for line_no, line in enumerate(markdown_text.splitlines()[:20000], 1):
        report_ws.append([line_no, line])
    report_ws.freeze_panes = "A2"
    report_ws.column_dimensions["A"].width = 10
    report_ws.column_dimensions["B"].width = 120
    for row in report_ws.iter_rows():
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    wb.save(path)


def _pdf_escape(text: str) -> str:
    # Minimal PDF fallback. It keeps ASCII best; Vietnamese may be simplified if
    # reportlab is unavailable. DOCX/MD should be preferred for perfect accents.
    ascii_text = unicodedata.normalize("NFKD", text or "").encode("ascii", "ignore").decode("ascii")
    return ascii_text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _write_pdf(path: Path, markdown_text: str) -> None:
    """Write a Unicode PDF when reportlab and a system font are available."""
    try:
        from reportlab.lib import colors  # type: ignore
        from reportlab.lib.enums import TA_CENTER  # type: ignore
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore
        from reportlab.lib.units import mm  # type: ignore
        from reportlab.pdfbase import pdfmetrics  # type: ignore
        from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle  # type: ignore

        font_name = "Helvetica"
        bold_name = "Helvetica-Bold"
        font_candidates = [
            ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
            ("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf", "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf"),
        ]
        for regular_path, bold_path in font_candidates:
            if Path(regular_path).exists():
                pdfmetrics.registerFont(TTFont("FiniipUnicode", regular_path))
                font_name = "FiniipUnicode"
                if Path(bold_path).exists():
                    pdfmetrics.registerFont(TTFont("FiniipUnicodeBold", bold_path))
                    bold_name = "FiniipUnicodeBold"
                else:
                    bold_name = font_name
                break

        doc = SimpleDocTemplate(
            str(path), pagesize=A4, rightMargin=16 * mm, leftMargin=16 * mm,
            topMargin=15 * mm, bottomMargin=16 * mm,
            title="Finiip report",
        )
        styles = getSampleStyleSheet()
        body = ParagraphStyle("FiniipBody", parent=styles["BodyText"], fontName=font_name, fontSize=9.5, leading=13, spaceAfter=5)
        h1 = ParagraphStyle("FiniipH1", parent=styles["Heading1"], fontName=bold_name, fontSize=18, leading=22, alignment=TA_CENTER, textColor=colors.HexColor("#243B64"), spaceAfter=12)
        h2 = ParagraphStyle("FiniipH2", parent=styles["Heading2"], fontName=bold_name, fontSize=13, leading=17, textColor=colors.HexColor("#365F91"), spaceBefore=8, spaceAfter=6)
        h3 = ParagraphStyle("FiniipH3", parent=styles["Heading3"], fontName=bold_name, fontSize=11, leading=15, textColor=colors.HexColor("#365F91"), spaceBefore=6, spaceAfter=4)
        bullet = ParagraphStyle("FiniipBullet", parent=body, leftIndent=12, firstLineIndent=-6, bulletIndent=4)
        story: List[Any] = []
        lines = markdown_text.splitlines()
        index = 0
        while index < len(lines):
            stripped = lines[index].strip()
            if not stripped:
                story.append(Spacer(1, 3))
                index += 1
                continue
            if stripped.startswith("|"):
                rows, next_index = _markdown_table_rows(lines, index)
                if rows:
                    table_data = [[Paragraph(html.escape(cell), body) for cell in row] for row in rows]
                    table = Table(table_data, repeatRows=1, hAlign="LEFT")
                    table.setStyle(TableStyle([
                        ("FONTNAME", (0, 0), (-1, -1), font_name),
                        ("FONTNAME", (0, 0), (-1, 0), bold_name),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF7")),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAB7C4")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]))
                    story.extend([table, Spacer(1, 6)])
                    index = next_index
                    continue
            if stripped.startswith("# "):
                story.append(Paragraph(html.escape(_clean_markdown_inline(stripped[2:])), h1))
            elif stripped.startswith("## "):
                story.append(Paragraph(html.escape(_clean_markdown_inline(stripped[3:])), h2))
            elif stripped.startswith("### "):
                story.append(Paragraph(html.escape(_clean_markdown_inline(stripped[4:])), h3))
            elif stripped.startswith("- "):
                story.append(Paragraph("• " + html.escape(_clean_markdown_inline(stripped[2:])), bullet))
            elif re.match(r"^\d+\.\s+", stripped):
                story.append(Paragraph(html.escape(_clean_markdown_inline(stripped)), bullet))
            else:
                story.append(Paragraph(html.escape(_clean_markdown_inline(stripped)), body))
            index += 1
        doc.build(story)
        return
    except Exception:
        pass

    # Minimal dependency-free fallback. Vietnamese accents may be simplified.
    lines = []
    for raw in markdown_text.splitlines():
        line = re.sub(r"^#+\s*", "", raw).replace("**", "").replace("`", "")
        if not line.strip():
            lines.append("")
        else:
            while len(line) > 90:
                lines.append(line[:90]); line = line[90:]
            lines.append(line)
    pages = [lines[i:i + 48] for i in range(0, len(lines), 48)] or [["Finiip report"]]
    objects: List[bytes] = []
    page_ids = []
    def add_obj(content: bytes) -> int:
        objects.append(content)
        return len(objects)
    font_id = add_obj(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    for page_lines in pages:
        y = 780
        text_ops = ["BT", "/F1 10 Tf", "1 0 0 1 40 800 Tm"]
        for line in page_lines:
            text_ops.append(f"1 0 0 1 40 {y} Tm ({_pdf_escape(line)}) Tj")
            y -= 15
        text_ops.append("ET")
        stream = "\n".join(text_ops).encode("latin-1", "ignore")
        content_id = add_obj(b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream")
        page_id = add_obj(f"<< /Type /Page /Parent 0 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_id} 0 R >>".encode())
        page_ids.append(page_id)
    kids = " ".join(f"{pid} 0 R" for pid in page_ids)
    pages_id = add_obj(f"<< /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >>".encode())
    for pid in page_ids:
        objects[pid - 1] = objects[pid - 1].replace(b"/Parent 0 0 R", f"/Parent {pages_id} 0 R".encode())
    catalog_id = add_obj(f"<< /Type /Catalog /Pages {pages_id} 0 R >>".encode())
    data = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for idx, obj in enumerate(objects, 1):
        offsets.append(len(data))
        data.extend(f"{idx} 0 obj\n".encode()); data.extend(obj); data.extend(b"\nendobj\n")
    xref = len(data)
    data.extend(f"xref\n0 {len(objects)+1}\n0000000000 65535 f \n".encode())
    for off in offsets[1:]:
        data.extend(f"{off:010d} 00000 n \n".encode())
    data.extend(f"trailer << /Size {len(objects)+1} /Root {catalog_id} 0 R >>\nstartxref\n{xref}\n%%EOF".encode())
    path.write_bytes(bytes(data))


def create_and_run_sync(
    *,
    files: List[FileReportInput],
    instruction: str = "",
    question: str = "",
    task_type: str = "auto_report",
    output_format: str = "docx",
    report_style: str = "detailed",
    workspace_id: str = "default",
    user_id: str = "anonymous",
    title: str = "",
) -> Dict[str, Any]:
    job = create_file_report_job(
        files=files,
        instruction=instruction,
        question=question,
        task_type=task_type,
        output_format=output_format,
        report_style=report_style,
        workspace_id=workspace_id,
        user_id=user_id,
        title=title,
        save_inputs=True,
    )
    return run_file_report_job(job["job_id"])


def get_job_status(job_id: str) -> Dict[str, Any]:
    job = _read_job(job_id)
    public = dict(job)
    for f in public.get("files") or []:
        f.pop("stored_path", None)
    public.pop("output_path", None)
    public.pop("output_dir", None)
    return public


def resolve_job_output(job_id: str) -> Dict[str, Any]:
    job = _read_job(job_id)
    if job.get("status") != "done":
        raise FileNotFoundError(f"Job chưa có file tải xuống, status={job.get('status')}")
    path = Path(job.get("output_path") or "")
    if not path.exists() or _job_dir(job_id) not in path.parents:
        raise FileNotFoundError("Không tìm thấy output file")
    return {"path": path, "filename": job.get("output_filename") or path.name, "metadata": job}


def list_file_report_history(workspace_id: str = "", user_id: str = "", limit: int = 50, include_deleted: bool = False) -> Dict[str, Any]:
    store = _load_history()
    out = []
    for item in store.get("items") or []:
        if not include_deleted and item.get("deleted"):
            continue
        if workspace_id and item.get("workspace_id") != workspace_id:
            continue
        if user_id and item.get("user_id") != user_id:
            continue
        out.append(item)
        if len(out) >= max(1, min(limit, 200)):
            break
    return {"version": FILE_REPORT_VERSION, "count": len(out), "items": out}


def delete_file_report_job(job_id: str, hard_delete: bool = False) -> Dict[str, Any]:
    job = _read_job(job_id)
    if hard_delete:
        folder = _job_dir(job_id)
        if folder.exists():
            shutil.rmtree(folder)
        job["deleted"] = True
        job["status"] = "deleted"
        _upsert_history(job)
        return {"deleted": True, "hard_delete": True, "job_id": job_id}
    job["deleted"] = True
    job["status"] = "deleted"
    _write_job(job)
    return {"deleted": True, "hard_delete": False, "job_id": job_id}


def capabilities() -> Dict[str, Any]:
    return {
        "version": FILE_REPORT_VERSION,
        "features": {
            "v68_frontend_api": True,
            "v69_async_jobs": True,
            "v70_history": True,
            "v71_pdf_export": True,
            "v72_multiple_files": True,
            "v110_llm_grounded_report": True,
            "v110_long_file_chunk_selection": True,
            "v110_unicode_pdf": True,
            "v110_professional_docx_xlsx": True,
            "v110_image_ocr": True,
        },
        "endpoints": [
            "POST /ai/v68/file-report/create-sync",
            "POST /ai/v69/file-report/jobs",
            "GET /ai/v69/file-report/jobs/{job_id}",
            "GET /ai/v69/file-report/jobs/{job_id}/download",
            "GET /ai/v70/file-report/history",
            "DELETE /ai/v70/file-report/history/{job_id}",
        ],
        "input_formats": sorted(ALLOWED_EXTENSIONS),
        "output_formats": FILE_REPORT_OUTPUT_FORMATS,
        "task_types": FILE_REPORT_TASK_TYPES,
        "report_styles": FILE_REPORT_STYLES,
        "limits": {"max_file_mb": MAX_FILE_REPORT_MB, "max_files": MAX_FILE_REPORT_FILES},
    }

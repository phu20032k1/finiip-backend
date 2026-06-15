"""V66/V67 RAG helpers for Finiip backend.

Features:
- Read uploaded TXT/MD/CSV/JSON/PDF/DOCX/XLSX files.
- Chunk text for long regulation/procedure documents.
- Save a local JSON RAG store for dev/offline use.
- Optionally save documents/chunks to Supabase Postgres tables rag_documents/rag_chunks.
- Search with lightweight keyword scoring so the app works without paid embeddings.

Supabase requirement:
- Set DATABASE_URL to the Supabase Postgres connection string.
- Run the SQL in docs/V67_SUPABASE_RAG_STORAGE_GUIDE.md first.
"""
from __future__ import annotations

import csv
import hashlib
import io
import json
import math
import os
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

DATA_DIR = Path(__file__).resolve().parents[1] / "data"
UPLOAD_DIR = DATA_DIR / "rag_uploads"
LOCAL_RAG_FILE = DATA_DIR / "v66_rag_store.json"
DATA_DIR.mkdir(exist_ok=True)
UPLOAD_DIR.mkdir(exist_ok=True)

VI_STOPWORDS = {
    "và", "của", "cho", "các", "một", "những", "được", "theo", "trong", "khi", "này", "đó",
    "với", "hoặc", "thì", "là", "có", "không", "phải", "về", "để", "từ", "đến", "trên",
    "dưới", "nếu", "như", "cần", "hỏi", "giúp", "tôi", "tớ", "bạn", "anh", "chị",
}


def load_dotenv_light(path: Optional[Path] = None) -> None:
    """Tiny .env loader so local VS Code runs without python-dotenv."""
    env_path = path or (Path(__file__).resolve().parents[1] / ".env")
    if not env_path.exists():
        return
    for raw in env_path.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


load_dotenv_light()


def now_iso() -> str:
    return datetime.utcnow().isoformat(timespec="seconds") + "Z"


def safe_filename(name: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", name or "upload")[:120]
    return stem or "upload"


def normalize_text(text: str) -> str:
    text = (text or "").lower()
    text = re.sub(r"[^\w\sÀ-ỹđĐ]", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def tokenize(text: str) -> List[str]:
    words = normalize_text(text).split()
    return [w for w in words if len(w) >= 3 and w not in VI_STOPWORDS]


def split_tags(tags: Optional[str | List[str]]) -> List[str]:
    if not tags:
        return []
    if isinstance(tags, list):
        parts = tags
    else:
        parts = re.split(r"[,;\n]+", tags)
    return [p.strip() for p in parts if p and p.strip()]


def load_local_store() -> Dict[str, Any]:
    if not LOCAL_RAG_FILE.exists():
        return {"version": "v66_local_rag", "documents": [], "chunks": []}
    try:
        return json.loads(LOCAL_RAG_FILE.read_text(encoding="utf-8"))
    except Exception:
        return {"version": "v66_local_rag", "documents": [], "chunks": []}


def save_local_store(store: Dict[str, Any]) -> None:
    LOCAL_RAG_FILE.write_text(json.dumps(store, ensure_ascii=False, indent=2), encoding="utf-8")


def _ocr_pil_image(image: Any) -> str:
    try:
        import pytesseract  # type: ignore
        return pytesseract.image_to_string(image, lang=os.getenv("FINIIP_OCR_LANG", "vie+eng")) or ""
    except Exception:
        return ""


def read_image_bytes(raw: bytes) -> str:
    try:
        from PIL import Image  # type: ignore
        image = Image.open(io.BytesIO(raw)).convert("RGB")
        text = _ocr_pil_image(image)
        if not text.strip():
            raise ValueError("OCR không nhận diện được chữ trong ảnh")
        return text
    except Exception as exc:
        raise ValueError(f"Không đọc được ảnh/OCR: {exc}") from exc


def read_pdf_bytes(raw: bytes) -> str:
    """Read text PDFs and OCR scanned pages when PyMuPDF/Tesseract are present."""
    page_texts: List[str] = []
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw))
        for index, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            page_texts.append(f"[TRANG {index}]\n{text}".strip())
    except Exception:
        page_texts = []

    extracted = "\n\n".join(page_texts).strip()
    visible_chars = len(re.sub(r"\s+", "", extracted))
    if visible_chars >= int(os.getenv("FINIIP_PDF_TEXT_MIN_CHARS", "120")):
        return extracted

    # Scanned PDF fallback. PyMuPDF renders pages without requiring poppler.
    try:
        import fitz  # type: ignore
        from PIL import Image  # type: ignore
        doc = fitz.open(stream=raw, filetype="pdf")
        ocr_pages: List[str] = []
        max_pages = int(os.getenv("FINIIP_OCR_MAX_PAGES", "80"))
        for index, page in enumerate(doc, 1):
            if index > max_pages:
                ocr_pages.append(f"[CẢNH BÁO] Chỉ OCR {max_pages} trang đầu.")
                break
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = _ocr_pil_image(image)
            ocr_pages.append(f"[TRANG {index} - OCR]\n{text}".strip())
        ocr_text = "\n\n".join(ocr_pages).strip()
        if ocr_text:
            return ocr_text
    except Exception:
        pass

    if extracted:
        return extracted
    raise ValueError("PDF không có lớp text và OCR chưa đọc được. Hãy kiểm tra Tesseract/PyMuPDF.")


def read_docx_bytes(raw: bytes) -> str:
    try:
        import docx  # type: ignore
        doc = docx.Document(io.BytesIO(raw))
        lines: List[str] = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                lines.append(paragraph.text.strip())
        for table_index, table in enumerate(doc.tables, 1):
            lines.append(f"# BẢNG {table_index}")
            for row in table.rows:
                values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(values):
                    lines.append(" | ".join(values))
        return "\n".join(lines)
    except ImportError as exc:
        raise ValueError("Thiếu thư viện python-docx. Chạy: pip install python-docx") from exc
    except Exception as exc:
        raise ValueError(f"Không đọc được DOCX: {exc}") from exc


def read_xlsx_bytes(raw: bytes) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=False)
        lines: List[str] = []
        max_rows = int(os.getenv("FINIIP_XLSX_MAX_ROWS_PER_SHEET", "20000"))
        max_cols = int(os.getenv("FINIIP_XLSX_MAX_COLS", "200"))
        for ws in wb.worksheets:
            lines.append(f"# SHEET: {ws.title}")
            for row_index, row in enumerate(ws.iter_rows(values_only=False), 1):
                if row_index > max_rows:
                    lines.append(f"[CẢNH BÁO] Sheet {ws.title} chỉ đọc {max_rows} dòng đầu.")
                    break
                values: List[str] = []
                for cell in row[:max_cols]:
                    value = cell.value
                    if value is None:
                        values.append("")
                    elif isinstance(value, str) and value.startswith("="):
                        values.append(f"{cell.coordinate}:{value}")
                    else:
                        values.append(str(value))
                while values and not values[-1]:
                    values.pop()
                if any(values):
                    lines.append(" | ".join(values))
        return "\n".join(lines)
    except Exception as exc:
        raise ValueError(f"Không đọc được XLSX: {exc}") from exc


def read_upload_bytes(filename: str, raw: bytes) -> str:
    ext = Path(filename or "").suffix.lower()
    if ext == ".pdf":
        return read_pdf_bytes(raw)
    if ext == ".docx":
        return read_docx_bytes(raw)
    if ext in {".xlsx", ".xlsm"}:
        return read_xlsx_bytes(raw)
    if ext in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}:
        return read_image_bytes(raw)
    if ext in {".txt", ".md", ".csv", ".json", ".html", ".xml"}:
        return raw.decode("utf-8", errors="ignore")
    # Best-effort fallback for plain text files with unknown extension.
    text = raw.decode("utf-8", errors="ignore")
    if text.strip():
        return text
    raise ValueError(f"Chưa hỗ trợ định dạng file: {ext or filename}")


def chunk_text(text: str, chunk_size: int = 1200, overlap: int = 180) -> List[str]:
    cleaned = re.sub(r"\r\n?", "\n", text or "")
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    if not cleaned:
        return []
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", cleaned) if p.strip()]
    chunks: List[str] = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) + 2 <= chunk_size:
            current = (current + "\n\n" + para).strip()
        else:
            if current:
                chunks.append(current)
            if len(para) <= chunk_size:
                current = para
            else:
                start = 0
                while start < len(para):
                    chunks.append(para[start:start + chunk_size].strip())
                    start += max(1, chunk_size - overlap)
                current = ""
    if current:
        chunks.append(current)
    return chunks


def keyword_score(question: str, content: str) -> float:
    q = tokenize(question)
    c = tokenize(content)
    if not q or not c:
        return 0.0
    qset = set(q)
    c_counts = {w: c.count(w) for w in set(c)}
    score = 0.0
    for w in qset:
        if w in c_counts:
            score += 1.0 + math.log(1 + c_counts[w])
    # phrase boost
    nq = normalize_text(question)
    nc = normalize_text(content)
    for phrase in re.findall(r"\w+(?:\s+\w+){1,4}", nq):
        if len(phrase) >= 8 and phrase in nc:
            score += 2.0
    return round(score, 4)


def save_local_rag_document(
    *,
    title: str,
    content: str,
    filename: str,
    category: str = "general",
    document_type: str = "document",
    source: str = "admin_upload",
    tags: Optional[List[str]] = None,
    storage_path: Optional[str] = None,
) -> Dict[str, Any]:
    store = load_local_store()
    doc_id = str(uuid.uuid4())
    chunks = chunk_text(content)
    document = {
        "id": doc_id,
        "title": title,
        "filename": filename,
        "document_type": document_type,
        "category": category,
        "source": source,
        "tags": tags or [],
        "storage_path": storage_path,
        "created_at": now_iso(),
        "content_length": len(content),
        "chunk_count": len(chunks),
    }
    store.setdefault("documents", []).append(document)
    for i, chunk in enumerate(chunks):
        store.setdefault("chunks", []).append({
            "id": str(uuid.uuid4()),
            "document_id": doc_id,
            "chunk_index": i,
            "title": title,
            "content": chunk,
            "metadata": {
                "filename": filename,
                "category": category,
                "document_type": document_type,
                "source": source,
                "tags": tags or [],
            },
            "created_at": now_iso(),
        })
    save_local_store(store)
    return {"document": document, "chunks": chunks}


def search_local_rag(question: str, limit: int = 5, category: Optional[str] = None) -> Dict[str, Any]:
    store = load_local_store()
    scored = []
    docs_by_id = {d.get("id"): d for d in store.get("documents", [])}
    for chunk in store.get("chunks", []):
        meta = chunk.get("metadata") or {}
        if category and meta.get("category") != category:
            continue
        score = keyword_score(question, chunk.get("title", "") + "\n" + chunk.get("content", ""))
        if score > 0:
            scored.append((score, chunk))
    scored.sort(key=lambda x: x[0], reverse=True)
    results = []
    for score, chunk in scored[: max(1, limit)]:
        doc = docs_by_id.get(chunk.get("document_id"), {})
        results.append({
            "score": score,
            "document_id": chunk.get("document_id"),
            "chunk_id": chunk.get("id"),
            "chunk_index": chunk.get("chunk_index"),
            "title": doc.get("title") or chunk.get("title"),
            "source": doc.get("source") or (chunk.get("metadata") or {}).get("source"),
            "category": (chunk.get("metadata") or {}).get("category"),
            "content": chunk.get("content"),
        })
    answer = build_answer_from_chunks(question, results)
    return {"question": question, "matched": len(results), "results": results, "answer": answer}


def build_answer_from_chunks(question: str, chunks: List[Dict[str, Any]]) -> str:
    if not chunks:
        return "Chưa tìm thấy đoạn tài liệu phù hợp trong RAG. Hãy upload tài liệu liên quan trước hoặc kiểm tra category."
    lines = ["Dựa trên các đoạn tài liệu tìm được, các ý liên quan là:"]
    for idx, item in enumerate(chunks[:5], start=1):
        content = re.sub(r"\s+", " ", item.get("content") or "").strip()
        if len(content) > 700:
            content = content[:700].rstrip() + "..."
        lines.append(f"\n[{idx}] {item.get('title') or 'Tài liệu'} — chunk {item.get('chunk_index')}: {content}")
    lines.append("\nLưu ý: đây là câu trả lời RAG trích theo tài liệu đã upload; cần kiểm tra quy định hiện hành trước khi dùng chính thức.")
    return "\n".join(lines)


def database_url() -> str:
    return os.getenv("DATABASE_URL", "")


def supabase_enabled() -> bool:
    url = database_url()
    return bool(url and url.startswith(("postgresql://", "postgres://")))


def pgvector_literal(text: str, dims: int = 1536) -> str:
    """Deterministic hash fallback embedding, good enough to test vector plumbing."""
    vec = [0.0] * dims
    words = tokenize(text)[:3000]
    if not words:
        return "[" + ",".join("0" for _ in range(dims)) + "]"
    for word in words:
        digest = hashlib.sha256(word.encode("utf-8")).digest()
        idx = int.from_bytes(digest[:4], "big") % dims
        sign = 1.0 if digest[4] % 2 == 0 else -1.0
        vec[idx] += sign
    norm = math.sqrt(sum(v * v for v in vec)) or 1.0
    vec = [round(v / norm, 6) for v in vec]
    return "[" + ",".join(str(v) for v in vec) + "]"


def supabase_status() -> Dict[str, Any]:
    status = {
        "database_url_configured": bool(database_url()),
        "database_url_is_postgres": supabase_enabled(),
        "rag_documents": None,
        "rag_chunks": None,
        "ok": False,
    }
    if not supabase_enabled():
        status["message"] = "DATABASE_URL chưa phải PostgreSQL/Supabase. Local RAG vẫn dùng được."
        return status
    try:
        import psycopg2
        with psycopg2.connect(database_url()) as conn:
            with conn.cursor() as cur:
                cur.execute("select count(*) from rag_documents")
                status["rag_documents"] = cur.fetchone()[0]
                cur.execute("select count(*) from rag_chunks")
                status["rag_chunks"] = cur.fetchone()[0]
        status["ok"] = True
        return status
    except Exception as exc:
        status["error"] = str(exc)
        return status


def save_supabase_rag_document(
    *,
    title: str,
    chunks: List[str],
    filename: str,
    category: str,
    document_type: str,
    source: str,
    tags: Optional[List[str]] = None,
    storage_path: Optional[str] = None,
    uploaded_by: str = "admin",
) -> Dict[str, Any]:
    if not supabase_enabled():
        return {"enabled": False, "saved": False, "message": "DATABASE_URL chưa cấu hình PostgreSQL/Supabase"}
    try:
        import psycopg2
        import psycopg2.extras
        with psycopg2.connect(database_url()) as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(
                    """
                    insert into rag_documents (title, document_type, category, source, storage_path, uploaded_by)
                    values (%s, %s, %s, %s, %s, %s)
                    returning id
                    """,
                    (title, document_type, category, source, storage_path or filename, uploaded_by),
                )
                document_id = str(cur.fetchone()["id"])
                for i, chunk in enumerate(chunks):
                    metadata = {
                        "filename": filename,
                        "category": category,
                        "document_type": document_type,
                        "source": source,
                        "tags": tags or [],
                        "embedding_provider": "v67_hash_fallback",
                    }
                    embedding = pgvector_literal(title + "\n" + chunk)
                    cur.execute(
                        """
                        insert into rag_chunks (document_id, chunk_index, content, metadata, embedding)
                        values (%s, %s, %s, %s::jsonb, %s::vector)
                        """,
                        (document_id, i, chunk, json.dumps(metadata, ensure_ascii=False), embedding),
                    )
            conn.commit()
        return {"enabled": True, "saved": True, "document_id": document_id, "chunks": len(chunks)}
    except Exception as exc:
        return {"enabled": True, "saved": False, "error": str(exc)}


def search_supabase_rag(question: str, limit: int = 5, category: Optional[str] = None) -> Dict[str, Any]:
    if not supabase_enabled():
        return {"enabled": False, **search_local_rag(question, limit=limit, category=category)}
    try:
        import psycopg2
        import psycopg2.extras
        # Use full-text-ish fallback. pgvector hash fallback is mainly for storage plumbing;
        # lexical scoring is more predictable for Vietnamese legal/accounting text.
        with psycopg2.connect(database_url()) as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(
                    """
                    select c.id as chunk_id, c.document_id, c.chunk_index, c.content, c.metadata,
                           d.title, d.source, d.category
                    from rag_chunks c
                    join rag_documents d on d.id = c.document_id
                    where (%s is null or d.category = %s)
                    order by c.created_at desc
                    limit 800
                    """,
                    (category, category),
                )
                rows = [dict(r) for r in cur.fetchall()]
        scored = []
        for row in rows:
            score = keyword_score(question, (row.get("title") or "") + "\n" + (row.get("content") or ""))
            if score > 0:
                scored.append((score, row))
        scored.sort(key=lambda x: x[0], reverse=True)
        results = []
        for score, row in scored[: max(1, limit)]:
            results.append({
                "score": score,
                "document_id": str(row.get("document_id")),
                "chunk_id": str(row.get("chunk_id")),
                "chunk_index": row.get("chunk_index"),
                "title": row.get("title"),
                "source": row.get("source"),
                "category": row.get("category"),
                "content": row.get("content"),
            })
        return {"enabled": True, "question": question, "matched": len(results), "results": results, "answer": build_answer_from_chunks(question, results)}
    except Exception as exc:
        local = search_local_rag(question, limit=limit, category=category)
        local["enabled"] = True
        local["supabase_error"] = str(exc)
        local["fallback"] = "local_rag"
        return local

# ============================================================
# V68/V69/V70/V71 - Production backend RAG upgrades
# ============================================================
# These functions intentionally keep working without OpenAI/Supabase:
# - If OPENAI_API_KEY is configured, use real embeddings / optional answer generation.
# - Otherwise use deterministic hash embeddings so pgvector plumbing can be tested.
# - If DATABASE_URL is missing, fall back to local JSON RAG.

EMBEDDING_DIM = int(os.getenv("RAG_EMBEDDING_DIM", "1536"))
DEFAULT_EMBEDDING_MODEL = os.getenv("OPENAI_EMBEDDING_MODEL", "text-embedding-3-small")
DEFAULT_CHAT_MODEL = os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini")
MAX_UPLOAD_BYTES = int(os.getenv("RAG_MAX_UPLOAD_BYTES", str(20 * 1024 * 1024)))
ALLOWED_RAG_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".pdf", ".docx", ".xlsx", ".xlsm", ".html", ".xml", ".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}


def embedding_provider() -> Dict[str, Any]:
    use_openai = bool(os.getenv("OPENAI_API_KEY"))
    return {
        "provider": "openai" if use_openai else "hash_fallback",
        "model": DEFAULT_EMBEDDING_MODEL if use_openai else "deterministic_hash_v68",
        "dimensions": EMBEDDING_DIM,
        "semantic_quality": "high" if use_openai else "test_only",
    }


def embedding_literal(text: str) -> str:
    """Return a pgvector literal like [0.1,0.2,...]."""
    if os.getenv("OPENAI_API_KEY"):
        try:
            from openai import OpenAI  # type: ignore
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            kwargs: Dict[str, Any] = {"model": DEFAULT_EMBEDDING_MODEL, "input": (text or "")[:24000]}
            # text-embedding-3-small supports dimensions. Other models may not, so fallback on error.
            if DEFAULT_EMBEDDING_MODEL.startswith("text-embedding-3"):
                kwargs["dimensions"] = EMBEDDING_DIM
            resp = client.embeddings.create(**kwargs)
            vec = resp.data[0].embedding
            return "[" + ",".join(str(round(float(v), 8)) for v in vec) + "]"
        except Exception:
            # Never block upload because external embedding failed; fallback still stores a vector.
            pass
    return pgvector_literal(text, dims=EMBEDDING_DIM)


def validate_rag_file(filename: str, size_bytes: int) -> Dict[str, Any]:
    ext = Path(filename or "").suffix.lower()
    if ext and ext not in ALLOWED_RAG_EXTENSIONS:
        return {"ok": False, "error": f"Định dạng {ext} chưa được phép upload vào RAG", "allowed": sorted(ALLOWED_RAG_EXTENSIONS)}
    if size_bytes > MAX_UPLOAD_BYTES:
        return {"ok": False, "error": f"File vượt quá giới hạn {MAX_UPLOAD_BYTES} bytes", "max_upload_bytes": MAX_UPLOAD_BYTES}
    return {"ok": True, "extension": ext or "unknown", "size_bytes": size_bytes}


def _pg_conn():
    import psycopg2
    return psycopg2.connect(database_url())


def ensure_supabase_rag_schema() -> Dict[str, Any]:
    """Create/upgrade minimal Supabase Postgres schema for RAG."""
    if not supabase_enabled():
        return {"ok": False, "enabled": False, "message": "DATABASE_URL chưa cấu hình PostgreSQL/Supabase"}
    statements = []
    try:
        with _pg_conn() as conn:
            with conn.cursor() as cur:
                try:
                    cur.execute("create extension if not exists vector")
                    statements.append("vector_extension_ok")
                except Exception as exc:
                    # Some Supabase setups require enabling vector in dashboard first.
                    conn.rollback()
                    statements.append(f"vector_extension_warning: {exc}")
                cur.execute(
                    """
                    create table if not exists rag_documents (
                      id uuid primary key default gen_random_uuid(),
                      title text not null,
                      document_type text,
                      category text,
                      source text,
                      storage_path text,
                      uploaded_by text,
                      created_at timestamptz default now(),
                      updated_at timestamptz default now()
                    )
                    """
                )
                cur.execute("alter table rag_documents add column if not exists updated_at timestamptz default now()")
                cur.execute(
                    """
                    create table if not exists rag_chunks (
                      id uuid primary key default gen_random_uuid(),
                      document_id uuid references rag_documents(id) on delete cascade,
                      chunk_index int not null,
                      content text not null,
                      metadata jsonb default '{}'::jsonb,
                      embedding vector(1536),
                      created_at timestamptz default now()
                    )
                    """
                )
                cur.execute("create index if not exists rag_documents_category_idx on rag_documents(category)")
                cur.execute("create index if not exists rag_chunks_document_id_idx on rag_chunks(document_id)")
                cur.execute("create index if not exists rag_chunks_metadata_gin_idx on rag_chunks using gin(metadata)")
                try:
                    cur.execute(
                        """
                        create index if not exists rag_chunks_embedding_idx
                        on rag_chunks using ivfflat (embedding vector_cosine_ops) with (lists = 100)
                        """
                    )
                    statements.append("vector_index_ok")
                except Exception as exc:
                    conn.rollback()
                    statements.append(f"vector_index_warning: {exc}")
                conn.commit()
        return {"ok": True, "enabled": True, "statements": statements, "embedding": embedding_provider()}
    except Exception as exc:
        return {"ok": False, "enabled": True, "error": str(exc), "statements": statements}


# Override V67 save with richer metadata/embedding provider.
def save_supabase_rag_document(
    *,
    title: str,
    chunks: List[str],
    filename: str,
    category: str,
    document_type: str,
    source: str,
    tags: Optional[List[str]] = None,
    storage_path: Optional[str] = None,
    uploaded_by: str = "admin",
) -> Dict[str, Any]:
    if not supabase_enabled():
        return {"enabled": False, "saved": False, "message": "DATABASE_URL chưa cấu hình PostgreSQL/Supabase"}
    schema = ensure_supabase_rag_schema()
    if not schema.get("ok"):
        return {"enabled": True, "saved": False, "schema": schema}
    try:
        import psycopg2.extras
        provider = embedding_provider()
        with _pg_conn() as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(
                    """
                    insert into rag_documents (title, document_type, category, source, storage_path, uploaded_by, updated_at)
                    values (%s, %s, %s, %s, %s, %s, now())
                    returning id
                    """,
                    (title, document_type, category, source, storage_path or filename, uploaded_by),
                )
                document_id = str(cur.fetchone()["id"])
                for i, chunk in enumerate(chunks):
                    metadata = {
                        "filename": filename,
                        "category": category,
                        "document_type": document_type,
                        "source": source,
                        "tags": tags or [],
                        "embedding_provider": provider["provider"],
                        "embedding_model": provider["model"],
                    }
                    emb = embedding_literal(title + "\n" + chunk)
                    cur.execute(
                        """
                        insert into rag_chunks (document_id, chunk_index, content, metadata, embedding)
                        values (%s, %s, %s, %s::jsonb, %s::vector)
                        """,
                        (document_id, i, chunk, json.dumps(metadata, ensure_ascii=False), emb),
                    )
            conn.commit()
        return {"enabled": True, "saved": True, "document_id": document_id, "chunks": len(chunks), "embedding": provider}
    except Exception as exc:
        return {"enabled": True, "saved": False, "error": str(exc)}


def _supabase_rows_for_search(question: str, limit: int, category: Optional[str] = None) -> Tuple[List[Dict[str, Any]], Optional[str]]:
    """Try vector search first; fallback to newest rows for lexical rerank."""
    import psycopg2.extras
    vec = embedding_literal(question)
    try:
        with _pg_conn() as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(
                    """
                    select c.id as chunk_id, c.document_id, c.chunk_index, c.content, c.metadata,
                           d.title, d.source, d.category, d.document_type,
                           (c.embedding <=> %s::vector) as vector_distance
                    from rag_chunks c
                    join rag_documents d on d.id = c.document_id
                    where (%s is null or d.category = %s)
                    order by c.embedding <=> %s::vector
                    limit %s
                    """,
                    (vec, category, category, vec, max(20, limit * 8)),
                )
                return [dict(r) for r in cur.fetchall()], None
    except Exception as exc:
        vector_error = str(exc)
    with _pg_conn() as conn:
        with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(
                """
                select c.id as chunk_id, c.document_id, c.chunk_index, c.content, c.metadata,
                       d.title, d.source, d.category, d.document_type,
                       null::float as vector_distance
                from rag_chunks c
                join rag_documents d on d.id = c.document_id
                where (%s is null or d.category = %s)
                order by c.created_at desc
                limit 1000
                """,
                (category, category),
            )
            return [dict(r) for r in cur.fetchall()], vector_error


# Override V67 search with vector + lexical hybrid scoring.
def search_supabase_rag(question: str, limit: int = 5, category: Optional[str] = None) -> Dict[str, Any]:
    if not supabase_enabled():
        return {"enabled": False, **search_local_rag(question, limit=limit, category=category), "mode": "local_keyword"}
    try:
        rows, vector_error = _supabase_rows_for_search(question, limit, category)
        scored: List[Tuple[float, Dict[str, Any]]] = []
        for row in rows:
            lexical = keyword_score(question, (row.get("title") or "") + "\n" + (row.get("content") or ""))
            vd = row.get("vector_distance")
            semantic = 0.0 if vd is None else max(0.0, 1.0 - float(vd)) * 10.0
            score = round((semantic * 0.65) + (lexical * 0.35), 4)
            # keep vector top results even if lexical is low; helpful for real OpenAI embeddings.
            if score > 0 or vd is not None:
                scored.append((score, row))
        scored.sort(key=lambda x: x[0], reverse=True)
        results = []
        for score, row in scored[: max(1, limit)]:
            results.append({
                "score": score,
                "vector_distance": None if row.get("vector_distance") is None else float(row.get("vector_distance")),
                "document_id": str(row.get("document_id")),
                "chunk_id": str(row.get("chunk_id")),
                "chunk_index": row.get("chunk_index"),
                "title": row.get("title"),
                "source": row.get("source"),
                "category": row.get("category"),
                "document_type": row.get("document_type"),
                "content": row.get("content"),
            })
        return {
            "enabled": True,
            "mode": "hybrid_vector_keyword",
            "embedding": embedding_provider(),
            "vector_error": vector_error,
            "question": question,
            "matched": len(results),
            "results": results,
            "answer": build_answer_from_chunks(question, results),
        }
    except Exception as exc:
        local = search_local_rag(question, limit=limit, category=category)
        local.update({"enabled": True, "supabase_error": str(exc), "fallback": "local_rag", "mode": "local_keyword"})
        return local


def rag_health() -> Dict[str, Any]:
    local = load_local_store()
    status = supabase_status()
    schema = ensure_supabase_rag_schema() if supabase_enabled() else {"ok": False, "enabled": False}
    return {
        "ok": bool(status.get("ok")) or not supabase_enabled(),
        "supabase": status,
        "schema": schema,
        "embedding": embedding_provider(),
        "limits": {"max_upload_bytes": MAX_UPLOAD_BYTES, "allowed_extensions": sorted(ALLOWED_RAG_EXTENSIONS)},
        "local_store": {
            "documents": len(local.get("documents", [])),
            "chunks": len(local.get("chunks", [])),
            "path": str(LOCAL_RAG_FILE),
        },
    }


def list_rag_documents(limit: int = 50, offset: int = 0, category: Optional[str] = None, source: Optional[str] = None) -> Dict[str, Any]:
    if not supabase_enabled():
        docs = load_local_store().get("documents", [])
        if category:
            docs = [d for d in docs if d.get("category") == category]
        if source:
            docs = [d for d in docs if d.get("source") == source]
        return {"enabled": False, "total": len(docs), "documents": docs[offset: offset + limit]}
    try:
        import psycopg2.extras
        with _pg_conn() as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(
                    """
                    select d.id, d.title, d.document_type, d.category, d.source, d.storage_path,
                           d.uploaded_by, d.created_at, d.updated_at, count(c.id) as chunk_count
                    from rag_documents d
                    left join rag_chunks c on c.document_id = d.id
                    where (%s is null or d.category = %s) and (%s is null or d.source = %s)
                    group by d.id
                    order by d.created_at desc
                    limit %s offset %s
                    """,
                    (category, category, source, source, limit, offset),
                )
                docs = [dict(r) for r in cur.fetchall()]
                cur.execute(
                    "select count(*) from rag_documents where (%s is null or category = %s) and (%s is null or source = %s)",
                    (category, category, source, source),
                )
                total = int(cur.fetchone()[0])
        for d in docs:
            d["id"] = str(d["id"])
            if d.get("created_at"):
                d["created_at"] = d["created_at"].isoformat()
            if d.get("updated_at"):
                d["updated_at"] = d["updated_at"].isoformat()
        return {"enabled": True, "total": total, "documents": docs}
    except Exception as exc:
        return {"enabled": True, "error": str(exc), "documents": []}


def get_rag_document(document_id: str, include_chunks: bool = True, chunk_limit: int = 200) -> Dict[str, Any]:
    if not supabase_enabled():
        store = load_local_store()
        doc = next((d for d in store.get("documents", []) if str(d.get("id")) == document_id), None)
        chunks = [c for c in store.get("chunks", []) if str(c.get("document_id")) == document_id][:chunk_limit]
        return {"enabled": False, "document": doc, "chunks": chunks if include_chunks else []}
    try:
        import psycopg2.extras
        with _pg_conn() as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute("select * from rag_documents where id = %s", (document_id,))
                doc = cur.fetchone()
                if not doc:
                    return {"enabled": True, "document": None, "chunks": []}
                chunks: List[Dict[str, Any]] = []
                if include_chunks:
                    cur.execute(
                        "select id, chunk_index, content, metadata, created_at from rag_chunks where document_id = %s order by chunk_index limit %s",
                        (document_id, chunk_limit),
                    )
                    chunks = [dict(r) for r in cur.fetchall()]
        doc = dict(doc)
        doc["id"] = str(doc["id"])
        for key in ["created_at", "updated_at"]:
            if doc.get(key):
                doc[key] = doc[key].isoformat()
        for c in chunks:
            c["id"] = str(c["id"])
            if c.get("created_at"):
                c["created_at"] = c["created_at"].isoformat()
        return {"enabled": True, "document": doc, "chunks": chunks}
    except Exception as exc:
        return {"enabled": True, "error": str(exc), "document": None, "chunks": []}


def delete_rag_document(document_id: str) -> Dict[str, Any]:
    if not supabase_enabled():
        store = load_local_store()
        before_docs = len(store.get("documents", []))
        before_chunks = len(store.get("chunks", []))
        store["documents"] = [d for d in store.get("documents", []) if str(d.get("id")) != document_id]
        store["chunks"] = [c for c in store.get("chunks", []) if str(c.get("document_id")) != document_id]
        save_local_store(store)
        return {"enabled": False, "deleted_documents": before_docs - len(store["documents"]), "deleted_chunks": before_chunks - len(store["chunks"])}
    try:
        with _pg_conn() as conn:
            with conn.cursor() as cur:
                cur.execute("delete from rag_documents where id = %s", (document_id,))
                deleted = cur.rowcount
            conn.commit()
        return {"enabled": True, "deleted": deleted > 0, "deleted_documents": deleted}
    except Exception as exc:
        return {"enabled": True, "deleted": False, "error": str(exc)}


def reindex_rag_document(document_id: str, chunk_size: int = 1200, overlap: int = 180) -> Dict[str, Any]:
    doc_data = get_rag_document(document_id, include_chunks=False)
    doc = doc_data.get("document")
    if not doc:
        return {"ok": False, "error": "Không tìm thấy document"}
    storage_path = doc.get("storage_path") or ""
    path = Path(storage_path)
    if not path.exists():
        return {"ok": False, "error": f"Không tìm thấy file gốc để reindex: {storage_path}"}
    raw = path.read_bytes()
    text = read_upload_bytes(path.name, raw)
    chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap)
    if not chunks:
        return {"ok": False, "error": "File gốc không có nội dung để chunk"}
    if not supabase_enabled():
        return {"ok": False, "error": "Reindex theo document_id hiện chỉ hỗ trợ Supabase; local có thể upload lại file."}
    try:
        with _pg_conn() as conn:
            with conn.cursor() as cur:
                cur.execute("delete from rag_chunks where document_id = %s", (document_id,))
                provider = embedding_provider()
                for i, chunk in enumerate(chunks):
                    metadata = {
                        "filename": path.name,
                        "category": doc.get("category"),
                        "document_type": doc.get("document_type"),
                        "source": doc.get("source"),
                        "reindexed_at": now_iso(),
                        "embedding_provider": provider["provider"],
                        "embedding_model": provider["model"],
                    }
                    cur.execute(
                        """
                        insert into rag_chunks (document_id, chunk_index, content, metadata, embedding)
                        values (%s, %s, %s, %s::jsonb, %s::vector)
                        """,
                        (document_id, i, chunk, json.dumps(metadata, ensure_ascii=False), embedding_literal((doc.get("title") or "") + "\n" + chunk)),
                    )
                cur.execute("update rag_documents set updated_at = now() where id = %s", (document_id,))
            conn.commit()
        return {"ok": True, "document_id": document_id, "chunks": len(chunks), "embedding": embedding_provider()}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def build_rag_answer(question: str, results: List[Dict[str, Any]], style: str = "detailed", use_llm: bool = False) -> Dict[str, Any]:
    citations = []
    for i, r in enumerate(results, start=1):
        citations.append({
            "ref": f"[{i}]",
            "document_id": r.get("document_id"),
            "chunk_id": r.get("chunk_id"),
            "title": r.get("title"),
            "chunk_index": r.get("chunk_index"),
            "score": r.get("score"),
        })
    context = "\n\n".join(
        f"[{i}] {r.get('title')} / chunk {r.get('chunk_index')}\n{(r.get('content') or '')[:2500]}"
        for i, r in enumerate(results[:8], start=1)
    )
    if use_llm and os.getenv("OPENAI_API_KEY") and results:
        try:
            from openai import OpenAI  # type: ignore
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            system = (
                "Bạn là trợ lý RAG cho tài liệu luật, thông tư, quy trình và nghiệp vụ. "
                "Chỉ dựa vào CONTEXT. Nếu thiếu căn cứ, nói rõ chưa đủ tài liệu. "
                "Trả lời tiếng Việt, có các mục: Kết luận, Căn cứ, Các bước/giải thích, Lưu ý. "
                "Luôn trích nguồn dạng [1], [2] theo context."
            )
            user = f"CÂU HỎI:\n{question}\n\nCONTEXT:\n{context}"
            resp = client.chat.completions.create(
                model=DEFAULT_CHAT_MODEL,
                messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
                temperature=0.2,
            )
            answer = resp.choices[0].message.content or ""
            return {"answer": answer, "mode": "llm_grounded", "citations": citations}
        except Exception as exc:
            # Fall through to extractive answer but expose the error for debugging.
            llm_error = str(exc)
    else:
        llm_error = None
    if not results:
        return {
            "answer": "Chưa tìm thấy căn cứ phù hợp trong kho RAG. Hãy upload thêm tài liệu đúng category hoặc kiểm tra câu hỏi/từ khóa.",
            "mode": "extractive_no_match",
            "citations": [],
        }
    lines = ["Kết luận tạm thời dựa trên tài liệu đã upload:"]
    for i, r in enumerate(results[:5], start=1):
        text = re.sub(r"\s+", " ", r.get("content") or "").strip()
        lines.append(f"\n[{i}] {r.get('title') or 'Tài liệu'} — chunk {r.get('chunk_index')}: {text[:900]}{'...' if len(text) > 900 else ''}")
    lines.append("\nCăn cứ: " + ", ".join(c["ref"] for c in citations[:5]))
    lines.append("Lưu ý: bản extractive này chưa gọi LLM; đặt OPENAI_API_KEY để sinh câu trả lời diễn giải tốt hơn.")
    if llm_error:
        lines.append(f"LLM fallback reason: {llm_error}")
    return {"answer": "\n".join(lines), "mode": "extractive_grounded", "citations": citations}


def answer_rag_question(question: str, limit: int = 6, category: Optional[str] = None, style: str = "detailed", use_llm: bool = False) -> Dict[str, Any]:
    search = search_supabase_rag(question, limit=limit, category=category)
    results = search.get("results", [])
    answer = build_rag_answer(question, results, style=style, use_llm=use_llm)
    return {
        "question": question,
        "category": category,
        "search_mode": search.get("mode"),
        "matched": search.get("matched", 0),
        "embedding": search.get("embedding") or embedding_provider(),
        "answer": answer.get("answer"),
        "answer_mode": answer.get("mode"),
        "citations": answer.get("citations", []),
        "sources": results,
        "warnings": [w for w in [search.get("supabase_error"), search.get("vector_error")] if w],
    }


# ============================================================
# V81 - RAG Sources & Document Metadata Management
# ============================================================
# Backend only. Frontend can be built separately against these endpoints.
# Adds legal/procedure metadata, workspace/user fields, richer source objects,
# and keeps backward compatibility with older V66/V70 callers.

V81_DOCUMENT_METADATA_FIELDS = [
    "filename", "document_number", "issued_date", "effective_date", "authority",
    "status", "version", "workspace_id", "user_id", "language", "jurisdiction",
    "tags", "source", "category", "document_type",
]


def build_document_metadata(
    *,
    filename: str,
    category: str = "general",
    document_type: str = "document",
    source: str = "admin_upload",
    tags: Optional[List[str]] = None,
    document_number: Optional[str] = None,
    issued_date: Optional[str] = None,
    effective_date: Optional[str] = None,
    authority: Optional[str] = None,
    status: str = "active",
    version: Optional[str] = None,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    language: str = "vi",
    jurisdiction: Optional[str] = None,
    extra_metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    metadata: Dict[str, Any] = {
        "filename": filename,
        "category": category,
        "document_type": document_type,
        "source": source,
        "tags": tags or [],
        "document_number": document_number or None,
        "issued_date": issued_date or None,
        "effective_date": effective_date or None,
        "authority": authority or None,
        "status": status or "active",
        "version": version or None,
        "workspace_id": workspace_id or None,
        "user_id": user_id or None,
        "language": language or "vi",
        "jurisdiction": jurisdiction or None,
    }
    if extra_metadata:
        metadata.update({k: v for k, v in extra_metadata.items() if v is not None})
    return metadata


def _json_or_none(value: Optional[str | Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    if not value:
        return None
    if isinstance(value, dict):
        return value
    try:
        parsed = json.loads(value)
        return parsed if isinstance(parsed, dict) else {"raw": parsed}
    except Exception:
        return {"note": str(value)}


def ensure_supabase_rag_schema() -> Dict[str, Any]:  # type: ignore[override]
    """Create/upgrade Supabase schema for cited RAG + document metadata."""
    if not supabase_enabled():
        return {"ok": False, "enabled": False, "message": "DATABASE_URL chưa cấu hình PostgreSQL/Supabase"}
    statements: List[str] = []
    try:
        with _pg_conn() as conn:
            with conn.cursor() as cur:
                try:
                    cur.execute("create extension if not exists vector")
                    statements.append("vector_extension_ok")
                except Exception as exc:
                    conn.rollback()
                    statements.append(f"vector_extension_warning: {exc}")
                cur.execute(
                    """
                    create table if not exists rag_documents (
                      id uuid primary key default gen_random_uuid(),
                      title text not null,
                      document_type text,
                      category text,
                      source text,
                      storage_path text,
                      uploaded_by text,
                      created_at timestamptz default now(),
                      updated_at timestamptz default now()
                    )
                    """
                )
                for stmt in [
                    "alter table rag_documents add column if not exists filename text",
                    "alter table rag_documents add column if not exists document_number text",
                    "alter table rag_documents add column if not exists issued_date date",
                    "alter table rag_documents add column if not exists effective_date date",
                    "alter table rag_documents add column if not exists authority text",
                    "alter table rag_documents add column if not exists status text default 'active'",
                    "alter table rag_documents add column if not exists version text",
                    "alter table rag_documents add column if not exists workspace_id text",
                    "alter table rag_documents add column if not exists user_id text",
                    "alter table rag_documents add column if not exists language text default 'vi'",
                    "alter table rag_documents add column if not exists jurisdiction text",
                    "alter table rag_documents add column if not exists tags jsonb default '[]'::jsonb",
                    "alter table rag_documents add column if not exists metadata jsonb default '{}'::jsonb",
                    "alter table rag_documents add column if not exists content_hash text",
                    "alter table rag_documents add column if not exists updated_at timestamptz default now()",
                ]:
                    cur.execute(stmt)
                cur.execute(
                    """
                    create table if not exists rag_chunks (
                      id uuid primary key default gen_random_uuid(),
                      document_id uuid references rag_documents(id) on delete cascade,
                      chunk_index int not null,
                      content text not null,
                      metadata jsonb default '{}'::jsonb,
                      embedding vector(1536),
                      created_at timestamptz default now()
                    )
                    """
                )
                cur.execute("alter table rag_chunks add column if not exists token_count int")
                cur.execute("alter table rag_chunks add column if not exists page_start int")
                cur.execute("alter table rag_chunks add column if not exists page_end int")
                cur.execute("create index if not exists rag_documents_category_idx on rag_documents(category)")
                cur.execute("create index if not exists rag_documents_source_idx on rag_documents(source)")
                cur.execute("create index if not exists rag_documents_status_idx on rag_documents(status)")
                cur.execute("create index if not exists rag_documents_workspace_idx on rag_documents(workspace_id)")
                cur.execute("create index if not exists rag_documents_user_idx on rag_documents(user_id)")
                cur.execute("create index if not exists rag_documents_metadata_gin_idx on rag_documents using gin(metadata)")
                cur.execute("create index if not exists rag_documents_tags_gin_idx on rag_documents using gin(tags)")
                cur.execute("create index if not exists rag_chunks_document_id_idx on rag_chunks(document_id)")
                cur.execute("create index if not exists rag_chunks_metadata_gin_idx on rag_chunks using gin(metadata)")
                try:
                    cur.execute(
                        """
                        create index if not exists rag_chunks_embedding_idx
                        on rag_chunks using ivfflat (embedding vector_cosine_ops) with (lists = 100)
                        """
                    )
                    statements.append("vector_index_ok")
                except Exception as exc:
                    conn.rollback()
                    statements.append(f"vector_index_warning: {exc}")
                conn.commit()
        return {"ok": True, "enabled": True, "version": "v81", "statements": statements, "embedding": embedding_provider()}
    except Exception as exc:
        return {"ok": False, "enabled": True, "version": "v81", "error": str(exc), "statements": statements}


def save_local_rag_document(  # type: ignore[override]
    *,
    title: str,
    content: str,
    filename: str,
    category: str = "general",
    document_type: str = "document",
    source: str = "admin_upload",
    tags: Optional[List[str]] = None,
    storage_path: Optional[str] = None,
    document_number: Optional[str] = None,
    issued_date: Optional[str] = None,
    effective_date: Optional[str] = None,
    authority: Optional[str] = None,
    status: str = "active",
    version: Optional[str] = None,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    language: str = "vi",
    jurisdiction: Optional[str] = None,
    extra_metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    store = load_local_store()
    doc_id = str(uuid.uuid4())
    chunks = chunk_text(content)
    metadata = build_document_metadata(
        filename=filename,
        category=category,
        document_type=document_type,
        source=source,
        tags=tags,
        document_number=document_number,
        issued_date=issued_date,
        effective_date=effective_date,
        authority=authority,
        status=status,
        version=version,
        workspace_id=workspace_id,
        user_id=user_id,
        language=language,
        jurisdiction=jurisdiction,
        extra_metadata=extra_metadata,
    )
    document = {
        "id": doc_id,
        "title": title,
        "filename": filename,
        "document_type": document_type,
        "category": category,
        "source": source,
        "tags": tags or [],
        "storage_path": storage_path,
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "content_length": len(content),
        "content_hash": hashlib.sha256((content or "").encode("utf-8")).hexdigest(),
        "chunk_count": len(chunks),
        "metadata": metadata,
        **{k: metadata.get(k) for k in ["document_number", "issued_date", "effective_date", "authority", "status", "version", "workspace_id", "user_id", "language", "jurisdiction"]},
    }
    store.setdefault("documents", []).append(document)
    for i, chunk in enumerate(chunks):
        chunk_meta = {**metadata, "chunk_index": i, "token_count": len(tokenize(chunk))}
        store.setdefault("chunks", []).append({
            "id": str(uuid.uuid4()),
            "document_id": doc_id,
            "chunk_index": i,
            "title": title,
            "content": chunk,
            "metadata": chunk_meta,
            "token_count": len(tokenize(chunk)),
            "created_at": now_iso(),
        })
    save_local_store(store)
    return {"document": document, "chunks": chunks}


def save_supabase_rag_document(  # type: ignore[override]
    *,
    title: str,
    chunks: List[str],
    filename: str,
    category: str,
    document_type: str,
    source: str,
    tags: Optional[List[str]] = None,
    storage_path: Optional[str] = None,
    uploaded_by: str = "admin",
    document_number: Optional[str] = None,
    issued_date: Optional[str] = None,
    effective_date: Optional[str] = None,
    authority: Optional[str] = None,
    status: str = "active",
    version: Optional[str] = None,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    language: str = "vi",
    jurisdiction: Optional[str] = None,
    extra_metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    if not supabase_enabled():
        return {"enabled": False, "saved": False, "message": "DATABASE_URL chưa cấu hình PostgreSQL/Supabase"}
    schema = ensure_supabase_rag_schema()
    if not schema.get("ok"):
        return {"enabled": True, "saved": False, "schema": schema}
    try:
        import psycopg2.extras
        provider = embedding_provider()
        metadata = build_document_metadata(
            filename=filename,
            category=category,
            document_type=document_type,
            source=source,
            tags=tags,
            document_number=document_number,
            issued_date=issued_date,
            effective_date=effective_date,
            authority=authority,
            status=status,
            version=version,
            workspace_id=workspace_id,
            user_id=user_id,
            language=language,
            jurisdiction=jurisdiction,
            extra_metadata=extra_metadata,
        )
        content_hash = hashlib.sha256("\n\n".join(chunks).encode("utf-8")).hexdigest()
        with _pg_conn() as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(
                    """
                    insert into rag_documents (
                        title, filename, document_type, category, source, storage_path, uploaded_by,
                        document_number, issued_date, effective_date, authority, status, version,
                        workspace_id, user_id, language, jurisdiction, tags, metadata, content_hash, updated_at
                    ) values (
                        %s, %s, %s, %s, %s, %s, %s,
                        %s, nullif(%s, '')::date, nullif(%s, '')::date, %s, %s, %s,
                        %s, %s, %s, %s, %s::jsonb, %s::jsonb, %s, now()
                    )
                    returning id
                    """,
                    (
                        title, filename, document_type, category, source, storage_path or filename, uploaded_by,
                        document_number, issued_date, effective_date, authority, status, version,
                        workspace_id, user_id, language, jurisdiction,
                        json.dumps(tags or [], ensure_ascii=False),
                        json.dumps(metadata, ensure_ascii=False),
                        content_hash,
                    ),
                )
                document_id = str(cur.fetchone()["id"])
                for i, chunk in enumerate(chunks):
                    chunk_meta = {
                        **metadata,
                        "chunk_index": i,
                        "token_count": len(tokenize(chunk)),
                        "embedding_provider": provider["provider"],
                        "embedding_model": provider["model"],
                    }
                    emb = embedding_literal(title + "\n" + chunk)
                    cur.execute(
                        """
                        insert into rag_chunks (document_id, chunk_index, content, metadata, embedding, token_count)
                        values (%s, %s, %s, %s::jsonb, %s::vector, %s)
                        """,
                        (document_id, i, chunk, json.dumps(chunk_meta, ensure_ascii=False), emb, chunk_meta["token_count"]),
                    )
            conn.commit()
        return {"enabled": True, "saved": True, "document_id": document_id, "chunks": len(chunks), "metadata": metadata, "embedding": provider}
    except Exception as exc:
        return {"enabled": True, "saved": False, "error": str(exc)}


def _source_from_row(row: Dict[str, Any], score: float) -> Dict[str, Any]:
    meta = row.get("metadata") or {}
    if isinstance(meta, str):
        try:
            meta = json.loads(meta)
        except Exception:
            meta = {}
    return {
        "score": score,
        "vector_distance": None if row.get("vector_distance") is None else float(row.get("vector_distance")),
        "document_id": str(row.get("document_id")),
        "chunk_id": str(row.get("chunk_id")),
        "chunk_index": row.get("chunk_index"),
        "page_start": row.get("page_start") or meta.get("page_start"),
        "page_end": row.get("page_end") or meta.get("page_end"),
        "title": row.get("title"),
        "filename": row.get("filename") or meta.get("filename"),
        "source": row.get("source") or meta.get("source"),
        "category": row.get("category") or meta.get("category"),
        "document_type": row.get("document_type") or meta.get("document_type"),
        "document_number": row.get("document_number") or meta.get("document_number"),
        "issued_date": str(row.get("issued_date")) if row.get("issued_date") else meta.get("issued_date"),
        "effective_date": str(row.get("effective_date")) if row.get("effective_date") else meta.get("effective_date"),
        "authority": row.get("authority") or meta.get("authority"),
        "status": row.get("status") or meta.get("status"),
        "version": row.get("version") or meta.get("version"),
        "workspace_id": row.get("workspace_id") or meta.get("workspace_id"),
        "user_id": row.get("user_id") or meta.get("user_id"),
        "tags": row.get("tags") or meta.get("tags") or [],
        "metadata": meta,
        "content": row.get("content"),
    }


def _supabase_rows_for_search(question: str, limit: int, category: Optional[str] = None) -> Tuple[List[Dict[str, Any]], Optional[str]]:  # type: ignore[override]
    import psycopg2.extras
    vec = embedding_literal(question)
    select_sql = """
        select c.id as chunk_id, c.document_id, c.chunk_index, c.content, c.metadata,
               c.page_start, c.page_end,
               d.title, d.filename, d.source, d.category, d.document_type, d.document_number,
               d.issued_date, d.effective_date, d.authority, d.status, d.version,
               d.workspace_id, d.user_id, d.tags,
               (c.embedding <=> %s::vector) as vector_distance
        from rag_chunks c
        join rag_documents d on d.id = c.document_id
        where (%s is null or d.category = %s)
    """
    try:
        with _pg_conn() as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(select_sql + " order by c.embedding <=> %s::vector limit %s", (vec, category, category, vec, max(20, limit * 8)))
                return [dict(r) for r in cur.fetchall()], None
    except Exception as exc:
        vector_error = str(exc)
    with _pg_conn() as conn:
        with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            fallback_sql = select_sql.replace("(c.embedding <=> %s::vector) as vector_distance", "null::float as vector_distance")
            cur.execute(fallback_sql + " order by c.created_at desc limit 1000", (category, category))
            return [dict(r) for r in cur.fetchall()], vector_error


def search_supabase_rag(question: str, limit: int = 5, category: Optional[str] = None) -> Dict[str, Any]:  # type: ignore[override]
    if not supabase_enabled():
        local = search_local_rag(question, limit=limit, category=category)
        # Local chunks already contain metadata after V81 uploads; expose them as sources.
        for item in local.get("results", []):
            item.setdefault("filename", (item.get("metadata") or {}).get("filename"))
            item.setdefault("metadata", item.get("metadata") or {})
        return {"enabled": False, **local, "mode": "local_keyword_v81"}
    try:
        rows, vector_error = _supabase_rows_for_search(question, limit, category)
        scored: List[Tuple[float, Dict[str, Any]]] = []
        for row in rows:
            lexical = keyword_score(question, (row.get("title") or "") + "\n" + (row.get("content") or ""))
            vd = row.get("vector_distance")
            semantic = 0.0 if vd is None else max(0.0, 1.0 - float(vd)) * 10.0
            status_boost = 0.5 if (row.get("status") or "active") in {"active", "valid", "current", "con_hieu_luc", "còn hiệu lực"} else 0.0
            score = round((semantic * 0.62) + (lexical * 0.35) + status_boost, 4)
            if score > 0 or vd is not None:
                scored.append((score, row))
        scored.sort(key=lambda x: x[0], reverse=True)
        results = [_source_from_row(row, score) for score, row in scored[: max(1, limit)]]
        return {
            "enabled": True,
            "mode": "hybrid_vector_keyword_v81",
            "embedding": embedding_provider(),
            "vector_error": vector_error,
            "question": question,
            "matched": len(results),
            "results": results,
            "answer": build_answer_from_chunks(question, results),
        }
    except Exception as exc:
        local = search_local_rag(question, limit=limit, category=category)
        local.update({"enabled": True, "supabase_error": str(exc), "fallback": "local_rag", "mode": "local_keyword_v81"})
        return local


def list_rag_documents(  # type: ignore[override]
    limit: int = 50,
    offset: int = 0,
    category: Optional[str] = None,
    source: Optional[str] = None,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    status: Optional[str] = None,
) -> Dict[str, Any]:
    if not supabase_enabled():
        docs = load_local_store().get("documents", [])
        for key, value in [("category", category), ("source", source), ("workspace_id", workspace_id), ("user_id", user_id), ("status", status)]:
            if value:
                docs = [d for d in docs if d.get(key) == value or (d.get("metadata") or {}).get(key) == value]
        return {"enabled": False, "total": len(docs), "documents": docs[offset: offset + limit]}
    try:
        import psycopg2.extras
        with _pg_conn() as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(
                    """
                    select d.id, d.title, d.filename, d.document_type, d.category, d.source, d.storage_path,
                           d.uploaded_by, d.document_number, d.issued_date, d.effective_date, d.authority,
                           d.status, d.version, d.workspace_id, d.user_id, d.language, d.jurisdiction,
                           d.tags, d.metadata, d.content_hash, d.created_at, d.updated_at,
                           count(c.id) as chunk_count
                    from rag_documents d
                    left join rag_chunks c on c.document_id = d.id
                    where (%s is null or d.category = %s)
                      and (%s is null or d.source = %s)
                      and (%s is null or d.workspace_id = %s)
                      and (%s is null or d.user_id = %s)
                      and (%s is null or d.status = %s)
                    group by d.id
                    order by d.created_at desc
                    limit %s offset %s
                    """,
                    (category, category, source, source, workspace_id, workspace_id, user_id, user_id, status, status, limit, offset),
                )
                docs = [dict(r) for r in cur.fetchall()]
                cur.execute(
                    """
                    select count(*) from rag_documents
                    where (%s is null or category = %s)
                      and (%s is null or source = %s)
                      and (%s is null or workspace_id = %s)
                      and (%s is null or user_id = %s)
                      and (%s is null or status = %s)
                    """,
                    (category, category, source, source, workspace_id, workspace_id, user_id, user_id, status, status),
                )
                total = int(cur.fetchone()[0])
        for d in docs:
            d["id"] = str(d["id"])
            for key in ["created_at", "updated_at", "issued_date", "effective_date"]:
                if d.get(key):
                    d[key] = d[key].isoformat()
        return {"enabled": True, "total": total, "documents": docs}
    except Exception as exc:
        return {"enabled": True, "error": str(exc), "documents": []}


def build_rag_answer(question: str, results: List[Dict[str, Any]], style: str = "detailed", use_llm: bool = False) -> Dict[str, Any]:  # type: ignore[override]
    citations = []
    for i, r in enumerate(results, start=1):
        citations.append({
            "ref": f"[{i}]",
            "document_id": r.get("document_id"),
            "chunk_id": r.get("chunk_id"),
            "title": r.get("title"),
            "filename": r.get("filename"),
            "document_number": r.get("document_number"),
            "authority": r.get("authority"),
            "issued_date": r.get("issued_date"),
            "effective_date": r.get("effective_date"),
            "status": r.get("status"),
            "page_start": r.get("page_start"),
            "page_end": r.get("page_end"),
            "chunk_index": r.get("chunk_index"),
            "score": r.get("score"),
        })
    context = "\n\n".join(
        f"[{i}] {r.get('title')} / {r.get('document_number') or ''} / chunk {r.get('chunk_index')}\n{(r.get('content') or '')[:2500]}"
        for i, r in enumerate(results[:8], start=1)
    )
    if use_llm and os.getenv("OPENAI_API_KEY") and results:
        try:
            from openai import OpenAI  # type: ignore
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            system = (
                "Bạn là trợ lý RAG cho tài liệu luật, thông tư, quy trình và nghiệp vụ. "
                "Chỉ dựa vào CONTEXT. Nếu thiếu căn cứ, nói rõ chưa đủ tài liệu. "
                "Trả lời tiếng Việt có cấu trúc: Kết luận, Căn cứ, Giải thích/quy trình, Lưu ý. "
                "Luôn trích nguồn dạng [1], [2]. Không bịa điều/khoản ngoài context."
            )
            user = f"CÂU HỎI:\n{question}\n\nCONTEXT:\n{context}"
            resp = client.chat.completions.create(
                model=DEFAULT_CHAT_MODEL,
                messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
                temperature=0.2,
            )
            return {"answer": resp.choices[0].message.content or "", "mode": "llm_grounded_v81", "citations": citations}
        except Exception as exc:
            llm_error = str(exc)
    else:
        llm_error = None
    if not results:
        return {
            "answer": "Chưa tìm thấy căn cứ phù hợp trong kho RAG. Hãy upload thêm tài liệu đúng category/workspace hoặc kiểm tra câu hỏi/từ khóa.",
            "mode": "extractive_no_match_v81",
            "citations": [],
        }
    lines = ["Kết luận tạm thời dựa trên tài liệu đã upload:"]
    for i, r in enumerate(results[:5], start=1):
        text = re.sub(r"\s+", " ", r.get("content") or "").strip()
        meta_bits = []
        if r.get("document_number"):
            meta_bits.append(str(r.get("document_number")))
        if r.get("authority"):
            meta_bits.append(str(r.get("authority")))
        if r.get("effective_date"):
            meta_bits.append(f"hiệu lực {r.get('effective_date')}")
        if r.get("status"):
            meta_bits.append(f"trạng thái {r.get('status')}")
        meta_text = " — " + "; ".join(meta_bits) if meta_bits else ""
        lines.append(f"\n[{i}] {r.get('title') or r.get('filename') or 'Tài liệu'}{meta_text} — chunk {r.get('chunk_index')}: {text[:900]}{'...' if len(text) > 900 else ''}")
    lines.append("\nCăn cứ: " + ", ".join(c["ref"] for c in citations[:5]))
    lines.append("Lưu ý: đây là câu trả lời trích xuất theo nguồn; dùng use_llm=true + OPENAI_API_KEY để diễn giải mạch lạc hơn.")
    if llm_error:
        lines.append(f"LLM fallback reason: {llm_error}")
    return {"answer": "\n".join(lines), "mode": "extractive_grounded_v81", "citations": citations}


def search_local_rag(question: str, limit: int = 5, category: Optional[str] = None) -> Dict[str, Any]:  # type: ignore[override]
    """V81 local JSON search with rich metadata in sources."""
    store = load_local_store()
    scored = []
    docs_by_id = {d.get("id"): d for d in store.get("documents", [])}
    for chunk in store.get("chunks", []):
        meta = chunk.get("metadata") or {}
        doc = docs_by_id.get(chunk.get("document_id"), {})
        if category and (meta.get("category") or doc.get("category")) != category:
            continue
        score = keyword_score(question, chunk.get("title", "") + "\n" + chunk.get("content", ""))
        if score > 0:
            scored.append((score, chunk, doc))
    scored.sort(key=lambda x: x[0], reverse=True)
    results = []
    for score, chunk, doc in scored[: max(1, limit)]:
        meta = chunk.get("metadata") or doc.get("metadata") or {}
        results.append({
            "score": score,
            "document_id": chunk.get("document_id"),
            "chunk_id": chunk.get("id"),
            "chunk_index": chunk.get("chunk_index"),
            "title": doc.get("title") or chunk.get("title"),
            "filename": doc.get("filename") or meta.get("filename"),
            "source": doc.get("source") or meta.get("source"),
            "category": doc.get("category") or meta.get("category"),
            "document_type": doc.get("document_type") or meta.get("document_type"),
            "document_number": doc.get("document_number") or meta.get("document_number"),
            "issued_date": doc.get("issued_date") or meta.get("issued_date"),
            "effective_date": doc.get("effective_date") or meta.get("effective_date"),
            "authority": doc.get("authority") or meta.get("authority"),
            "status": doc.get("status") or meta.get("status"),
            "version": doc.get("version") or meta.get("version"),
            "workspace_id": doc.get("workspace_id") or meta.get("workspace_id"),
            "user_id": doc.get("user_id") or meta.get("user_id"),
            "tags": doc.get("tags") or meta.get("tags") or [],
            "metadata": meta,
            "content": chunk.get("content"),
        })
    answer = build_answer_from_chunks(question, results)
    return {"question": question, "matched": len(results), "results": results, "answer": answer}

# ============================================================
# V82/V83 - Workspace isolation + stronger hybrid retrieval
# ============================================================
# V82: every answer/search/document-management call can be scoped by
#      workspace_id/user_id/status so different frontends/users do not mix data.
# V83: hybrid retrieval combines vector candidates + keyword candidates, then
#      reranks by semantic_score, keyword_score, exact-phrase boost and status.

DEFAULT_SEARCH_STATUS = os.getenv("RAG_DEFAULT_STATUS", "active")


def _scope_matches(meta: Dict[str, Any], doc: Dict[str, Any], *, workspace_id: Optional[str] = None, user_id: Optional[str] = None, status: Optional[str] = None) -> bool:
    if workspace_id and (doc.get("workspace_id") or meta.get("workspace_id")) != workspace_id:
        return False
    if user_id and (doc.get("user_id") or meta.get("user_id")) != user_id:
        return False
    if status and (doc.get("status") or meta.get("status") or "active") != status:
        return False
    return True


def _keyword_terms(question: str, max_terms: int = 8) -> List[str]:
    terms = []
    for word in tokenize(question):
        if word not in terms:
            terms.append(word)
        if len(terms) >= max_terms:
            break
    return terms


def _scope_where_sql(
    *,
    category: Optional[str] = None,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    status: Optional[str] = None,
) -> Tuple[str, List[Any]]:
    clauses = ["(%s is null or d.category = %s)", "(%s is null or d.workspace_id = %s)", "(%s is null or d.user_id = %s)", "(%s is null or d.status = %s)"]
    params: List[Any] = [category, category, workspace_id, workspace_id, user_id, user_id, status, status]
    return " and ".join(clauses), params


def _hybrid_score(
    question: str,
    row: Dict[str, Any],
    *,
    vector_weight: float = 0.60,
    keyword_weight: float = 0.40,
) -> Tuple[float, float, float]:
    text = (row.get("title") or "") + "\n" + (row.get("document_number") or "") + "\n" + (row.get("content") or "")
    lexical = float(keyword_score(question, text))
    vd = row.get("vector_distance")
    semantic = 0.0 if vd is None else max(0.0, 1.0 - float(vd)) * 10.0
    normalized_keyword = min(10.0, lexical)
    status = (row.get("status") or "active").lower()
    status_boost = 0.6 if status in {"active", "valid", "current", "con_hieu_luc", "còn hiệu lực"} else 0.0
    exact_boost = 0.0
    q_norm = normalize_text(question)
    c_norm = normalize_text(text)
    for phrase in re.findall(r"\w+(?:\s+\w+){1,5}", q_norm):
        if len(phrase) >= 8 and phrase in c_norm:
            exact_boost += 0.5
    score = round((semantic * vector_weight) + (normalized_keyword * keyword_weight) + status_boost + min(exact_boost, 2.0), 4)
    return score, semantic, lexical


def search_local_rag(  # type: ignore[override]
    question: str,
    limit: int = 5,
    category: Optional[str] = None,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    status: Optional[str] = None,
) -> Dict[str, Any]:
    """V82 local JSON search with category/workspace/user/status isolation."""
    store = load_local_store()
    scored: List[Tuple[float, Dict[str, Any], Dict[str, Any], float]] = []
    docs_by_id = {d.get("id"): d for d in store.get("documents", [])}
    for chunk in store.get("chunks", []):
        meta = chunk.get("metadata") or {}
        doc = docs_by_id.get(chunk.get("document_id"), {})
        if category and (meta.get("category") or doc.get("category")) != category:
            continue
        if not _scope_matches(meta, doc, workspace_id=workspace_id, user_id=user_id, status=status):
            continue
        lexical = float(keyword_score(question, chunk.get("title", "") + "\n" + chunk.get("content", "")))
        if lexical > 0:
            scored.append((lexical, chunk, doc, lexical))
    scored.sort(key=lambda x: x[0], reverse=True)
    results = []
    for score, chunk, doc, lexical in scored[: max(1, limit)]:
        meta = chunk.get("metadata") or doc.get("metadata") or {}
        results.append({
            "score": round(score, 4),
            "semantic_score": 0.0,
            "keyword_score": round(lexical, 4),
            "vector_distance": None,
            "document_id": chunk.get("document_id"),
            "chunk_id": chunk.get("id"),
            "chunk_index": chunk.get("chunk_index"),
            "title": doc.get("title") or chunk.get("title"),
            "filename": doc.get("filename") or meta.get("filename"),
            "source": doc.get("source") or meta.get("source"),
            "category": doc.get("category") or meta.get("category"),
            "document_type": doc.get("document_type") or meta.get("document_type"),
            "document_number": doc.get("document_number") or meta.get("document_number"),
            "issued_date": doc.get("issued_date") or meta.get("issued_date"),
            "effective_date": doc.get("effective_date") or meta.get("effective_date"),
            "authority": doc.get("authority") or meta.get("authority"),
            "status": doc.get("status") or meta.get("status"),
            "version": doc.get("version") or meta.get("version"),
            "workspace_id": doc.get("workspace_id") or meta.get("workspace_id"),
            "user_id": doc.get("user_id") or meta.get("user_id"),
            "tags": doc.get("tags") or meta.get("tags") or [],
            "metadata": meta,
            "content": chunk.get("content"),
        })
    return {
        "question": question,
        "matched": len(results),
        "results": results,
        "answer": build_answer_from_chunks(question, results),
        "scope": {"category": category, "workspace_id": workspace_id, "user_id": user_id, "status": status},
    }


def _supabase_rows_for_search(  # type: ignore[override]
    question: str,
    limit: int,
    category: Optional[str] = None,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    status: Optional[str] = None,
) -> Tuple[List[Dict[str, Any]], Optional[str]]:
    import psycopg2.extras
    vec = embedding_literal(question)
    where_sql, scope_params = _scope_where_sql(category=category, workspace_id=workspace_id, user_id=user_id, status=status)
    select_sql = f"""
        select c.id as chunk_id, c.document_id, c.chunk_index, c.content, c.metadata,
               c.page_start, c.page_end,
               d.title, d.filename, d.source, d.category, d.document_type, d.document_number,
               d.issued_date, d.effective_date, d.authority, d.status, d.version,
               d.workspace_id, d.user_id, d.tags,
               (c.embedding <=> %s::vector) as vector_distance
        from rag_chunks c
        join rag_documents d on d.id = c.document_id
        where {where_sql}
    """
    rows_by_chunk: Dict[str, Dict[str, Any]] = {}
    vector_error: Optional[str] = None
    try:
        with _pg_conn() as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(select_sql + " order by c.embedding <=> %s::vector limit %s", [vec] + scope_params + [vec, max(30, limit * 10)])
                for r in cur.fetchall():
                    row = dict(r)
                    rows_by_chunk[str(row.get("chunk_id"))] = row
    except Exception as exc:
        vector_error = str(exc)

    fallback_select = select_sql.replace("(c.embedding <=> %s::vector) as vector_distance", "null::float as vector_distance")
    terms = _keyword_terms(question)
    keyword_clauses = []
    keyword_params: List[Any] = []
    for term in terms:
        pattern = f"%{term}%"
        keyword_clauses.append("(c.content ilike %s or d.title ilike %s or coalesce(d.document_number, '') ilike %s)")
        keyword_params.extend([pattern, pattern, pattern])
    keyword_sql = " and (" + " or ".join(keyword_clauses) + ")" if keyword_clauses else ""
    with _pg_conn() as conn:
        with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(fallback_select + keyword_sql + " order by c.created_at desc limit %s", scope_params + keyword_params + [max(100, limit * 30)])
            for r in cur.fetchall():
                row = dict(r)
                rows_by_chunk.setdefault(str(row.get("chunk_id")), row)
            if len(rows_by_chunk) < max(20, limit * 5):
                cur.execute(fallback_select + " order by c.created_at desc limit %s", scope_params + [max(100, limit * 20)])
                for r in cur.fetchall():
                    row = dict(r)
                    rows_by_chunk.setdefault(str(row.get("chunk_id")), row)
    return list(rows_by_chunk.values()), vector_error


def search_supabase_rag(  # type: ignore[override]
    question: str,
    limit: int = 5,
    category: Optional[str] = None,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    status: Optional[str] = None,
    search_mode: str = "hybrid",
    vector_weight: float = 0.60,
    keyword_weight: float = 0.40,
) -> Dict[str, Any]:
    """V83 hybrid search. search_mode: hybrid | vector | keyword."""
    status_filter = status if status is not None else None
    if not supabase_enabled():
        local = search_local_rag(question, limit=limit, category=category, workspace_id=workspace_id, user_id=user_id, status=status_filter)
        for item in local.get("results", []):
            item.setdefault("filename", (item.get("metadata") or {}).get("filename"))
            item.setdefault("metadata", item.get("metadata") or {})
        return {"enabled": False, **local, "mode": "local_keyword_v82", "search_mode": search_mode}
    try:
        rows, vector_error = _supabase_rows_for_search(question, limit, category, workspace_id, user_id, status_filter)
        scored: List[Tuple[float, Dict[str, Any], float, float]] = []
        mode = (search_mode or "hybrid").lower()
        for row in rows:
            score, semantic, lexical = _hybrid_score(question, row, vector_weight=vector_weight, keyword_weight=keyword_weight)
            if mode == "keyword":
                score = round(min(10.0, lexical), 4)
            elif mode == "vector":
                score = round(semantic, 4)
            if score > 0 or row.get("vector_distance") is not None:
                scored.append((score, row, semantic, lexical))
        scored.sort(key=lambda x: x[0], reverse=True)
        results = []
        for score, row, semantic, lexical in scored[: max(1, limit)]:
            source = _source_from_row(row, score)
            source["semantic_score"] = round(semantic, 4)
            source["keyword_score"] = round(lexical, 4)
            results.append(source)
        return {
            "enabled": True,
            "mode": f"{mode}_workspace_scoped_v83",
            "embedding": embedding_provider(),
            "vector_error": vector_error,
            "question": question,
            "matched": len(results),
            "results": results,
            "answer": build_answer_from_chunks(question, results),
            "scope": {"category": category, "workspace_id": workspace_id, "user_id": user_id, "status": status_filter},
            "weights": {"vector": vector_weight, "keyword": keyword_weight},
        }
    except Exception as exc:
        local = search_local_rag(question, limit=limit, category=category, workspace_id=workspace_id, user_id=user_id, status=status_filter)
        local.update({"enabled": True, "supabase_error": str(exc), "fallback": "local_rag", "mode": "local_keyword_v82"})
        return local


def answer_rag_question(  # type: ignore[override]
    question: str,
    limit: int = 6,
    category: Optional[str] = None,
    style: str = "detailed",
    use_llm: bool = False,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    status: Optional[str] = None,
    search_mode: str = "hybrid",
    vector_weight: float = 0.60,
    keyword_weight: float = 0.40,
) -> Dict[str, Any]:
    search = search_supabase_rag(
        question,
        limit=limit,
        category=category,
        workspace_id=workspace_id,
        user_id=user_id,
        status=status,
        search_mode=search_mode,
        vector_weight=vector_weight,
        keyword_weight=keyword_weight,
    )
    results = search.get("results", [])
    answer = build_rag_answer(question, results, style=style, use_llm=use_llm)
    return {
        "question": question,
        "category": category,
        "workspace_id": workspace_id,
        "user_id": user_id,
        "status": status,
        "search_mode": search.get("mode"),
        "matched": search.get("matched", 0),
        "embedding": search.get("embedding") or embedding_provider(),
        "answer": answer.get("answer"),
        "answer_mode": answer.get("mode"),
        "citations": answer.get("citations", []),
        "sources": results,
        "scope": search.get("scope"),
        "weights": search.get("weights"),
        "warnings": [w for w in [search.get("supabase_error"), search.get("vector_error")] if w],
    }


def get_rag_document(  # type: ignore[override]
    document_id: str,
    include_chunks: bool = True,
    chunk_limit: int = 200,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
) -> Dict[str, Any]:
    if not supabase_enabled():
        store = load_local_store()
        doc = next((d for d in store.get("documents", []) if str(d.get("id")) == document_id), None)
        if doc and not _scope_matches(doc.get("metadata") or {}, doc, workspace_id=workspace_id, user_id=user_id):
            return {"enabled": False, "document": None, "chunks": [], "error": "Document nằm ngoài workspace/user scope"}
        chunks = [c for c in store.get("chunks", []) if str(c.get("document_id")) == document_id][:chunk_limit]
        return {"enabled": False, "document": doc, "chunks": chunks if include_chunks else []}
    try:
        import psycopg2.extras
        with _pg_conn() as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(
                    """
                    select * from rag_documents
                    where id = %s
                      and (%s is null or workspace_id = %s)
                      and (%s is null or user_id = %s)
                    """,
                    (document_id, workspace_id, workspace_id, user_id, user_id),
                )
                doc = cur.fetchone()
                if not doc:
                    return {"enabled": True, "document": None, "chunks": []}
                chunks: List[Dict[str, Any]] = []
                if include_chunks:
                    cur.execute(
                        "select id, chunk_index, content, metadata, page_start, page_end, token_count, created_at from rag_chunks where document_id = %s order by chunk_index limit %s",
                        (document_id, chunk_limit),
                    )
                    chunks = [dict(r) for r in cur.fetchall()]
        doc = dict(doc)
        doc["id"] = str(doc["id"])
        for key in ["created_at", "updated_at", "issued_date", "effective_date"]:
            if doc.get(key):
                doc[key] = doc[key].isoformat()
        for c in chunks:
            c["id"] = str(c["id"])
            if c.get("created_at"):
                c["created_at"] = c["created_at"].isoformat()
        return {"enabled": True, "document": doc, "chunks": chunks}
    except Exception as exc:
        return {"enabled": True, "error": str(exc), "document": None, "chunks": []}


def delete_rag_document(  # type: ignore[override]
    document_id: str,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
) -> Dict[str, Any]:
    if not supabase_enabled():
        store = load_local_store()
        doc = next((d for d in store.get("documents", []) if str(d.get("id")) == document_id), None)
        if doc and not _scope_matches(doc.get("metadata") or {}, doc, workspace_id=workspace_id, user_id=user_id):
            return {"enabled": False, "deleted": False, "error": "Document nằm ngoài workspace/user scope"}
        before_docs = len(store.get("documents", []))
        before_chunks = len(store.get("chunks", []))
        store["documents"] = [d for d in store.get("documents", []) if str(d.get("id")) != document_id]
        store["chunks"] = [c for c in store.get("chunks", []) if str(c.get("document_id")) != document_id]
        save_local_store(store)
        return {"enabled": False, "deleted_documents": before_docs - len(store["documents"]), "deleted_chunks": before_chunks - len(store["chunks"])}
    try:
        with _pg_conn() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    delete from rag_documents
                    where id = %s
                      and (%s is null or workspace_id = %s)
                      and (%s is null or user_id = %s)
                    """,
                    (document_id, workspace_id, workspace_id, user_id, user_id),
                )
                deleted = cur.rowcount
            conn.commit()
        return {"enabled": True, "deleted": deleted > 0, "deleted_documents": deleted, "scope": {"workspace_id": workspace_id, "user_id": user_id}}
    except Exception as exc:
        return {"enabled": True, "deleted": False, "error": str(exc)}


def reindex_rag_document(  # type: ignore[override]
    document_id: str,
    chunk_size: int = 1200,
    overlap: int = 180,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
) -> Dict[str, Any]:
    doc_data = get_rag_document(document_id, include_chunks=False, workspace_id=workspace_id, user_id=user_id)
    doc = doc_data.get("document")
    if not doc:
        return {"ok": False, "error": "Không tìm thấy document trong scope hiện tại"}
    storage_path = doc.get("storage_path") or ""
    path = Path(storage_path)
    if not path.exists():
        return {"ok": False, "error": f"Không tìm thấy file gốc để reindex: {storage_path}"}
    raw = path.read_bytes()
    text = read_upload_bytes(path.name, raw)
    chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap)
    if not chunks:
        return {"ok": False, "error": "File gốc không có nội dung để chunk"}
    if not supabase_enabled():
        return {"ok": False, "error": "Reindex theo document_id hiện chỉ hỗ trợ Supabase; local có thể upload lại file."}
    try:
        with _pg_conn() as conn:
            with conn.cursor() as cur:
                cur.execute("delete from rag_chunks where document_id = %s", (document_id,))
                provider = embedding_provider()
                for i, chunk in enumerate(chunks):
                    metadata = {
                        "filename": path.name,
                        "category": doc.get("category"),
                        "document_type": doc.get("document_type"),
                        "source": doc.get("source"),
                        "workspace_id": doc.get("workspace_id"),
                        "user_id": doc.get("user_id"),
                        "reindexed_at": now_iso(),
                        "embedding_provider": provider["provider"],
                        "embedding_model": provider["model"],
                        "token_count": len(tokenize(chunk)),
                    }
                    cur.execute(
                        """
                        insert into rag_chunks (document_id, chunk_index, content, metadata, embedding, token_count)
                        values (%s, %s, %s, %s::jsonb, %s::vector, %s)
                        """,
                        (document_id, i, chunk, json.dumps(metadata, ensure_ascii=False), embedding_literal((doc.get("title") or "") + "\n" + chunk), metadata["token_count"]),
                    )
                cur.execute("update rag_documents set updated_at = now() where id = %s", (document_id,))
            conn.commit()
        return {"ok": True, "document_id": document_id, "chunks": len(chunks), "embedding": embedding_provider(), "scope": {"workspace_id": workspace_id, "user_id": user_id}}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}

# ============================================================
# V84 - Backend API helpers for frontend document management
# ============================================================
# Frontend can call stable V84 endpoints for: upload/list/detail/status/update/delete/reindex.
# These helpers keep the frontend contract clean while reusing V81-V83 RAG internals.

EDITABLE_DOCUMENT_FIELDS = {
    "title", "filename", "document_type", "category", "source", "uploaded_by",
    "document_number", "issued_date", "effective_date", "authority", "status", "version",
    "workspace_id", "user_id", "language", "jurisdiction", "tags", "metadata",
}


def normalize_document_for_frontend(doc: Dict[str, Any]) -> Dict[str, Any]:
    """Small, predictable shape for frontend tables/cards."""
    meta = doc.get("metadata") or {}
    if isinstance(meta, str):
        try:
            meta = json.loads(meta)
        except Exception:
            meta = {}
    return {
        "id": str(doc.get("id")) if doc.get("id") is not None else None,
        "title": doc.get("title"),
        "filename": doc.get("filename") or meta.get("filename"),
        "document_type": doc.get("document_type") or meta.get("document_type"),
        "category": doc.get("category") or meta.get("category"),
        "source": doc.get("source") or meta.get("source"),
        "status": doc.get("status") or meta.get("status") or "active",
        "version": doc.get("version") or meta.get("version"),
        "workspace_id": doc.get("workspace_id") or meta.get("workspace_id"),
        "user_id": doc.get("user_id") or meta.get("user_id"),
        "document_number": doc.get("document_number") or meta.get("document_number"),
        "issued_date": doc.get("issued_date") or meta.get("issued_date"),
        "effective_date": doc.get("effective_date") or meta.get("effective_date"),
        "authority": doc.get("authority") or meta.get("authority"),
        "language": doc.get("language") or meta.get("language"),
        "jurisdiction": doc.get("jurisdiction") or meta.get("jurisdiction"),
        "tags": doc.get("tags") or meta.get("tags") or [],
        "chunk_count": int(doc.get("chunk_count") or 0),
        "content_hash": doc.get("content_hash"),
        "created_at": doc.get("created_at"),
        "updated_at": doc.get("updated_at"),
        "metadata": meta,
    }


def list_rag_documents_frontend(
    *,
    limit: int = 50,
    offset: int = 0,
    category: Optional[str] = None,
    source: Optional[str] = None,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    status: Optional[str] = None,
    q: Optional[str] = None,
) -> Dict[str, Any]:
    """V84 frontend-friendly document list with optional local search q."""
    data = list_rag_documents(
        limit=500 if q else limit,
        offset=0 if q else offset,
        category=category,
        source=source,
        workspace_id=workspace_id,
        user_id=user_id,
        status=status,
    )
    docs = [normalize_document_for_frontend(d) for d in data.get("documents", [])]
    if q:
        nq = normalize_text(q)
        docs = [
            d for d in docs
            if nq in normalize_text(" ".join(str(d.get(k) or "") for k in ["title", "filename", "document_number", "authority", "category", "source", "status"]))
        ]
        docs = docs[offset: offset + limit]
    return {
        "ok": not bool(data.get("error")),
        "version": "v84_frontend_document_api",
        "enabled": data.get("enabled"),
        "total": data.get("total", len(docs)) if not q else len(docs),
        "limit": limit,
        "offset": offset,
        "filters": {"category": category, "source": source, "workspace_id": workspace_id, "user_id": user_id, "status": status, "q": q},
        "documents": docs,
        "error": data.get("error"),
    }


def rag_document_status_summary(
    *,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    category: Optional[str] = None,
) -> Dict[str, Any]:
    """Counts for dashboard: statuses, categories, workspaces, total chunks."""
    if not supabase_enabled():
        store = load_local_store()
        docs = store.get("documents", [])
        filtered = []
        for d in docs:
            meta = d.get("metadata") or {}
            if workspace_id and (d.get("workspace_id") or meta.get("workspace_id")) != workspace_id:
                continue
            if user_id and (d.get("user_id") or meta.get("user_id")) != user_id:
                continue
            if category and (d.get("category") or meta.get("category")) != category:
                continue
            filtered.append(d)
        by_status: Dict[str, int] = {}
        by_category: Dict[str, int] = {}
        by_workspace: Dict[str, int] = {}
        for d in filtered:
            meta = d.get("metadata") or {}
            st = d.get("status") or meta.get("status") or "active"
            cat = d.get("category") or meta.get("category") or "general"
            ws = d.get("workspace_id") or meta.get("workspace_id") or "default"
            by_status[st] = by_status.get(st, 0) + 1
            by_category[cat] = by_category.get(cat, 0) + 1
            by_workspace[ws] = by_workspace.get(ws, 0) + 1
        return {
            "ok": True,
            "enabled": False,
            "total_documents": len(filtered),
            "total_chunks": sum(int(d.get("chunk_count") or 0) for d in filtered),
            "by_status": by_status,
            "by_category": by_category,
            "by_workspace": by_workspace,
            "embedding": embedding_provider(),
        }
    try:
        import psycopg2.extras
        with _pg_conn() as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                where = "where (%s is null or d.workspace_id = %s) and (%s is null or d.user_id = %s) and (%s is null or d.category = %s)"
                params = (workspace_id, workspace_id, user_id, user_id, category, category)
                cur.execute(f"select count(*) as total_documents from rag_documents d {where}", params)
                total_documents = int(cur.fetchone()["total_documents"])
                cur.execute(
                    f"select count(c.id) as total_chunks from rag_documents d left join rag_chunks c on c.document_id = d.id {where}",
                    params,
                )
                total_chunks = int(cur.fetchone()["total_chunks"])
                cur.execute(f"select coalesce(d.status,'active') as key, count(*) as value from rag_documents d {where} group by coalesce(d.status,'active')", params)
                by_status = {str(r["key"]): int(r["value"]) for r in cur.fetchall()}
                cur.execute(f"select coalesce(d.category,'general') as key, count(*) as value from rag_documents d {where} group by coalesce(d.category,'general')", params)
                by_category = {str(r["key"]): int(r["value"]) for r in cur.fetchall()}
                cur.execute(f"select coalesce(d.workspace_id,'default') as key, count(*) as value from rag_documents d {where} group by coalesce(d.workspace_id,'default')", params)
                by_workspace = {str(r["key"]): int(r["value"]) for r in cur.fetchall()}
        return {
            "ok": True,
            "enabled": True,
            "total_documents": total_documents,
            "total_chunks": total_chunks,
            "by_status": by_status,
            "by_category": by_category,
            "by_workspace": by_workspace,
            "embedding": embedding_provider(),
            "scope": {"workspace_id": workspace_id, "user_id": user_id, "category": category},
        }
    except Exception as exc:
        return {"ok": False, "enabled": True, "error": str(exc)}


def update_rag_document_metadata(
    document_id: str,
    updates: Dict[str, Any],
    *,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
) -> Dict[str, Any]:
    """Patch editable document metadata. Used by frontend edit modal."""
    clean = {k: v for k, v in (updates or {}).items() if k in EDITABLE_DOCUMENT_FIELDS}
    if not clean:
        return {"ok": False, "updated": False, "error": "Không có field hợp lệ để cập nhật", "allowed_fields": sorted(EDITABLE_DOCUMENT_FIELDS)}
    if not supabase_enabled():
        store = load_local_store()
        found = False
        for d in store.get("documents", []):
            if str(d.get("id")) != document_id:
                continue
            if not _scope_matches(d.get("metadata") or {}, d, workspace_id=workspace_id, user_id=user_id):
                return {"ok": False, "updated": False, "error": "Document nằm ngoài workspace/user scope"}
            found = True
            meta = d.setdefault("metadata", {})
            for k, v in clean.items():
                if k == "metadata" and isinstance(v, dict):
                    meta.update(v)
                else:
                    d[k] = v
                    meta[k] = v
            d["updated_at"] = now_iso()
        if found:
            save_local_store(store)
        return {"ok": found, "enabled": False, "updated": found, "document_id": document_id}
    try:
        import psycopg2.extras
        with _pg_conn() as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(
                    """
                    select metadata from rag_documents
                    where id = %s and (%s is null or workspace_id = %s) and (%s is null or user_id = %s)
                    """,
                    (document_id, workspace_id, workspace_id, user_id, user_id),
                )
                row = cur.fetchone()
                if not row:
                    return {"ok": False, "enabled": True, "updated": False, "error": "Không tìm thấy document trong scope hiện tại"}
                meta = row.get("metadata") or {}
                if isinstance(meta, str):
                    meta = json.loads(meta)
                set_parts = []
                params: List[Any] = []
                for k, v in clean.items():
                    if k == "metadata" and isinstance(v, dict):
                        meta.update(v)
                    elif k in {"tags"}:
                        set_parts.append(f"{k} = %s::jsonb")
                        params.append(json.dumps(v if isinstance(v, list) else split_tags(str(v)), ensure_ascii=False))
                        meta[k] = v if isinstance(v, list) else split_tags(str(v))
                    elif k in {"issued_date", "effective_date"}:
                        set_parts.append(f"{k} = nullif(%s, '')::date")
                        params.append(v or "")
                        meta[k] = v
                    elif k != "metadata":
                        set_parts.append(f"{k} = %s")
                        params.append(v)
                        meta[k] = v
                set_parts.append("metadata = %s::jsonb")
                params.append(json.dumps(meta, ensure_ascii=False))
                set_parts.append("updated_at = now()")
                sql = f"update rag_documents set {', '.join(set_parts)} where id = %s and (%s is null or workspace_id = %s) and (%s is null or user_id = %s)"
                params.extend([document_id, workspace_id, workspace_id, user_id, user_id])
                cur.execute(sql, params)
                updated = cur.rowcount
            conn.commit()
        return {"ok": updated > 0, "enabled": True, "updated": updated > 0, "document_id": document_id, "updated_fields": sorted(clean.keys())}
    except Exception as exc:
        return {"ok": False, "enabled": True, "updated": False, "error": str(exc)}


def set_rag_document_status(
    document_id: str,
    status: str,
    *,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
) -> Dict[str, Any]:
    allowed = {"active", "draft", "archived", "replaced", "inactive"}
    if status not in allowed:
        return {"ok": False, "error": f"status phải thuộc {sorted(allowed)}"}
    return update_rag_document_metadata(document_id, {"status": status}, workspace_id=workspace_id, user_id=user_id)
